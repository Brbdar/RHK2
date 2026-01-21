#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""RHK Befundassistent – Report Builder (split from rhk_app_web_master.py).

Enthält:
- Arztbericht, Patientenbericht, interner Bericht
- Input-Summary, JSON Export/Import

Hinweis: Inhalt ist weitgehend 1:1 aus der Master-Datei extrahiert.
"""

from __future__ import annotations

from functools import lru_cache

import json
import hashlib
import threading
from collections import OrderedDict

# =============================================================================
# Performance: bounded report caching (no functional changes)
# =============================================================================

REPORT_CACHE_MAXSIZE = 256

_report_cache_lock = threading.RLock()

# Each cache maps (kind, case_fingerprint) -> rendered string/dict
_report_cache: 'OrderedDict[tuple, object]' = OrderedDict()


def _case_fingerprint(case: dict) -> str:
    """Stable fingerprint for a case dict.

    Case is JSON-serializable by design (ui/derived/scores/decision/env/warnings/debug).
    We hash the sorted JSON to keep keys compact and avoid memory blowups.
    """
    try:
        js = json.dumps(case, ensure_ascii=False, sort_keys=True, separators=(',', ':'))
    except Exception:
        # Fallback: best-effort; should not happen in practice
        js = str(case)
    return hashlib.blake2b(js.encode('utf-8', errors='ignore'), digest_size=16).hexdigest()


def _cache_get(kind: str, fp: str):
    key = (kind, fp)
    with _report_cache_lock:
        if key in _report_cache:
            _report_cache.move_to_end(key)
            return _report_cache[key]
    return None


def _cache_set(kind: str, fp: str, value):
    key = (kind, fp)
    with _report_cache_lock:
        _report_cache[key] = value
        _report_cache.move_to_end(key)
        while len(_report_cache) > REPORT_CACHE_MAXSIZE:
            _report_cache.popitem(last=False)


from rhk_base import *  # noqa: F401,F403

# PH Therapieepisoden (restart-fähig)
from rhk_ph_tx import (  # noqa: F401
    parse_ph_tx_table_rows,
    legacy_lists_to_episodes,
    format_ph_tx_episode_line,
)

# Optional: local phrase/rule DB (DSGVO-safe). App must run without it.
try:
    from rhk_report_db import select_phrases  # type: ignore
except Exception:  # pragma: no cover
    select_phrases = None  # type: ignore

# Einige Render-Helpers liegen im Case-Modul (im Flat-Master waren sie vorher "weiter oben").
from rhk_case import build_render_ctx, render_p01_dynamic, filter_module_text  # noqa: F401

# Optional study pre-screen checks. App must run without it.
try:
    from rhk_study_checks import get_study_hints  # type: ignore
except Exception:  # pragma: no cover
    def get_study_hints(case):  # type: ignore
        return []
# ---------------------------------------------------------------------------
# Small, conservative post-filters for narrative blocks
# ---------------------------------------------------------------------------

def _no_congestion_context(ui: Dict[str, Any], der: Dict[str, Any]) -> bool:
    """Return True if available inputs support "no (central or pulmonary venous) congestion"."""
    congestion_likely = bool((der or {}).get("congestion_likely"))
    pawp = _safe_float((ui or {}).get("pawp_rest"))
    pv_stauung_likely = bool(pawp is not None and pawp > 15)
    return (not congestion_likely) and (not pv_stauung_likely)


def _filter_narrative_block(text: str, ui: Dict[str, Any], der: Dict[str, Any]) -> str:
    """Conservatively removes known contradictory stock phrases from narrative bundles.

    IMPORTANT: This must never invent content; it only removes text that is internally
    contradictory to the current case (e.g., "Bei führender Stauung" while no congestion).
    """
    t = str(text or "")
    if not t.strip():
        return ""

    # If no congestion is likely, remove the optional clause that proposes congestion-first
    # management, because it contradicts the explicitly stated absence of congestion.
    if _no_congestion_context(ui, der):
        if "Bei führender Stauung" in t:
            # Remove the sentence segment starting at "Bei führender Stauung:" up to the next period.
            import re
            t = re.sub(r"\s*Bei führender Stauung:\s*[^.]*\.?\s*", " ", t, flags=re.IGNORECASE)

    # Reduce redundant one-liners that often duplicate earlier context.
    t = t.replace("Aktuell RHK in Ruhe.", "").replace("Aktuell RHK in Ruhe", "").strip()

    # In-house center reports must not recommend referral to itself.
    # Remove any sentences that suggest discussion/presentation in an external expert board/center.
    import re
    t = re.sub(r"\s*Diskussion\s+im\s+Expert[^.]*\.\s*", " ", t, flags=re.IGNORECASE)
    t = re.sub(r"\s*Diskussion\s+im\s+PH-Board[^.]*\.\s*", " ", t, flags=re.IGNORECASE)
    t = re.sub(r"\s*Vorstellung\s+im\s+PH-Zentrum[^.]*\.\s*", " ", t, flags=re.IGNORECASE)
    t = re.sub(r"\s*Vorstellung\s+im\s+Referenzzentrum[^.]*\.\s*", " ", t, flags=re.IGNORECASE)
    # Avoid overly broad regexes that could remove unrelated content.

    # Clean multiple spaces introduced by removals.
    t = " ".join(t.split())
    return t



# =============================================================================
# Warnings / Hinweise – formatting helpers
# =============================================================================

def _format_warning_item(w: Any) -> str:
    """Render warning/hint items robustly.

    Some logic paths attach structured dicts like:
    {"code": ..., "severity": ..., "message": ..., "fields": ..., "values": ...}
    We must never leak raw dict representations into the UI/report.
    """
    if w is None:
        return ""
    if isinstance(w, dict):
        msg = str((w.get("message") or "")).strip()
        if not msg:
            # fallback to code if message missing
            msg = str((w.get("code") or "")).strip()
        # Optional tiny suffix if we have a key value pair and it's non sensitive
        try:
            vals = w.get("values") or {}
            if isinstance(vals, dict) and vals:
                # show at most one numeric value in a compact, human-readable form
                k0 = next(iter(vals.keys()))
                v0 = vals.get(k0)
                if isinstance(v0, (int, float)):
                    key_map = {
                        "pasp_echo": "sPAP (Echo)",
                        "trv_ms": "TRV",
                        "rap_rest": "RAP",
                        "pawp_rest": "PAWP",
                        "mpap_rest": "mPAP",
                        "co_rest": "CO",
                        "ci_rest": "CI",
                    }
                    lab = key_map.get(str(k0), str(k0))
                    # avoid clutter for unknown internal keys
                    if lab:
                        msg = f"{msg} ({lab}: {fmt_float(v0, 1)})"
        except Exception:
            pass
        return msg
    return str(w).strip()

# =============================================================================
# Befund – input summary block
# =============================================================================


def _md_kv(label: str, value: str) -> str:
    return f"- **{label}:** {value}"


def _md_section(title: str, lines: List[str], *, add_colon: bool = False) -> str:
    """Small helper to build a section from a list of list-items.

    add_colon=True yields plain-text headers like 'Klinik:' (used for Arztbericht summary blocks).
    """
    if not lines:
        return ""
    if add_colon:
        return f"{title}:\n" + "\n".join(lines)
    return f"### {title}\n" + "\n".join(lines)



def _build_risk_lines(case: Dict[str, Any]) -> List[str]:
    """Build risk stratification lines (doctor-facing) as markdown list items."""
    sc = case.get("scores") or {}
    der = case.get("derived") or {}
    lines: List[str] = []
    if sc.get("esc_ers_4s"):
        lines.append(_md_kv("ESC/ERS 4-Strata", str(sc["esc_ers_4s"])))
    if sc.get("esc_ers_3s"):
        lines.append(_md_kv("ESC/ERS 3-Strata", str(sc["esc_ers_3s"])))
    if sc.get("reveal_lite2"):
        cat = sc.get("reveal_lite2")
        pts = sc.get("reveal_lite2_points")
        if cat == "nicht berechenbar":
            missing = sc.get("reveal_lite2_missing") or []
            miss_txt = ", ".join(missing) if missing else "Parameter unvollständig"
            lines.append(_md_kv("REVEAL Lite 2", f"nicht berechenbar (fehlend: {miss_txt})"))
        else:
            cat_de = {"low": "niedrig", "intermediate": "intermediär", "high": "hoch"}.get(str(cat), str(cat))
            pts_txt = str(pts) if pts is not None else "—"
            lines.append(_md_kv("REVEAL Lite 2", f"{pts_txt} Punkte ({cat_de})"))
    if der.get("hfpef_category"):
        lines.append(_md_kv("HFpEF (H2FPEF)", f"{der['hfpef_category']} (~{_fmt(der.get('hfpef_percent'),0)}%)"))
    return lines


def _strip_procedere_from_text(text: str) -> str:
    """Remove procedural / repetitive paragraphs from a narrative block used outside Procedere."""
    if not isinstance(text, str):
        return ""
    raw = text.strip()
    if not raw:
        return ""
    # Split by blank lines (markdown paragraphs)
    paras = [p.strip() for p in re.split(r"\n\s*\n", raw) if p.strip()]
    keep: List[str] = []
    bad_kw = [
        "Empfohlen:", "Empfohlen", "Procedere", "Vorstellung", "Board", "Mitbeurteilung",
        "Therapieoptimierung", "Abklärung", "PH-Zentrum", "Autoimmun", "HIV", "Infektiologie",
        "Genetik", "angeborene", "HRCT", "V/Q", "Pulmonalisangiographie", "Therapie",
        "Diagnose/Einordnung", "ESC/ERS", "REVEAL",
    ]
    for p in paras:
        if any(k.lower() in p.lower() for k in bad_kw):
            continue
        keep.append(p)
    out = "\n\n".join(keep).strip()
    return out


def _sanitize_concluding(text: str) -> str:
    """Keep etiological conclusion, but strip any recommendation-like content.

    For the Arztbericht interpretation block we must NOT output:
    - "Zusatzhinweis:" paragraphs
    - Any explicit recommendations ("Empfohlen", "Mitbeurteilung", "Therapieplanung",
      "PH-Zentrum", "Abklärung gemäß Leitlinie", etc.)

    The dedicated "Procedere" section is the only place where recommendations belong.
    """
    if not isinstance(text, str):
        return ""
    t = (text or "").strip()
    if not t:
        return ""

    # 1) Drop any paragraph that starts with "Zusatzhinweis:" (case-insensitive)
    paras = [p.strip() for p in re.split(r"\n\s*\n", t) if p.strip()]
    cleaned_paras: List[str] = []
    for p in paras:
        if re.match(r"^zusatzhinweis\s*:\s*", p.strip(), flags=re.IGNORECASE):
            continue
        cleaned_paras.append(p)
    t = "\n\n".join(cleaned_paras).strip()
    if not t:
        return ""

    # 2) Remove sentences/lines that contain recommendation cues.
    # Keep this conservative: only remove when clearly procedural.
    rec_markers = [
        "empfehl",          # Empfohlen/Empfehlung
        "ph-zentrum",       # PH-Zentrum
        "mitbeurteilung",   # Mitbeurteilung
        "therapieplanung",  # Therapieplanung
        "therapie nach",    # Therapie nach Risikoprofil
        "strukturierte komplettierung",  # list-like diagnostic completion
        "abklärung gemäß",  # guideline wording
        "leitlinie",        # guideline references in procedural form
        "risikoadaptiert",  # risk-adapted therapy planning
        "komplette",        # "komplette ... Abklärung" (procedural)
        "optimierung/abklärung",
    ]

    # Work line-wise first to catch "Empfohlen:" and "PH-Zentrum:" fragments.
    lines = [ln.strip() for ln in t.splitlines()]
    kept_lines: List[str] = []
    for ln in lines:
        s = ln.lower()
        if any(m in s for m in rec_markers):
            continue
        kept_lines.append(ln)
    t = " ".join([ln for ln in kept_lines if ln]).strip()
    if not t:
        return ""

    # Finally, sentence split and filter again (covers inline "Empfohlen:" within long sentences)
    sents = re.split(r"(?<=[.!?])\s+", t)
    keep_sents = [s for s in sents if not any(m in s.lower() for m in rec_markers)]
    return " ".join([s.strip() for s in keep_sents if s.strip()]).strip()


def _sanitize_interpretation_block(text: str) -> str:
    """Final safety net for the Arztbericht interpretation block.

    Goal:
    - Keep *diagnostic/etiology interpretation* sentences.
    - Reliably remove any *procedural recommendations* (Empfohlen/PH-Zentrum/etc.)
      even if they appear inline within the same paragraph.
    - Drop any dedicated 'Zusatzhinweis:' paragraph.
    """
    if not isinstance(text, str):
        return ""
    t = (text or "").strip()
    if not t:
        return ""

    # Procedural markers (lowercase matching). Keep conservative and German-centric.
    rec_markers = [
        "empfehl",               # Empfohlen/Empfehlung
        "ph-zentrum",            # PH-Zentrum
        "mitbeurteilung",        # Mitbeurteilung
        "therapieplanung",       # Therapieplanung
        "therapie nach",         # Therapie nach Risikoprofil
        "abklärung gemäß",       # guideline procedural wording
        "leitlinie",             # guideline references in procedural form
        "risikoadaptiert",       # risk-adapted therapy planning
        "komplette",             # "komplette ... Abklärung"
        "optimierung/abklärung",
        "interdisziplin",        # interdisziplinäre Einordnung/Board (often procedural)
        "diagnostik werden empfohlen",
        "werden empfohlen",
        "empfohlen wird",
    ]

    def _truncate_at_marker(sentence: str) -> str:
        """If a procedural marker appears, keep the part before it (if meaningful)."""
        s = (sentence or "").strip()
        if not s:
            return ""
        sl = s.lower()
        # Find earliest marker occurrence
        cut_pos = None
        for mk in rec_markers:
            p = sl.find(mk)
            if p != -1:
                if cut_pos is None or p < cut_pos:
                    cut_pos = p
        if cut_pos is None:
            return s

        # Try to cut at the marker boundary, but also remove preceding separators like ';' or '–'
        prefix = s[:cut_pos].rstrip(" ;,–-:\n\t")
        # If prefix is too short / non-informative, drop entirely
        if len(prefix) < 20:
            return ""
        return prefix

    # 1) Drop any paragraph that starts with 'Zusatzhinweis:' (case-insensitive)
    paras = [p.strip() for p in re.split(r"\n\s*\n", t) if p.strip()]
    paras = [p for p in paras if not re.match(r"^zusatzhinweis\s*:\s*", p, flags=re.IGNORECASE)]

    kept_sents: List[str] = []
    for p in paras:
        # Split sentences (simple but robust enough for our German blocks)
        sents = re.split(r"(?<=[.!?])\s+", p.strip())
        for s in sents:
            cleaned = _truncate_at_marker(s)
            if cleaned:
                kept_sents.append(cleaned)

    # 2) Re-join into paragraphs (single block is fine)
    out = " ".join([s.strip() for s in kept_sents if s.strip()]).strip()
    # 3) Final cleanup: collapse whitespace
    out = re.sub(r"\s+", " ", out).strip()
    return out
def _build_relevante_vorerkrankungen_line(ui: Dict[str, Any]) -> str:
    """Build a single-line 'Relevante Vorerkrankungen' string for the Arztbericht.

    Includes ONLY items explicitly captured as relevant comorbidities in the UI:
    - Freitext 'comorbidities'
    - Virologie/Infektiologie (e.g., HIV/Hepatitis) when marked positive
    - Immunologie/Autoimmun when marked positive
    - Angeborener Herzfehler/Shunt when marked positive
    """
    items: List[str] = []
    comorb = (ui.get("comorbidities") or "").strip()
    if comorb:
        items.append(comorb)

    # CHD/Shunt
    if ui.get("chd_pos") is True:
        chd_type = (ui.get("chd_type") or "").strip()
        chd_desc = (ui.get("chd_desc") or "").strip()
        txt = "Angeborener Herzfehler/Shunt"
        bits = []
        if chd_type:
            bits.append(chd_type)
        if chd_desc:
            bits.append(chd_desc)
        if bits:
            txt += f" ({' – '.join(bits)})"
        items.append(txt)

    # Virology
    if ui.get("virology_pos") is True:
        v_items = ui.get("virology_items")
        v_desc = (ui.get("virology_desc") or "").strip()
        parts: List[str] = []
        if isinstance(v_items, list) and v_items:
            parts.extend([str(x).strip() for x in v_items if str(x).strip()])
        elif isinstance(v_items, str) and v_items.strip():
            parts.append(v_items.strip())
        if v_desc:
            parts.append(v_desc)
        items.append("Virologie/Infektiologie: " + (", ".join(parts) if parts else "positiv"))

    # Immunology
    if ui.get("immunology_pos") is True:
        i_items = ui.get("immunology_items")
        i_desc = (ui.get("immunology_desc") or "").strip()
        parts = []
        if isinstance(i_items, list) and i_items:
            parts.extend([str(x).strip() for x in i_items if str(x).strip()])
        elif isinstance(i_items, str) and i_items.strip():
            parts.append(i_items.strip())
        if i_desc:
            parts.append(i_desc)
        items.append("Immunologie/Autoimmun: " + (", ".join(parts) if parts else "positiv"))

    joined = "; ".join([x for x in items if x])
    return joined if joined else "-"



def _get_ph_tx_episodes(ui: Dict[str, Any], derived: Optional[Dict[str, Any]] = None) -> List[Dict[str, str]]:
    """Return PH therapy episodes (prefer derived, else parse UI table, else legacy lists)."""
    der = derived or {}
    if isinstance(der.get("ph_tx_episodes"), list):
        eps = [e for e in (der.get("ph_tx_episodes") or []) if isinstance(e, dict)]
        if eps:
            return eps
    rows = ui.get("ph_tx_table")
    eps = parse_ph_tx_table_rows(rows)
    if eps:
        return eps
    return legacy_lists_to_episodes(ui)


def _build_ph_therapieverlauf_block(ui: Dict[str, Any], derived: Optional[Dict[str, Any]] = None) -> str:
    """Build 'PH-Therapieverlauf' block for Arztbericht.

    Deterministic ordering:
    - Historie (früher/abgesetzt/pausiert)
    - Aktuell
    - Geplant
    """
    eps = _get_ph_tx_episodes(ui, derived)
    if not eps:
        return ""

    def _collect(statuses: set[str]) -> List[str]:
        out: List[str] = []
        for e in eps:
            try:
                st = str(e.get("status") or "").strip().lower()
            except Exception:
                st = ""
            if st not in statuses:
                continue
            s = format_ph_tx_episode_line(e)
            if s and s not in out:
                out.append(s)
        return out

    lines: List[str] = []
    hist = _collect({"früher", "abgesetzt", "pausiert"})
    cur = _collect({"aktuell"})
    planned = _collect({"geplant"})

    if hist:
        lines.append(_md_kv("Historie", ", ".join(hist)))
    if cur:
        lines.append(_md_kv("Aktuell", ", ".join(cur)))
    if planned:
        lines.append(_md_kv("Geplant", ", ".join(planned)))

    return "PH-Therapieverlauf:\n" + "\n".join([l for l in lines if l]) + "\n"

def summarize_inputs(case: Dict[str, Any], *, mode: str = "default") -> str:
    """Creates a compact, structured overview of the raw input data (Markdown)."""
    ui = case.get("ui") or {}
    env = case.get("env") or {}
    der = case.get("derived") or {}

    parts: List[str] = []

    is_doctor = (mode == "doctor")

    # ---------------------------------------------------------------------
    # Klinik
    # ---------------------------------------------------------------------
    klinik_lines: List[str] = []
    story = (ui.get("story") or "").strip()
    if story:
        klinik_lines.append(_md_kv("Kurz-Anamnese", story))

    comorb = (ui.get("comorbidities") or "").strip()
    if comorb and (not is_doctor):
        klinik_lines.append(_md_kv("Relevante Vorerkrankungen", comorb))

    # Angeborener Herzfehler / Shunt (DD Gruppe 1)
    if (not is_doctor) and ui.get("chd_pos") is True:
        chd_type = ui.get("chd_type")
        chd_desc = (ui.get("chd_desc") or "").strip()
        txt = "ja"
        if chd_type:
            txt += f" ({chd_type})"
        if chd_desc:
            txt += f" – {chd_desc}"
        klinik_lines.append(_md_kv("Angeborener Herzfehler/Shunt", txt))

    # Virologie/Infektiologie (z.B. HIV; DD Gruppe 1)
    if (not is_doctor) and ui.get("virology_pos") is True:
        items = ui.get("virology_items")
        desc = (ui.get("virology_desc") or "").strip()
        parts: List[str] = []
        if isinstance(items, list) and items:
            parts.append(", ".join([str(x) for x in items if str(x).strip()]))
        elif isinstance(items, str) and items.strip():
            parts.append(items.strip())
        if desc:
            parts.append(desc)
        klinik_lines.append(_md_kv("Virologie/Infektiologie", " / ".join(parts) if parts else "positiv"))

    # Immunologie/Autoimmun (z.B. CTD; DD Gruppe 1)
    if (not is_doctor) and ui.get("immunology_pos") is True:
        items = ui.get("immunology_items")
        desc = (ui.get("immunology_desc") or "").strip()
        parts = []
        if isinstance(items, list) and items:
            parts.append(", ".join([str(x) for x in items if str(x).strip()]))
        elif isinstance(items, str) and items.strip():
            parts.append(items.strip())
        if desc:
            parts.append(desc)
        klinik_lines.append(_md_kv("Immunologie/Autoimmun", " / ".join(parts) if parts else "positiv"))

    # Genetik/Mutation (DD Gruppe 1)
    if ui.get("mutation_pos") is True:
        items = ui.get("mutation_items")
        desc = (ui.get("mutation_desc") or "").strip()
        parts = []
        if isinstance(items, list) and items:
            parts.append(", ".join([str(x) for x in items if str(x).strip()]))
        elif isinstance(items, str) and items.strip():
            parts.append(items.strip())
        if desc:
            parts.append(desc)
        klinik_lines.append(_md_kv("Genetik/Mutation", " / ".join(parts) if parts else "positiv"))

    if (not is_doctor) and ui.get("ph_known") is True:
        klinik_lines.append(_md_kv("PH-Diagnose", "bekannt"))

        # Details zur bekannten PH (falls angegeben)
        dx = (ui.get("ph_known_dx") or "").strip()
        if dx:
            klinik_lines.append(_md_kv("Bekannte PH-Diagnose", dx))

        first_dx = (ui.get("ph_first_dx") or "").strip()
        if first_dx:
            klinik_lines.append(_md_kv("Erstdiagnose", first_dx))

        reason = (ui.get("ph_reason_rhk") or "").strip()
        if reason:
            klinik_lines.append(_md_kv("Aktueller Anlass", reason))

        subtype = (ui.get("ph_known_subtype") or "").strip()
        if subtype:
            klinik_lines.append(_md_kv("Subtyp/Kontext", subtype))

        eps = _get_ph_tx_episodes(ui, der)
        if eps:
            cur = [format_ph_tx_episode_line(e) for e in eps if str(e.get("status") or "").strip().lower() == "aktuell"]
            hist = [format_ph_tx_episode_line(e) for e in eps if str(e.get("status") or "").strip().lower() in ("früher", "abgesetzt", "pausiert")]
            planned = [format_ph_tx_episode_line(e) for e in eps if str(e.get("status") or "").strip().lower() == "geplant"]
            if hist:
                klinik_lines.append(_md_kv("PH Therapie Historie", ", ".join([x for x in hist if x])))
            if cur:
                klinik_lines.append(_md_kv("PH Therapie aktuell", ", ".join([x for x in cur if x])))
            if planned:
                klinik_lines.append(_md_kv("PH Therapie geplant", ", ".join([x for x in planned if x])))

        interv = ui.get("ph_interventions") or []
        if isinstance(interv, list) and interv:
            klinik_lines.append(_md_kv("Interventionen", ", ".join([str(x) for x in interv])))

    elif (not is_doctor) and ui.get("ph_suspected") is True:
        klinik_lines.append(_md_kv("PH-Verdachtsdiagnose", "ja"))

    # Vitals
    sbp = _safe_float(ui.get("bp_sys"))
    dbp = _safe_float(ui.get("bp_dia"))
    hr = _safe_float(ui.get("hr"))
    if sbp is not None or dbp is not None:
        if sbp is not None and dbp is not None:
            klinik_lines.append(_md_kv("Blutdruck", f"{fmt_int(sbp)}/{fmt_int(dbp)} mmHg"))
        elif sbp is not None:
            klinik_lines.append(_md_kv("Blutdruck", f"{fmt_int(sbp)} mmHg"))
        else:
            klinik_lines.append(_md_kv("Blutdruck", f"{fmt_int(dbp)} mmHg"))
    if hr is not None:
        klinik_lines.append(_md_kv("Herzfrequenz", f"{fmt_int(hr)}/min"))
    # EKG
    if ui.get("ekg_present") is True:
        signs = ui.get("ekg_rhs_signs") or []
        if isinstance(signs, list) and signs:
            other = (ui.get("ekg_other_text") or "").strip()
            items = [str(x) for x in signs if str(x).strip() and str(x).strip().lower() != "sonstiges/unklar"]
            if other:
                items.append(other)
            if items:
                klinik_lines.append(_md_kv("EKG Rechtsherzbelastungszeichen", ", ".join(items)))
    # LSB
    if ui.get("lsb_present") is True:
        klinik_lines.append(_md_kv("LSB", "ja"))

        lsb_reason = (ui.get("lsb_reason") or "").strip()
        if lsb_reason:
            klinik_lines.append(_md_kv("LSB Begründung", lsb_reason))
    # Nitrate/NO-Donor
    if ui.get("on_nitrates") is True:
        klinik_lines.append(_md_kv("Nitrate/NO-Donor", "ja"))

    # PDE-5 Härtefall (Dokumentation)
    if ui.get("pde5_hardship") is True:
        desc = (ui.get("pde5_hardship_desc") or "").strip()
        klinik_lines.append(_md_kv("PDE-5 Härtefall", desc if desc else "ja"))
    # Symptome / Funktion
    if ui.get("exertional_dyspnea") is True:
        klinik_lines.append(_md_kv("Belastungsdyspnoe", "ja"))

    syn = ui.get("syncope")
    syn_s: Optional[str] = None
    if isinstance(syn, bool):
        syn_s = "ja" if syn else None
    else:
        tmp = (syn or "").strip()
        if tmp and tmp.lower() not in ("keine", "nein"):
            syn_s = tmp
    if syn_s:
        klinik_lines.append(_md_kv("Synkope", syn_s))

    if ui.get("hemoptysis") is True:
        klinik_lines.append(_md_kv("Hämoptyse", "ja"))
    if ui.get("dizziness") is True:
        klinik_lines.append(_md_kv("Schwindel", "ja"))

    stairs = ui.get("stairs_flights")
    if stairs not in (None, "", 0):
        klinik_lines.append(_md_kv("Treppenstufen/Etagen (Alltag)", str(stairs)))

    who_fc = ui.get("who_fc")
    if who_fc:
        klinik_lines.append(_md_kv("WHO-FC", str(who_fc)))
    six = _safe_float(ui.get("six_mwd_m"))
    if six is not None:
        six_dt = (ui.get("six_mwd_date") or "").strip()
        if six_dt:
            klinik_lines.append(_md_kv("6MWD", f"{_fmt(six,0)} m (Datum: {six_dt})"))
        else:
            klinik_lines.append(_md_kv("6MWD", f"{_fmt(six,0)} m"))



    # Medikation / Zusatzangaben (falls erfasst)
    anticoag_status = (ui.get("anticoag_status") or "").strip()
    # "keine Angabe" darf niemals als Fakt in den Bericht geraten.
    if anticoag_status and anticoag_status.lower() not in ("keine angabe", "k. a.") and ((not is_doctor) or anticoag_status.lower() == "ja"):
        msg = anticoag_status
        if anticoag_status.lower() == "ja":
            bits: List[str] = []
            sub = (ui.get("anticoag_substance") or "").strip()
            ind = (ui.get("anticoag_indication") or "").strip()
            since = (ui.get("anticoag_since") or "").strip()
            if sub and sub.lower() not in ("keine angabe", "k. a."):
                bits.append(sub)
            if ind and ind.lower() not in ("keine angabe", "k. a."):
                bits.append(f"Indikation: {ind}")
            if since:
                bits.append(f"seit {since}")
            if bits:
                msg += " (" + "; ".join(bits) + ")"
        klinik_lines.append(_md_kv("Antikoagulation", msg))

    note = (ui.get("anticoag_note") or "").strip()
    if note and anticoag_status.lower() in ("ja", "nein") and (not is_doctor):
        klinik_lines.append(_md_kv("Antikoagulation – Bem.", note))

    antif_status = (ui.get("antifibrotic_status") or "").strip()
    if antif_status and antif_status.lower() not in ("keine angabe", "k. a."):
        msg = antif_status
        if antif_status.lower() == "ja":
            bits: List[str] = []
            drug = (ui.get("antifibrotic_drug") or "").strip()
            since = (ui.get("antifibrotic_since") or "").strip()
            if drug and drug.lower() not in ("keine angabe", "k. a."):
                bits.append(drug)
            if since:
                bits.append(f"seit {since}")
            if bits:
                msg += " (" + "; ".join(bits) + ")"
        klinik_lines.append(_md_kv("Antifibrotische Therapie", msg))

    antif_note = (ui.get("antifibrotic_note") or "").strip()
    if antif_note and antif_status.lower() in ("ja", "nein"):
        klinik_lines.append(_md_kv("Antifibrotika – Bem.", antif_note))

    ltx = (ui.get("ltx_eval") or "").strip()
    if ltx and ltx.lower() not in ("keine angabe", "k. a."):
        extra = ""
        ltx_date = (ui.get("ltx_eval_date") or "").strip()
        if ltx_date:
            extra = f" (Datum: {ltx_date})"
        klinik_lines.append(_md_kv("LTX-Evaluation", f"{ltx}{extra}"))

    if not klinik_lines:
        klinik_lines.append("Keine klinischen Angaben erfasst.")

    if is_doctor:
        _rl = _build_risk_lines(case)
        if _rl:
            klinik_lines.append("- **Risikostratifizierung:**")
            klinik_lines.extend(_rl)

    parts.append(_md_section("Klinik", klinik_lines, add_colon=is_doctor))

    # ---------------------------------------------------------------------
    # Labor (Fließtext; BNP/NT-proBNP separat)
    # ---------------------------------------------------------------------
    lab_items: List[str] = []
    lab_tail_lines: List[str] = []

    hb = _safe_float(ui.get("hb_g_dl"))
    if hb is not None:
        lab_items.append(f"Hb: {_fmt(hb,1)} g/dl" + (" (Anämie)" if der.get("anemia") else ""))

    crp = _safe_float(ui.get("crp_mg_l"))
    if crp is not None:
        lab_items.append(f"CRP: {_fmt(crp,1)} mg/l")

    crea = _safe_float(ui.get("creatinine_mg_dl"))
    if crea is not None:
        lab_items.append(f"Kreatinin: {_fmt(crea,2)} mg/dl")

    egfr = _safe_float(ui.get("egfr"))
    if egfr is not None:
        lab_items.append(f"eGFR: {fmt_int(egfr)} ml/min/1,73m²")

    inr = _safe_float(ui.get("inr"))
    if inr is not None:
        lab_items.append(f"INR: {_fmt(inr,2)}")

    ptt = _safe_float(ui.get("ptt_s"))
    if ptt is not None:
        lab_items.append(f"PTT: {_fmt(ptt,0)} s")

    thr = _safe_float(ui.get("platelets_g_l"))
    if thr is not None:
        lab_items.append(f"Thrombozyten: {_fmt(thr,0)} G/l")

    leuk = _safe_float(ui.get("leukocytes_g_l"))
    if leuk is not None:
        lab_items.append(f"Leukozyten: {_fmt(leuk,1)} G/l")

    # BNP/NT-proBNP bewusst separat unter dem Fließtext
    bnp_kind = ui.get("bnp_kind")
    bnp_val = _safe_float(ui.get("bnp_value"))
    if bnp_val is not None:
        extra = ""
        if ui.get("entresto") is True and isinstance(bnp_kind, str) and "BNP" in bnp_kind.upper() and "NT" not in bnp_kind.upper():
            extra = " (Hinweis: unter ARNI ist NT-proBNP typischerweise besser verwertbar)"
        lab_tail_lines.append(f"**{str(bnp_kind or 'BNP/NT-proBNP')}:** {_fmt(bnp_val,0)} pg/ml{extra}")

    cong_org = ui.get("congestive_organopathy")
    if isinstance(cong_org, str) and cong_org.lower().startswith("ja"):
        lab_tail_lines.append("Hinweis auf congestive Organopathie: ja")
    elif isinstance(cong_org, str) and cong_org.lower().startswith("nein"):
        lab_tail_lines.append("Hinweis auf congestive Organopathie: nein")

    lab_flow = "; ".join(lab_items) if lab_items else "Keine Laborwerte erfasst."
    lab_section = "### Labor\n" + lab_flow
    if lab_tail_lines:
        lab_section += "\n\n" + "\n".join(lab_tail_lines)

    parts.append(lab_section)

    # ---------------------------------------------------------------------
    # Bildgebung / Echo / CMR
    # ---------------------------------------------------------------------
    img_lines: List[str] = []

    if ui.get("ct_done"):
        findings = []
        for key, lab in [
            ("ct_ild", "ILD"),
            ("ct_emphysema", "Emphysem"),
            ("ct_embolie", "Embolie"),
            ("ct_mosaic", "Mosaikperfusion"),
            ("ct_koronarkalk", "Koronarkalk"),
        ]:
            if ui.get(key):
                findings.append(lab)
        if findings:
            img_lines.append(_md_kv("CT Thorax/Angio", ", ".join(findings)))
        else:
            img_lines.append(_md_kv("CT Thorax/Angio", "durchgeführt (keine pathologischen Befunde angegeben)"))
        ct_desc = (ui.get("ct_desc") or "").strip()
        if ct_desc:
            img_lines.append(_md_kv("CT Thorax Kurzbefund", ct_desc))
    else:
        img_lines.append(_md_kv("CT Thorax/Angio", "nicht angegeben"))

    if ui.get("vq_done"):
        vq_abn = "pathologisch" if ui.get("vq_defect") else "unauffällig/keine Defekte angegeben"
        img_lines.append(_md_kv("V/Q", vq_abn))
        vq_desc = (ui.get("vq_desc") or "").strip()
        if vq_desc:
            img_lines.append(_md_kv("V/Q Details", vq_desc))

        # Zusatz: PA Angio / CTEPH Konferenz (nur wenn aktiv dokumentiert)
        if ui.get("vq_pa_angio_done"):
            pa_desc = (ui.get("vq_pa_angio_desc") or "").strip()
            img_lines.append(_md_kv("PA Angio", pa_desc if pa_desc else "durchgeführt"))

        if ui.get("vq_cteph_conf_done"):
            dt = (ui.get("vq_cteph_conf_date") or "").strip()
            val = "erfolgt" + (f" ({dt})" if dt else "")
            img_lines.append(_md_kv("CTEPH Konferenz", val))
            dec_txt = (ui.get("vq_cteph_conf_decision") or "").strip()
            if dec_txt:
                img_lines.append(_md_kv("CTEPH Konferenz Beschluss", dec_txt))

    # Echo
    if ui.get("echo_done") or any(ui.get(k) not in (None, "", False) for k in ["lvef", "la_enlarged", "ee_ratio", "pasp_echo"]):
        echo_bits: List[str] = []
        lvef = _safe_float(ui.get("lvef"))
        if lvef is not None:
            echo_bits.append(f"LVEF {_fmt(lvef,0)}%")
        if ui.get("la_enlarged"):
            echo_bits.append("LA erweitert")
        ee = _safe_float(ui.get("ee_ratio"))
        if ee is not None:
            echo_bits.append(f"E/e' {_fmt(ee,1)}")
        pasp = _safe_float(ui.get("pasp_echo"))
        if pasp is not None:
            echo_bits.append(f"sPAP {_fmt(pasp,0)} mmHg")
        tapse = _safe_float(ui.get("tapse_mm"))
        if tapse is not None:
            echo_bits.append(f"TAPSE {_fmt(tapse,0)} mm")
        if der.get("tapse_spap") is not None:
            echo_bits.append(f"TAPSE/sPAP {_fmt(der.get('tapse_spap'),2)}")
        sprime = _safe_float(ui.get("s_prime_cm_s"))
        if sprime is not None:
            echo_bits.append(f"S' {_fmt(sprime,1)} cm/s")
        raesa = _safe_float(ui.get("ra_esa_cm2"))
        if raesa is not None:
            echo_bits.append(f"RA ESA {_fmt(raesa,0)} cm²")
        if der.get("raai") is not None:
            echo_bits.append(f"RAAI {_fmt(der.get('raai'),1)} cm²/m²")
        if der.get("s_prime_raai") is not None:
            echo_bits.append(f"S'/RAAI {_fmt(der.get('s_prime_raai'),2)}")
        ivcd = _safe_float(ui.get("ivc_diam_mm"))
        if ivcd is not None:
            echo_bits.append(f"IVC {_fmt(ivcd,0)} mm")
        ivcc = ui.get("ivc_collapse")
        if isinstance(ivcc, str) and ivcc:
            echo_bits.append(f"IVC Kollaps: {ivcc}")

        if echo_bits:
            img_lines.append(_md_kv("Echo", ", ".join(echo_bits)))
        else:
            img_lines.append(_md_kv("Echo", "durchgeführt (keine Details angegeben)"))

        echo_flags: List[str] = []
        if der.get("s_prime_raai_low") is True:
            echo_flags.append("S'/RAAI erniedrigt (<0,81)")
        if der.get("tapse_spap_reduced") is True:
            lbl = "TAPSE/sPAP vermindert"
            if der.get("tapse_spap_risk") == "hoch":
                lbl += " (hochgradig)"
            elif der.get("tapse_spap_risk") == "intermediär":
                lbl += " (mäßig)"
            echo_flags.append(lbl)
        if echo_flags:
            img_lines.append(_md_kv("Echo Zusatz", "; ".join(echo_flags)))
    # CMR
    if ui.get("cmr_done") or any(ui.get(k) not in (None, "", False) for k in ["rvef", "rvedv", "rvesv", "rvedvi", "rvesvi"]):
        cmr_bits: List[str] = []
        rvef = _safe_float(ui.get("rvef"))
        if rvef is not None:
            cmr_bits.append(f"RVEF {_fmt(rvef,0)}%")

        rvedv = _safe_float(ui.get("rvedv"))
        rvesv = _safe_float(ui.get("rvesv"))
        rvedvi_in = _safe_float(ui.get("rvedvi"))
        rvesvi_in = _safe_float(ui.get("rvesvi"))

        # Plausibilität: negative/0 Volumina als nicht vorhanden behandeln
        if rvedv is not None and rvedv <= 0:
            rvedv = None
        if rvesv is not None and rvesv <= 0:
            rvesv = None

        # Indexierung nur, wenn BSA vorhanden ist und Zielwerte leer sind
        bsa = _safe_float((case.get("derived") or {}).get("bsa_m2"))
        rvedvi = rvedvi_in
        rvesvi = rvesvi_in
        if rvedvi is None and rvedv is not None and bsa is not None and bsa > 0:
            rvedvi = rvedv / bsa
        if rvesvi is None and rvesv is not None and bsa is not None and bsa > 0:
            rvesvi = rvesv / bsa

        if rvedv is not None:
            cmr_bits.append(f"RVEDV {_fmt(rvedv,0)} ml")
        if rvesv is not None:
            cmr_bits.append(f"RVESV {_fmt(rvesv,0)} ml")
        if rvedvi is not None:
            cmr_bits.append(f"RVEDVi {_fmt(rvedvi,0)} ml/m²")
        if rvesvi is not None:
            cmr_bits.append(f"RVESVi {_fmt(rvesvi,0)} ml/m²")

        if cmr_bits:
            img_lines.append(_md_kv("CMR", ", ".join(cmr_bits)))
        else:
            img_lines.append(_md_kv("CMR", "durchgeführt (keine Details angegeben)"))
    if not img_lines:
        img_lines.append("Keine Bildgebung oder Echo oder CMR Angaben erfasst.")

    parts.append(_md_section("Bildgebung / Echo / CMR", img_lines, add_colon=is_doctor))

    # ---------------------------------------------------------------------
    # Lungenfunktion (Fließtext; Kommentar separat)
    # ---------------------------------------------------------------------
    if ui.get("lufu_done"):
        phen: List[str] = []
        if ui.get("lufu_obstructive"):
            phen.append("obstruktiv")
        if ui.get("lufu_restrictive"):
            phen.append("restriktiv")
        if ui.get("lufu_diffusion"):
            phen.append("Diffusionsstörung")

        lufu_items: List[str] = []
        if phen:
            lufu_items.append("Phänotyp: " + ", ".join(phen))

        # NOTE: UI was migrated to % Soll. Keep the historic keys for backward compatibility.
        fev1 = _safe_float(ui.get("fev1_l"))
        fvc = _safe_float(ui.get("fvc_l"))

        if fev1 is not None:
            lufu_items.append(f"FEV1: {_fmt(fev1,0)} %")
        if fvc is not None:
            lufu_items.append(f"FVC: {_fmt(fvc,0)} %")

        dlco = _safe_float(ui.get("dlco_sb"))
        if dlco is not None:
            lufu_items.append(f"DLCO: {_fmt(dlco,0)} %")

        dlco_va = _safe_float(ui.get("dlco_va"))
        if dlco_va is not None:
            lufu_items.append(f"DLCO/VA: {_fmt(dlco_va,0)} %")

        rv = _safe_float(ui.get("residual_volume_l"))
        if rv is not None:
            lufu_items.append(f"Residualvolumen (RV): {_fmt(rv,0)} %")

        lufu_flow = "; ".join(lufu_items) if lufu_items else "Lungenfunktion durchgeführt (Details nicht angegeben)."
        lufu_section = "### Lungenfunktion\n" + lufu_flow

        summ = (ui.get("lufu_summary") or "").strip()
        if summ:
            lufu_section += "\n\n**Kommentar:** " + summ

        parts.append(lufu_section)
    else:
        parts.append("### Lungenfunktion\nKeine Lungenfunktion erfasst.")

    # ---------------------------------------------------------------------
    # Spiroergometrie / CPET
    # ---------------------------------------------------------------------
    if ui.get("cpet_done"):
        cpet_items: List[str] = []
        v = _safe_float(ui.get("cpet_peak_vo2_ml_kg_min"))
        if v is not None:
            cpet_items.append(f"V'O2max/kg: {_fmt(v,1)} mL/min/kg")
        v = _safe_float(ui.get("cpet_peak_vo2_ml_min"))
        if v is not None:
            cpet_items.append(f"V'O2 Peak: {_fmt(v,0)} mL/min")
        v = _safe_float(ui.get("cpet_peak_vo2_pct_pred"))
        if v is not None:
            cpet_items.append(f"V'O2 Peak: {_fmt(v,0)} % Soll")
        v = _safe_float(ui.get("cpet_ve_vco2_slope"))
        if v is not None:
            cpet_items.append(f"V'E/V'CO2 Slope (VECO2s): {_fmt(v,1)}")
        v = _safe_float(ui.get("cpet_petco2_vt1_mmhg"))
        if v is not None:
            cpet_items.append(f"PETCO2 VT1: {_fmt(v,0)} mmHg")
        v = _safe_float(ui.get("cpet_ve_vco2_vt1"))
        if v is not None:
            cpet_items.append(f"VE/VCO2@VT1: {_fmt(v,1)}")
        v = _safe_float(ui.get("cpet_peak_o2_pulse_pct_pred"))
        if v is not None:
            cpet_items.append(f"Peak O2-Puls: {_fmt(v,0)} % Soll")
        v = _safe_float(ui.get("cpet_vo2_wr_slope_ml_min_w"))
        if v is not None:
            cpet_items.append(f"VO2Ws (ΔV'O2/ΔW): {_fmt(v,2)} mL/min/W")
        v = _safe_float(ui.get("cpet_vo2_vt1_ml_kg_min"))
        if v is not None:
            cpet_items.append(f"V'O2 VT1: {_fmt(v,1)} mL/min/kg")
        v = _safe_float(ui.get("cpet_vo2_vt1_ml_min"))
        if v is not None:
            cpet_items.append(f"V'O2 VT1: {_fmt(v,0)} mL/min")
        v = _safe_float(ui.get("cpet_vo2_vt2_ml_min"))
        if v is not None:
            cpet_items.append(f"V'O2 VT2: {_fmt(v,0)} mL/min")
        v = _safe_float(ui.get("cpet_spo2_nadir_pct"))
        if v is not None:
            cpet_items.append(f"SpO2 Nadir: {_fmt(v,0)} %")
        v = _safe_float(ui.get("cpet_spo2_rest_pct"))
        if v is not None:
            cpet_items.append(f"SpO2 Ruhe: {_fmt(v,0)} %")
        v = _safe_float(ui.get("cpet_spo2_peak_pct"))
        if v is not None:
            cpet_items.append(f"SpO2 Peak: {_fmt(v,0)} %")
        v = _safe_float(ui.get("cpet_o2_supp_l_min"))
        if v is not None and v > 0:
            cpet_items.append(f"O2 während CPET: {_fmt(v,1)} L/min")
        v = _safe_float(ui.get("cpet_rer_peak"))
        if v is not None:
            cpet_items.append(f"RER Peak: {_fmt(v,2)}")
        v = _safe_float(ui.get("cpet_hr_peak_bpm"))
        if v is not None:
            cpet_items.append(f"HF Peak: {_fmt(v,0)} 1/min")
        v = _safe_float(ui.get("cpet_hr_pct_pred"))
        if v is not None:
            cpet_items.append(f"HF Peak: {_fmt(v,0)} % Soll")
        pat = (ui.get("cpet_o2_pulse_pattern") or "").strip()
        if pat:
            cpet_items.append(f"O2Puls Verlauf: {pat}")

        v = _safe_float(ui.get("cpet_peak_o2_pulse_ml"))
        if v is not None:
            cpet_items.append(f"O2Puls Peak: {_fmt(v,1)} mL")
        v = _safe_float(ui.get("cpet_o2_pulse_slope"))
        if v is not None:
            cpet_items.append(f"O2Puls Slope: {_fmt(v,2)}")
        sys_v = _safe_float(ui.get("cpet_bp_sys_peak"))
        dia_v = _safe_float(ui.get("cpet_bp_dia_peak"))
        if sys_v is not None and dia_v is not None:
            cpet_items.append(f"RR Peak: {_fmt(sys_v,0)}/{_fmt(dia_v,0)} mmHg")
        sys_r = _safe_float(ui.get("cpet_bp_sys_rest"))
        dia_r = _safe_float(ui.get("cpet_bp_dia_rest"))
        if sys_r is not None and dia_r is not None:
            cpet_items.append(f"RR Ruhe: {_fmt(sys_r,0)}/{_fmt(dia_r,0)} mmHg")

        if bool(ui.get("cpet_angina")):
            cpet_items.append("Symptom: Angina")
        if bool(ui.get("cpet_dizziness")):
            cpet_items.append("Symptom: Schwindel/Präsynkope")
        if bool(ui.get("cpet_syncope")):
            cpet_items.append("Symptom: Synkope")
        if bool(ui.get("cpet_arrhythmia")):
            txt = (ui.get("cpet_arrhythmia_text") or "").strip()
            cpet_items.append("Arrhythmie" + (f" ({txt})" if txt else ""))
        st = (ui.get("cpet_st_changes") or "").strip()
        if st and st.lower() not in ("keine", "none"):
            cpet_items.append(f"ST/T: {st}")

        sr = (ui.get("cpet_stop_reason") or "").strip()
        if sr:
            cpet_items.append(f"Abbruchgrund: {sr}")
        v = _safe_float(ui.get("cpet_petco2_rest_mmhg"))
        if v is not None:
            cpet_items.append(f"PETCO2 Ruhe: {_fmt(v,0)} mmHg")
        v = _safe_float(ui.get("cpet_petco2_peak_mmhg"))
        if v is not None:
            cpet_items.append(f"PETCO2 Peak: {_fmt(v,0)} mmHg")
        v = _safe_float(ui.get("cpet_breathing_reserve_pct"))
        if v is not None:
            cpet_items.append(f"Atemreserve: {_fmt(v,0)} %")

        cpet_flow = "; ".join(cpet_items) if cpet_items else "CPET durchgeführt (Details nicht angegeben)."
        cpet_section = "### Spiroergometrie / CPET\n" + cpet_flow

        summ = (ui.get("cpet_summary") or "").strip()
        if summ:
            cpet_section += "\n\n**Kommentar:** " + summ

        # Optional: deterministic Spiro-Logic interpretation in doctor report
        if bool(ui.get("cpet_spiro_in_report")):
            try:
                import spiro_logic as _spiro
                res = _spiro.analyze(dict(ui))
                if res and res.report_text:
                    cpet_section += "\n\n**Spiro-Logic Interpretation:**\n" + res.report_text
            except Exception:
                pass

        parts.append(cpet_section)

    # Join sections
    return "\n\n".join([p for p in parts if p]).strip()



# =============================================================================
# Report DB phrase injection (optional)
# =============================================================================

def _report_db_text(case: Dict[str, Any], audience: str, section: str) -> str:
    """Return deterministic DB phrases for a given report section.

    - Never contains patient-identifiable data (DB is local + generic).
    - If DB is missing/unavailable, returns empty string.
    """
    if select_phrases is None:
        return ""
    try:
        env = case.get("env") or {}
        tags0 = []
        dec = case.get("decision") or {}
        if isinstance(dec.get("tags"), list):
            tags0 = [str(x) for x in (dec.get("tags") or []) if x]
        phrases, _tags = select_phrases(
            env=env,
            tags=tags0,
            audience=str(audience),
            section=str(section),
            safe_eval_bool_fn=safe_eval_bool,
        )
        phrases = [str(p).strip() for p in (phrases or []) if str(p).strip()]
        return "\n\n".join(phrases).strip()
    except Exception:
        return ""
# =============================================================================
# Doctor report (Markdown)
# =============================================================================

def build_doctor_report_template(case: Dict[str, Any], blocks: Dict[str, TextBlock]) -> str:
    fp = _case_fingerprint(case)
    cached = _cache_get('doctor_report_template', fp)
    if cached is not None:
        return cached

    """Arztbericht im Klinik-Layout (Muster-basiert, kompakt, nicht redundant).

    Ziele
    - feste Gliederung wie im Muster-DOCX
    - Arzt-Adressat (klinisch verwertbar, handlungsleitend)
    - keine Zahlenfriedhöfe, keine Dopplungen
    - CPET/Spiro-Logic nur als Kurzheadline + klinische Zusammenfassung (wenn vorhanden)
    """
    ui: Dict[str, Any] = case.get("ui", {}) or {}
    der: Dict[str, Any] = case.get("derived", {}) or {}
    sc: Dict[str, Any] = case.get("scores", {}) or {}
    dec: Dict[str, Any] = case.get("decision", {}) or {}
    env: Dict[str, Any] = case.get("env", {}) or {}

    ctx = build_render_ctx(case)

    def _par(line: str) -> None:
        line = (line or "").strip()
        if not line:
            return
        out.append(line)
        out.append("")  # blank line -> new paragraph in DOCX conversion

    def _clean_item(x: str) -> str:
        x = (x or "").strip()
        # strip common bullet markers the user might paste
        x = re.sub(r"^\s*[-•]\s+", "", x)
        return x.strip()

    def _split_items(s: str) -> List[str]:
        s = (s or "").strip()
        if not s or s in ("-", "—"):
            return []
        # Prefer newline-separated items
        if "\n" in s:
            items = [_clean_item(x) for x in s.splitlines()]
            return [x for x in items if x]
        # Fallback: semicolon-separated
        if ";" in s:
            items = [_clean_item(x) for x in s.split(";")]
            return [x for x in items if x]
        # Single item
        return [_clean_item(s)] if _clean_item(s) else []

    def _fmt_bp(sys_k: str, dia_k: str) -> Optional[str]:
        sbp = _safe_float(ui.get(sys_k))
        dbp = _safe_float(ui.get(dia_k))
        if sbp is None and dbp is None:
            return None
        if sbp is not None and dbp is not None:
            return f"{fmt_int(sbp)}/{fmt_int(dbp)} mmHg"
        if sbp is not None:
            return f"{fmt_int(sbp)} mmHg"
        return f"{fmt_int(dbp)} mmHg"

    def _kv(label: str, value: Any) -> Optional[str]:
        v = "" if value is None else str(value).strip()
        if not v or v in ("-", "—"):
            return None
        return f"{label}: {v}"

    def _bul(line: str, lvl: int = 0) -> None:
        line = (line or "").strip()
        if not line:
            return
        indent = "  " * max(0, int(lvl))
        out.append(f"{indent}- {line}")

    def _extract_section(md: str, title: str) -> List[str]:
        """Extract bullet lines from summarize_inputs section (best-effort)."""
        if not md:
            return []
        start = f"### {title}"
        if start not in md:
            return []
        chunk = md.split(start, 1)[1]
        # until next heading
        if "\n### " in chunk:
            chunk = chunk.split("\n### ", 1)[0]
        lines = []
        for ln in chunk.splitlines():
            ln = (ln or "").rstrip()
            if not ln.strip():
                continue
            if ln.strip().startswith("- "):
                lines.append(ln.strip())
        return lines

    out: List[str] = []

    # ------------------------------------------------------------------
    # Header block (fixed label line, as in the Muster)
    # ------------------------------------------------------------------
    _par("**Allgemeines, Klinik, Bildgebung und Funktion:**")

    # Kurz-Anamnese
    story = (ui.get("story") or "").strip()
    _par(f"Kurz-Anamnese: {story if story else '-'}")

    # Relevante Vorerkrankungen (prefer list formatting if multiple items)
    comorb_raw = (ui.get("comorbidities") or "").strip()
    comorb_items = _split_items(comorb_raw)
    if len(comorb_items) >= 2:
        _par("Relevante Vorerkrankungen: -")
        for it in comorb_items:
            _bul(it, 0)
        out.append("")
    elif len(comorb_items) == 1:
        _par(f"Relevante Vorerkrankungen: {comorb_items[0]}")
    else:
        _par("Relevante Vorerkrankungen: -")

    # Präprozedurale Sicherheitsangaben (werden in der UI erhoben; müssen im Bericht sichtbar sein)
    access_route = (ui.get("access_route") or "").strip()
    consent_done = ui.get("consent_done")
    anticoag_paused = ui.get("anticoag_paused")
    allergies_present = ui.get("allergies_present")
    allergies_list = ui.get("allergies_list") or []
    allergies_other = (ui.get("allergies_other_text") or "").strip()

    # Zugang
    if access_route:
        _par(f"Zugang (geplant): {access_route}")

    # Aufklärung/Einwilligung
    if consent_done is True:
        _par("Aufklärung/Einwilligung: erfolgt")
    elif consent_done is False:
        _par("Aufklärung/Einwilligung: nicht dokumentiert")

    # Antikoagulation pausiert?
    if anticoag_paused is True:
        _par("Antikoagulation: pausiert (bitte Periprozedur Plan prüfen)")
    elif anticoag_paused is False:
        _par("Antikoagulation: nicht pausiert (bitte Periprozedur Plan prüfen)")

    # Allergien
    if allergies_present is True:
        items = []
        if isinstance(allergies_list, list):
            items.extend([str(x).strip() for x in allergies_list if str(x).strip()])
        if allergies_other:
            items.append(allergies_other)
        _par(f"Allergien: {', '.join(items) if items else 'ja (nicht spezifiziert)'}")
    elif allergies_present is False:
        _par("Allergien: verneint")

    # Bekannte / vermutete PH
    if (not is_doctor) and ui.get("ph_known") is True:
        dx = (ui.get("ph_known_dx") or "").strip()
        _par(f"Bekannte PH-Diagnose: {dx if dx else '-'}")

        items: List[tuple[str, str]] = []
        first_dx = (ui.get("ph_first_dx") or "").strip()
        if first_dx:
            items.append(("Erstdiagnose", first_dx))
        reason = (ui.get("ph_reason_rhk") or "").strip()
        if reason:
            items.append(("Aktueller Anlass", reason))
        subtype = (ui.get("ph_known_subtype") or "").strip()
        if subtype:
            items.append(("Subtyp/Kontext", subtype))

        # Therapie (Episoden: Historie zuerst, dann aktuell, dann geplant)
        eps = _get_ph_tx_episodes(ui, der)
        if eps:
            hist = [format_ph_tx_episode_line(e) for e in eps if str(e.get("status") or "").strip().lower() in ("früher", "abgesetzt", "pausiert")]
            cur = [format_ph_tx_episode_line(e) for e in eps if str(e.get("status") or "").strip().lower() == "aktuell"]
            planned = [format_ph_tx_episode_line(e) for e in eps if str(e.get("status") or "").strip().lower() == "geplant"]
            if hist:
                items.append(("PH Therapie Historie", ", ".join([x for x in hist if x])))
            if cur:
                items.append(("PH Therapie aktuell", ", ".join([x for x in cur if x])))
            if planned:
                items.append(("PH Therapie geplant", ", ".join([x for x in planned if x])))

        # Optional legacy summary: freier Verlaufstext, falls dokumentiert
        tx_status = (ui.get("ph_tx_status") or "").strip()
        if tx_status and tx_status.lower() not in ("keine angabe", "-"):
            items.append(("Therapie-Verlauf (Kurz)", tx_status))

        for k, v in items:
            _bul(f"{k}: {v}", 0)
        out.append("")
    elif (not is_doctor) and ui.get("ph_suspected") is True:
        _par("PH-Verdachtsdiagnose: ja")

    # ------------------------------------------------------------------
    # Klinik (vitals + symptoms + function)
    # ------------------------------------------------------------------
    _par("Klinik")
    bp = _fmt_bp("bp_sys", "bp_dia")
    if bp:
        _bul(f"Blutdruck: {bp}")
    hr = _safe_float(ui.get("hr"))
    if hr is not None:
        _bul(f"Herzfrequenz: {fmt_int(hr)}/min")
    # Symptome (nur wenn aktiv)
    if ui.get("dizziness") is True:
        _bul("Schwindel: ja")
    syn = ui.get("syncope")
    if isinstance(syn, bool):
        if syn:
            _bul("Synkope: ja")
    else:
        syn_s = (syn or "").strip()
        if syn_s and syn_s.lower() not in ("nein", "keine"):
            _bul(f"Synkope: {syn_s}")
    if ui.get("exertional_dyspnea") is True:
        _bul("Belastungsdyspnoe: ja")
    who = (ui.get("who_fc") or "").strip()
    if who:
        _bul(f"WHO-FC: {who}")
    six = _safe_float(ui.get("six_mwd_m"))
    if six is not None:
        six_dt = (ui.get("six_mwd_date") or "").strip()
        if six_dt:
            _bul(f"6MWD: {_fmt(six,0)} m (Datum: {six_dt})")
        else:
            _bul(f"6MWD: {_fmt(six,0)} m")
    out.append("")

    # ------------------------------------------------------------------
    # Labor / Bildgebung / Funktion (as in Muster)
    # ------------------------------------------------------------------
    summ = summarize_inputs(case) or ""

    lab_lines = _extract_section(summ, "Labor")
    if lab_lines:
        _par("Labor:")
        for ln in lab_lines:
            _bul(ln[2:].strip(), 0)  # strip "- "
        out.append("")

    img_lines = _extract_section(summ, "Bildgebung / Echo / CMR")
    if img_lines:
        _par("Bildgebung / Echo / CMR:")
        for ln in img_lines:
            body = ln[2:].strip()
            # Nest detail lines for readability (matches Muster look)
            if body.lower().startswith(("v/q details:", "echo:", "cmr:", "mrt:", "v/q details")):
                _bul(body, 1)
            else:
                _bul(body, 0)
        out.append("")

    lufu_lines = _extract_section(summ, "Lungenfunktion")
    if lufu_lines:
        _par("Lungenfunktion:")
        for ln in lufu_lines:
            _bul(ln[2:].strip(), 0)
        out.append("")

    # CPET / Spiro-Logic (kompakt)
    cpet_present = any(_safe_float(ui.get(k)) is not None for k in ["cpet_peak_vo2_ml_kg_min", "cpet_ve_vco2_slope", "cpet_petco2_rest_mmhg", "cpet_hr_peak_bpm", "cpet_rer_peak"])
    if cpet_present or (ui.get("cpet_summary") or "").strip():
        _par("Spiroergometrie / CPET:")
        # Kontext (optional)
        prot = (ui.get("cpet_protocol") or "").strip()
        site = (ui.get("cpet_site") or "").strip()
        chrono = (ui.get("cpet_chrono_comment") or "").strip()
        if prot:
            _bul(f"Protokoll: {prot}", 0)
        if site:
            _bul(f"Ort/Setup: {site}", 0)
        if chrono:
            _bul(f"Chronotrope Limitierung: {chrono}", 0)
        try:
            import spiro_logic as _spiro
            res = _spiro.analyze(dict(ui))
        except Exception:
            res = None

        if res and (res.headline or res.clinical_summary):
            if res.headline:
                _bul(res.headline, 0)
            if res.clinical_summary:
                _bul(res.clinical_summary, 0)
        else:
            # Fallback: keep very short numeric cues
            v = _safe_float(ui.get("cpet_peak_vo2_ml_kg_min"))
            if v is not None:
                _bul(f"V'O2max/kg: {_fmt(v,1)} mL/min/kg", 0)
            v = _safe_float(ui.get("cpet_ve_vco2_slope"))
            if v is not None:
                _bul(f"V'E/V'CO2 Slope: {_fmt(v,1)}", 0)

        cpet_note = (ui.get("cpet_summary") or "").strip()
        if cpet_note:
            _bul(f"Kommentar: {cpet_note}", 0)
        out.append("")

    # ------------------------------------------------------------------
    # Beurteilung und Empfehlung (2 Absätze + Risikoscores)
    # ------------------------------------------------------------------
    _par("Beurteilung und Empfehlung:")

    # Use bundle blocks (compact narrative) – do NOT append numeric detail blocks here
    b_id = f"{dec.get('bundle','')}_B" if dec.get("bundle") else ""
    e_id = f"{dec.get('bundle','')}_E" if dec.get("bundle") else ""
    beur = render_block(blocks[b_id], ctx) if b_id and b_id in blocks else ""
    empf = render_block(blocks[e_id], ctx) if e_id and e_id in blocks else ""

    beur_p = _filter_narrative_block(markdown_to_plain(beur).strip(), ui, der)
    empf_p = _filter_narrative_block(markdown_to_plain(empf).strip(), ui, der)

    if beur_p:
        _par(beur_p)
    # second paragraph: keep it short; if empty, fall back to deterministic conclusion
    if empf_p:
        _par(empf_p)
    else:
        leading_cause = dec.get("leading_cause") or "unklaren Genese"
        leading_action = dec.get("leading_action") or ""
        _par(f"In der Zusammenschau der Befunde ergeben sich Hinweise auf mehrere mögliche Ursachen/Mechanismen ({leading_cause}). Eine eindeutige führende Zuordnung ist anhand der vorliegenden Angaben nicht sicher.")

    # Risk bullets (as in Muster)
    if sc.get("esc_ers_4s") or sc.get("esc_ers_3s") or sc.get("reveal_lite2") or sc.get("reveal_lite2_points"):
        if sc.get("esc_ers_4s"):
            _bul(f"ESC/ERS 4-Strata: {sc.get('esc_ers_4s')}", 0)
        if sc.get("esc_ers_3s"):
            _bul(f"ESC/ERS 3-Strata: {sc.get('esc_ers_3s')}", 0)
        if sc.get("reveal_lite2") is not None:
            cat = sc.get("reveal_lite2")
            pts = sc.get("reveal_lite2_points")
            if str(cat) == "nicht berechenbar":
                missing = sc.get("reveal_lite2_missing") or []
                miss_txt = ", ".join(missing) if missing else "Parameter unvollständig"
                _bul(f"REVEAL Lite 2: nicht berechenbar (fehlend: {miss_txt})", 0)
            else:
                cat_de = {"low": "niedrig", "intermediate": "intermediär", "high": "hoch"}.get(str(cat), str(cat))
                pts_txt = str(pts) if pts is not None else "—"
                _bul(f"REVEAL Lite 2: {pts_txt} Punkte ({cat_de})", 0)
        out.append("")

    # ------------------------------------------------------------------
    # Procedere (handlungsleitend, strukturiert)
    # ------------------------------------------------------------------
    _par("Procedere:")

    # Procedere muss ausschließlich die bewusst gewählten P-Module abbilden (Single Source of Truth).
    selected_mods = _normalize_module_ids(ui.get("modules") or [])

    policy = der.get("p_module_policy") or {}
    eff_policy = pmods_apply_overrides(policy, pmods_get_force_optional(ui))
    disabled_mods: Dict[str, str] = (eff_policy.get("disabled") or {})
    allowed_order: List[str] = eff_policy.get("allowed") or list(_ALL_P_MODULE_IDS)

    emitted_any = False
    for mid in allowed_order:
        if mid not in selected_mods:
            continue
        if mid in disabled_mods:
            # Modul ist nicht anwählbar oder klinisch nicht passend: im Bericht nicht ausgeben
            continue
        blk = blocks.get(mid)
        if not blk:
            continue

        # Templates können Platzhalter enthalten -> mit ctx rendern (SafeDict)
        if mid == "P01":
            txt = render_p01_dynamic(env)
        else:
            txt = render_block(blk, ctx)
            txt = filter_module_text(txt, env)
        txt = str(txt or "").strip()
        if not txt:
            continue

        # Bullet Rendering: Mehrzeiler bleiben als Unterpunkte strukturiert
        def _clean_bullet(s: str) -> str:
            s = str(s or "").strip()
            # Avoid double bullets like "- • ..." when templates already contain bullet glyphs.
            while s.startswith("•") or s.startswith("-") or s.startswith("–"):
                s = s[1:].lstrip()
            return s

        lines = [_clean_bullet(ln) for ln in txt.splitlines() if str(ln).strip()]
        if not lines:
            continue
        _bul(lines[0], 0)
        for sub in lines[1:]:
            _bul(sub, 1)
        emitted_any = True

    free = (ui.get("procedere_free") or "").strip()
    if free:
        _bul(free, 0)
        emitted_any = True

    if not emitted_any:
        _bul("Kein spezifisches Procedere ausgewählt oder ableitbar (Module nicht gewählt oder Daten fehlen).", 0)
    out.append("")

    # ------------------------------------------------------------------
    # Zusätzliche Hinweise (kontextbasiert, kurz)
    # ------------------------------------------------------------------
    hints: List[str] = []

    # DZL status (UI only, no inference)
    try:
        if bool(ui.get("dzl_flag")):
            _dzl = (ui.get("dzl_decision") or "").strip()
            if not _dzl:
                _dzl = "Noch nicht gefragt"
            hints.append(f"DZL: {_dzl}.")
    except Exception:
        pass

    bmi = _safe_float(der.get("bmi"))
    if bmi is not None and bmi >= 30:
        hints.append("Adipositas kann Dyspnoe und Leistungsfähigkeit beeinflussen; Belastbarkeit im Kontext (Training, Lagerung, Atemmuster) interpretieren.")

    # include selected warnings (non-technical)
    warn = case.get("warnings") or []
    if isinstance(warn, list):
        for w in warn:
            ww = _format_warning_item(w)
            if not ww:
                continue
            # keep only patient-facing actionable warnings out; clinician warning ok but short
            if "fehl" in ww.lower() or "unvoll" in ww.lower():
                continue
            hints.append(ww)

    # Study pre-screen (strictly based on existing captured fields)
    try:
        _study_hints = get_study_hints(case)
        if isinstance(_study_hints, list):
            for _h in _study_hints:
                if str(_h).strip():
                    hints.append(str(_h).strip())
    except Exception:
        pass

    if hints:
        _par("Zusätzliche Hinweise:")
        for h in list(dict.fromkeys([x for x in hints if x]))[:8]:
            _bul(h, 0)

    _res = "\n".join(out).rstrip()
    _cache_set('doctor_report_template', fp, _res)
    return _res
def build_doctor_report(case: Dict[str, Any], blocks: Dict[str, TextBlock]) -> str:
    fp = _case_fingerprint(case)
    cached = _cache_get('doctor_report', fp)
    if cached is not None:
        return cached

    ui = case["ui"]
    der = case["derived"]
    sc = case["scores"]
    dec = case["decision"]
    env = case["env"]

    ctx = build_render_ctx(case)

    def _build_ph_etiology_dd_block(d: Dict[str, Any]) -> str:
        """Return a compact differential etiology block for Interpretation.

        Principles
        - Uses only documented inputs (no assumptions, no imputation).
        - Never turns suggestions into a definitive diagnosis.
        - No procedural recommendations here (Procedere is a separate section).
        """
        try:
            et = (d or {}).get("ph_etiology") or {}
            cands = et.get("candidates") or []
            if not isinstance(cands, list) or not cands:
                return ""

            labels = [str(c.get("label_doc") or "").strip() for c in cands]
            labels = [x for x in labels if x]
            if not labels:
                return ""

            clear = bool(et.get("clear_leader"))
            head = "Ätiologische Einordnung: "
            if clear:
                head += f"Die Befunde sprechen am ehesten für eine führende {labels[0]}."
                if len(labels) > 1:
                    head += " Zusätzlich bestehen Hinweise auf " + ", ".join(labels[1:]) + "."
            else:
                head += (
                    "Es ergeben sich Hinweise auf mehrere mögliche Ursachen oder Mechanismen ("
                    + ", ".join(labels)
                    + "). Eine eindeutige führende Zuordnung ist anhand der vorliegenden Angaben nicht sicher."
                )

            # Evidence (short, deterministic; max 4 groups, max 3 cues per group)
            ev_lines: List[str] = []
            for c in cands[:4]:
                try:
                    g = c.get("group")
                    ev = c.get("evidence") or []
                    if not isinstance(ev, list):
                        continue
                    ev = [str(x).strip() for x in ev if str(x).strip()]
                    if not ev:
                        continue
                    ev_txt = "; ".join(ev[:3]) + ("; …" if len(ev) > 3 else "")
                    ev_lines.append(f"Gruppe {g}: {ev_txt}.")
                except Exception:
                    continue

            if ev_lines:
                # Use dot bullet char to avoid markdown hyphens in narrative sections
                ev_block = "\n" + "\n".join(["• " + ln for ln in ev_lines])
                return (head + ev_block).strip()
            return head.strip()
        except Exception:
            return ""

    def _hemo_interpretation_paragraph() -> str:
        """Generate a fluent, guideline-aligned interpretation paragraph placed *under* Beurteilung.

        Goals
        - Deterministic, compact, clinician-friendly narrative (not a bullet list).
        - Covers typical constellations (no PH, pre-capillary PH, IpcPH, CpcPH, unclassified).
        - Integrates provocation testing (exercise, fluid challenge) and key pathologic signals
          from the numeric summary (RAP/CI, PAC/PP).

        Notes on thresholds
        - Resting PH definition: mPAP >20 mmHg.
        - Pre-capillary PH: mPAP >20, PAWP ≤15, PVR >2 WU.
        - Post-capillary PH: mPAP >20, PAWP >15; IpcPH if PVR ≤2, CpcPH if PVR >2.
        - Exercise PH: mPAP/CO slope >3 WU; PAWP/CO slope >2 WU supports post-capillary component.
        - Fluid challenge (rapid ~500 mL saline): PAWP ≥18 mmHg suggests occult LHD/HFpEF.
        """
        d = der or {}

        # --- pull values (safe floats) ---
        mpap = _safe_float(d.get("mpap_rest"))
        pawp = _safe_float(d.get("pawp_rest"))
        pvr = _safe_float(d.get("pvr_rest"))
        rap = _safe_float(d.get("rap_rest"))
        ci = _safe_float(d.get("ci_rest"))
        pac = _safe_float(d.get("pac_rest_ml_per_mmhg"))
        pp = _safe_float(d.get("pp_pa_rest"))

        # If we do not have the core triad, do not invent.
        if mpap is None or pawp is None or pvr is None:
            return ""

        lines: List[str] = []

        # --- Rest classification ---
        # No PH
        if mpap <= 20:
            if pawp <= 15 and pvr < 2:
                lines.append(
                    "Die hämodynamischen Parameter in Ruhe liegen im Normbereich. "
                    "Es bestehen keine Kriterien für pulmonale Hypertonie."
                )
            elif pawp > 15:
                lines.append(
                    "Die Kriterien für pulmonale Hypertonie sind in Ruhe nicht erfüllt. "
                    "Auffällig sind erhöhte Linksherzfüllungsdrücke als Hinweis auf eine mögliche diastolische Dysfunktion/HFpEF Konstellation."
                )
            elif pvr >= 2:
                lines.append(
                    "Die Kriterien für pulmonale Hypertonie sind in Ruhe nicht erfüllt. "
                    "Bei erhöhter PVR ist die Konstellation im Kontext von Herzzeitvolumen und Messbedingungen zu interpretieren; eine frühe pulmonalvaskuläre Beteiligung kann nicht sicher ausgeschlossen werden."
                )
            else:
                lines.append("Die hämodynamischen Parameter in Ruhe liegen überwiegend im Normbereich; Kriterien für pulmonale Hypertonie sind nicht erfüllt.")

        # PH present (mPAP >20)
        else:
            # Pre-capillary
            if pawp <= 15 and pvr > 2:
                lines.append(
                    "Es liegen hämodynamische Kriterien für eine präkapilläre pulmonale Hypertonie vor."
                )
            # Unclassified (elevated mPAP, low PVR)
            elif pawp <= 15 and pvr <= 2:
                lines.append(
                    "Es besteht eine isolierte mPAP Erhöhung bei normalem PAWP und nicht erhöhter PVR. "
                    "Diese Konstellation erfüllt keine Kriterien einer präkapillären PH; eine Einordnung sollte im Kontext von Flow, Messbedingungen und klinischem Risiko erfolgen."
                )
            # Post-capillary
            elif pawp > 15:
                if pvr <= 2:
                    lines.append(
                        "Es liegen hämodynamische Kriterien für eine isolierte postkapilläre pulmonale Hypertonie vor, passend zu einer Linksherzerkrankung/HFpEF Konstellation."
                    )
                else:
                    lines.append(
                        "Es liegen hämodynamische Kriterien für eine kombinierte post und präkapilläre pulmonale Hypertonie vor. "
                        "Dies spricht für eine postkapilläre Komponente mit zusätzlicher pulmonalvaskulärer Beteiligung."
                    )
            else:
                lines.append("Es bestehen Kriterien für pulmonale Hypertonie. Die weitere Einordnung erfolgt anhand von PAWP und PVR im Gesamtkontext.")

        # --- Key hemodynamic signals (kept short, only if clearly abnormal) ---
        add_bits: List[str] = []
        if rap is not None and rap >= 10:
            add_bits.append("RAP erhöht als Hinweis auf rechtskardiale Füllungsdruckerhöhung")
        if ci is not None and ci < 2.0:
            add_bits.append("CI erniedrigt im Sinne einer Low output Konstellation")
        if pac is not None and pac < 2.0:
            if pp is not None and pp >= 30:
                add_bits.append("verminderte pulmonalarterielle Compliance mit erhöhter pulsatile RV Nachlast")
            else:
                add_bits.append("verminderte pulmonalarterielle Compliance als Hinweis auf erhöhte pulsatile RV Nachlast")
        if add_bits:
            lines.append("Zusätzlich zeigen sich " + "; ".join(add_bits) + ".")

        # --- PVOD/PCH DD (subtil; nur bei starker Konstellation) ---
        try:
            pvod_lvl = _safe_float(d.get("pvod_hint_level"))
            pvod_desc = str(d.get("pvod_hint_desc") or "").strip()
            if pvod_lvl is not None and pvod_lvl >= 2 and pvod_desc:
                # nur in präkapillärer Konstellation (PAWP ≤15) einblenden
                if mpap > 20 and pawp <= 15:
                    lines.append(f"Zusätzlich bestehen Red Flags für PVOD/PCH (DD): {pvod_desc}.")
        except Exception:
            pass

        # --- Fluid challenge ---
        if bool(d.get("volume_challenge_done")):
            pawp_post = _safe_float(d.get("vol_challenge_pawp_post"))
            if pawp_post is not None:
                if pawp_post >= 18:
                    lines.append(
                        "Die Volumenprovokation zeigt einen Anstieg der PAWP auf ≥18 mmHg und spricht damit für eine okkulte diastolische LV Dysfunktion/HFpEF. "
                        "Hinweis: Für die hämodynamische Antwort auf Fluid challenge bei PAH sind die Daten limitiert."
                    )
                else:
                    lines.append(
                        "Die Volumenprovokation zeigt keinen Anstieg der PAWP auf ≥18 mmHg und ergibt damit keinen Hinweis auf eine okkulte HFpEF Konstellation."
                    )

        # --- Exercise (optional; only if slopes are available) ---
        # User requirement: Belastungsbefunde (mPAP/CO-Slope, PAWP/CO-Slope) nur dann interpretieren,
        # wenn sie tatsächlich vorhanden sind. Keine "Leerstelle" im Text, nur weil die Checkbox gesetzt ist.
        if bool(d.get("exercise_done")):
            mpap_s = _safe_float(d.get("mpap_co_slope"))
            pawp_s = _safe_float(d.get("pawp_co_slope"))

            # If both slopes are missing, stay silent (no optional block).
            if not (mpap_s is None and pawp_s is None):
                # Build a dedicated optional block that is clearly separated.
                if mpap_s is not None and pawp_s is not None:
                    if mpap_s <= 3.0 and pawp_s <= 2.0:
                        lines.append(
                            "Unter Belastung zeigt sich keine abnorme pulmonale Druck Flow Reaktion und kein Hinweis auf eine belastungsassoziierte postkapilläre Komponente."
                        )
                    else:
                        ex_bits: List[str] = []
                        if mpap_s > 3.0:
                            ex_bits.append("eine abnorme pulmonale Druck Flow Reaktion")
                        if pawp_s > 2.0:
                            ex_bits.append("einen Hinweis auf eine belastungsassoziierte postkapilläre Komponente")
                        if ex_bits:
                            lines.append("Unter Belastung zeigt sich " + " und ".join(ex_bits) + ".")

                    # Steigungen dokumentieren, ohne Schwellenwerte zu wiederholen
                    lines.append(
                        "Die Steigungen betragen "
                        f"mPAP CO Slope {_fmt(mpap_s,1)} WU und PAWP CO Slope {_fmt(pawp_s,1)} WU."
                    )
                else:
                    # One slope only – keep phrasing conservative, but still a clear optional block.
                    ex_bits: List[str] = []
                    if mpap_s is not None:
                        ex_bits.append(
                            "abnorme pulmonale Druck Flow Reaktion"
                            if mpap_s > 3.0
                            else "keine abnorme pulmonale Druck Flow Reaktion"
                        )
                    if pawp_s is not None:
                        ex_bits.append(
                            "Hinweis auf eine belastungsassoziierte postkapilläre Komponente"
                            if pawp_s > 2.0
                            else "kein Hinweis auf eine belastungsassoziierte postkapilläre Komponente"
                        )
                    if ex_bits:
                        lines.append("Unter Belastung zeigt sich: " + "; ".join(ex_bits) + ".")

                    # Steigungen dokumentieren (wenn vorhanden), ohne Schwellenwerte zu wiederholen
                    s_bits: List[str] = []
                    if mpap_s is not None:
                        s_bits.append(f"mPAP CO Slope {_fmt(mpap_s,1)} WU")
                    if pawp_s is not None:
                        s_bits.append(f"PAWP CO Slope {_fmt(pawp_s,1)} WU")
                    if s_bits:
                        lines.append("Die Steigungen betragen " + " und ".join(s_bits) + ".")

        if not lines:
            return ""
        return "\n".join([l.strip() for l in lines if l.strip()])

    # Bundle blocks
    b_id = f"{dec['bundle']}_B"
    e_id = f"{dec['bundle']}_E"
    beurteilung = render_block(blocks[b_id], ctx) if b_id in blocks else f"[Fehlender Textblock: {b_id}]"
    beurteilung = _filter_narrative_block(beurteilung, ui, der)

    def _compose_rest_hemo_parenthetical(d: Dict[str, Any], ui: Dict[str, Any]) -> str:
        """Return a single parenthetical containing the mandatory rest hemodynamic parameters.

        Deterministic order; includes only values that are present. If core values are missing,
        returns an empty string (never invent).
        """
        mpap = d.get("mpap_rest")
        pawp = d.get("pawp_rest")
        pvr = d.get("pvr_rest")
        if mpap is None or pawp is None or pvr is None:
            return ""

        bits: List[str] = []
        spap = ui.get("spap_rest")
        dpap = ui.get("dpap_rest")
        if spap is not None and dpap is not None:
            bits.append(f"sPAP/dPAP {fmt_int(spap)}/{fmt_int(dpap)} mmHg")
        bits.append(f"mPAP {fmt_int(mpap)} mmHg")
        bits.append(f"PAWP {fmt_int(pawp)} mmHg")
        rap = d.get("rap_rest")
        if rap is not None:
            bits.append(f"RAP {fmt_int(rap)} mmHg")

        co = d.get("co_rest")
        if co is None:
            co = d.get("co")
        if co is not None:
            bits.append(f"CO {fmt_float(co, 2)} l/min")

        ci = d.get("ci_rest")
        if ci is None:
            ci = d.get("ci")
        if ci is not None:
            bits.append(f"CI {fmt_float(ci, 2)} l/min/m²")

        sv = d.get("sv_rest_ml")
        if sv is not None:
            bits.append(f"SV {fmt_int(sv)} ml")
        svi = d.get("svi_rest_ml_m2")
        if svi is not None:
            bits.append(f"SVI {fmt_int(svi)} ml/m²")

        bits.append(f"PVR {fmt_float(pvr, 2)} WU")
        pvri = d.get("pvri_rest")
        if pvri is not None:
            bits.append(f"PVRi {fmt_float(pvri, 2)} WU·m²")
        tpg = d.get("tpg_rest")
        if tpg is not None:
            bits.append(f"TPG {fmt_int(tpg)} mmHg")
        dpg = d.get("dpg_rest")
        if dpg is not None:
            bits.append(f"DPG {fmt_int(dpg)} mmHg")
        pp = d.get("pp_pa_rest")
        if pp is not None:
            bits.append(f"PP (PA) {fmt_int(pp)} mmHg")
        pac = d.get("pac_rest_ml_per_mmhg")
        if pac is not None:
            bits.append(f"PAC {fmt_int(pac)} ml/mmHg")
        rc = d.get("rc_time_rest_s")
        if rc is not None:
            bits.append(f"RC-Zeit {fmt_float(rc, 2)} s")

        return "(" + ", ".join(bits) + ")"
    # --- Dynamische Ergänzungen (Zahlen/Fakten) ---
    extra_lines: List[str] = []
    # systemische Hämodynamik / Oxygenierung (falls im Textblock nicht enthalten)
    if ctx.get("systemic_sentence") and "System" not in beurteilung:
        extra_lines.append(ctx["systemic_sentence"].strip())
    if ctx.get("oxygen_sentence") and "ox" not in beurteilung.lower():
        extra_lines.append(ctx["oxygen_sentence"].strip())

    # Integrate mandatory rest hemodynamics into the first assessment sentence to avoid redundancy.
    if der:
        import re
        rest_par = _compose_rest_hemo_parenthetical(der, ui)
        if rest_par:
            # Prefer replacing an existing parenthetical that starts with mPAP...
            if re.search(r"\(\s*mPAP[^)]*\)", beurteilung):
                beurteilung = re.sub(r"\(\s*mPAP[^)]*\)", rest_par, beurteilung, count=1)
            else:
                # Otherwise, inject right after the first sentence (deterministic).
                i = beurteilung.find(".")
                if i != -1:
                    beurteilung = (beurteilung[:i+1] + " " + rest_par + beurteilung[i+1:]).strip()

    # Belastung
    if der and der.get("exercise_done"):
        if ctx.get("exercise_protocol_sentence"):
            extra_lines.append(ctx["exercise_protocol_sentence"].strip())
        mpap_s = der.get("mpap_co_slope")
        pawp_s = der.get("pawp_co_slope")
        if mpap_s is not None or pawp_s is not None:
            s_bits = []
            if mpap_s is not None: s_bits.append(f"mPAP/CO-Slope {fmt_float(mpap_s, 2)} WU")
            if pawp_s is not None: s_bits.append(f"PAWP/CO-Slope {fmt_float(pawp_s, 2)} WU")
            extra_lines.append("Belastungshämodynamik: " + " / ".join(s_bits) + ".")
        d_spap = der.get("delta_spap")
        if d_spap is not None:
            extra_lines.append(f"ΔsPAP (Peak–Ruhe): {fmt_int(d_spap)} mmHg.")
        peak_ci = der.get("ci_peak")
        if peak_ci is not None:
            extra_lines.append(f"Peak CI: {fmt_float(peak_ci, 2)} l/min/m².")
        patt_desc = ctx.get("exercise_pattern_desc") or ""
        if patt_desc:
            extra_lines.append(f"Belastungsmuster: {patt_desc}.")

    # Vergleich (wenn vorhanden, aber im Textblock nicht schon enthalten)
    if ctx.get("comparison_sentence") and "Im Vergleich" not in beurteilung:
        extra_lines.append(ctx["comparison_sentence"].strip())

    # If no prior RHK comparison is available, say so explicitly (deterministic).
    if (not ctx.get("comparison_table_md")) and (not ctx.get("comparison_sentence")):
        extra_lines.append("Ein hämodynamischer Vorbefund zum Verlauf liegt nicht vor.")

    # If no exercise/volume/vasoreactivity testing was performed, say so explicitly.
    did_prov = bool(der and (der.get("exercise_done") or der.get("volume_done") or der.get("vaso_done")))
    if not did_prov:
        extra_lines.append("Keine Belastungs- oder Provokationsmanöver durchgeführt.")

    if extra_lines:
        beurteilung = (beurteilung.rstrip() + "\n\n" + "\n".join(extra_lines)).strip()

    # Guideline-aligned narrative interpretation (placed under the Beurteilung section).
    interpretation = _hemo_interpretation_paragraph().strip()

    # Optional: deepen the *course* interpretation (primary + secondary hemodynamic blocks)
    # Requirement: only add if values are available; no silent assumptions.
    try:
        from rhk_hemo_deep_interpretation import build_hemo_deep_interpretation

        _deep = build_hemo_deep_interpretation(ui, der)
        if str(_deep or "").strip():
            interpretation = (interpretation + "\n\n" + str(_deep).strip()).strip() if interpretation else str(_deep).strip()
    except Exception:
        pass

    empfehlung = render_block(blocks[e_id], ctx) if e_id in blocks else f"[Fehlender Textblock: {e_id}]"
    empfehlung = _filter_narrative_block(empfehlung, ui, der)

    # Empfehlung soll keine pathophysiologische Einordnung wiederholen.
    # Entferne Sätze, die explizit Schwellenwerte/Kriterien (mPAP/PAWP/PVR etc.) wiederholen.
    try:
        import re
        empfehlung = re.sub(r"\s*Es\s+liegen\s+hämodynamische\s+Kriterien[^.]*\.\s*", " ", empfehlung, flags=re.IGNORECASE)
        empfehlung = re.sub(r"\s*\([^)]*(mPAP|PAWP|PVR|TPG|DPG)[^)]*\)\s*", " ", empfehlung, flags=re.IGNORECASE)
        empfehlung = " ".join(str(empfehlung or "").split())
    except Exception:
        pass

    # RHK structured section
    # RHK Ruhehämodynamik: show a complete, clinician-friendly numeric summary when available.
    rest_lines = []
    rest_lines.append(f"- sPAP {_fmt(ui.get('spap_rest'),0)} / dPAP {_fmt(ui.get('dpap_rest'),0)} / mPAP {_fmt(der.get('mpap'),0)} mmHg")
    rest_lines.append(f"- PAWP {_fmt(ui.get('pawp_rest'),0)} mmHg, RAP {_fmt(ui.get('rap_rest'),0)} mmHg")
    rest_lines.append(f"- CO {_fmt(der.get('co'),2)} l/min, CI {_fmt(der.get('ci'),2)} l/min/m²")
    if der.get('sv_rest_ml') is not None or der.get('svi_rest_ml_m2') is not None:
        rest_lines.append(f"- SV {_fmt(der.get('sv_rest_ml'),0)} ml, SVI {_fmt(der.get('svi_rest_ml_m2'),0)} ml/m²")
    tail = []
    tail.append(f"PVR {_fmt(der.get('pvr'),2)} WU")
    if der.get('pvri') is not None:
        tail.append(f"PVRi {_fmt(der.get('pvri'),2)} WU·m²")
    if der.get('tpg') is not None:
        tail.append(f"TPG {_fmt(der.get('tpg'),0)} mmHg")
    if der.get('dpg') is not None:
        tail.append(f"DPG {_fmt(der.get('dpg'),0)} mmHg")
    if der.get('pp_pa_rest') is not None:
        tail.append(f"PP (PA) {_fmt(der.get('pp_pa_rest'),0)} mmHg")
    if der.get('pac_rest_ml_per_mmhg') is not None:
        tail.append(f"PAC {_fmt(der.get('pac_rest_ml_per_mmhg'),0)} ml/mmHg")
    if der.get('rc_time_rest_s') is not None:
        tail.append(f"RC-Zeit {_fmt(der.get('rc_time_rest_s'),2)} s")
    rest_lines.append("- " + ", ".join(tail))
    rest_line = "\n".join(rest_lines)

    exercise_block = ""
    if der.get("exercise_done"):
        ex_lines = []
        ex_lines.append(_md_kv("mPAP/CO-Slope", f"{_fmt(der.get('mpap_co_slope'),1)} WU"))
        ex_lines.append(_md_kv("PAWP/CO-Slope", f"{_fmt(der.get('pawp_co_slope'),1)} WU"))
        ex_lines.append(_md_kv("ΔsPAP (Peak–Ruhe)", f"{_fmt(der.get('delta_spap'),0)} mmHg"))
        ex_lines.append(_md_kv("peak CI", f"{_fmt(der.get('ci_peak'),2)} l/min/m²"))
        if der.get("adaptation_type"):
            ex_lines.append(_md_kv("Adaptionstyp", "homeometrisch" if der["adaptation_type"] == "homeometric" else "heterometrisch"))
        if der.get("exercise_pattern"):
            ex_lines.append(_md_kv("Belastungsmuster", describe_exercise_pattern(der.get("exercise_pattern"))))
        exercise_block = "#### Belastungshämodynamik\n" + "\n".join(ex_lines)

    volume_block = ""
    if der.get("volume_challenge_done"):
        vol_lines = []
        pawp_pre = der.get("vol_challenge_pawp_pre")
        pawp_post = der.get("vol_challenge_pawp_post")
        if pawp_pre is not None and pawp_post is not None:
            vol_lines.append(_md_kv("PAWP", f"{_fmt(pawp_pre,0)} → {_fmt(pawp_post,0)} mmHg"))
            if der.get("vol_challenge_delta_pawp") is not None:
                vol_lines.append(_md_kv("PAWP (Δ)", f"{_fmt(der.get('vol_challenge_delta_pawp'),0)} mmHg"))
        else:
            vol_lines.append(_md_kv("PAWP", "—"))

        if der.get("vol_challenge_delta_mpap") is not None:
            vol_lines.append(_md_kv("mPAP (Δ)", f"{_fmt(der.get('vol_challenge_delta_mpap'),0)} mmHg"))

        # Guideline-based endpoint: absolute PAWP response (≥18 mmHg after ~500 mL saline over 5–10 min)
        if pawp_post is not None:
            endp = "PAWP ≥18 mmHg (Hinweis okkulte HFpEF)" if bool(der.get("vol_challenge_pawp_ge_18")) else "PAWP <18 mmHg"
            vol_lines.append(_md_kv("Endpunkt", endp))
        volume_block = "#### Volumenchallenge\n" + "\n".join(vol_lines)

    vaso_block = ""
    if der.get("vaso_test_done"):
        vaso_lines = []
        vaso_lines.append(_md_kv("Agent", str(ui.get("vaso_agent") or "—")))
        if ui.get("vaso_response_desc"):
            vaso_lines.append(_md_kv("Antwort", str(ui.get("vaso_response_desc"))))
        vaso_block = "#### Vasoreaktivität\n" + "\n".join(vaso_lines)

    stepox_block = ""
    # Stufenoxymetrie only if meaningful: include only if >2 values are present.
    # IVC is intentionally not used.
    sat_keys = ["sat_svc", "sat_ra", "sat_rv", "sat_pa", "sat_ao"]
    sat_filled = sum(1 for k in sat_keys if _safe_float(ui.get(k)) is not None)
    if sat_filled >= 3:
        sat_lines = []
        for k, lab in [("sat_svc", "SVC"), ("sat_ra", "RA"), ("sat_rv", "RV"), ("sat_pa", "PA"), ("sat_ao", "AO")]:
            v = _safe_float(ui.get(k))
            if v is not None:
                sat_lines.append(_md_kv(lab, f"{_fmt(v,0)}%"))
        sat_lines.append(_md_kv("Interpretation", der.get("step_up_sentence") or "—"))
        stepox_block = "#### Stufenoxymetrie\n" + "\n".join(sat_lines)
    elif sat_filled > 0:
        # Do not interpret if too sparse; keep the report clean.
        pass

    curve_block = ""
    curve_flags = []
    if der.get("v_wave"):
        curve_flags.append("V-Welle (PAWP)")
    if der.get("a_wave"):
        curve_flags.append("A-Welle (PAWP)")
    if der.get("rap_a_wave_flag"):
        curve_flags.append("A-Welle (RAP)")
    if der.get("rap_v_wave_flag"):
        curve_flags.append("V-Welle (RAP)")
    if der.get("rv_pseudo_dip_flag"):
        curve_flags.append("Pseudo-Dip (RV)")
    if der.get("rv_dip_plateau_flag"):
        curve_flags.append("Dip-Plateau (RV)")
    if curve_flags:
        curve_block = "#### Kurvenmorphologie\n" + "\n".join([_md_kv("Befund", ", ".join(curve_flags))])

    # Risk lines (prominent, directly after dx)
    risk_lines = []
    if sc.get("esc_ers_4s"):
        risk_lines.append(_md_kv("ESC/ERS 4-Strata", sc["esc_ers_4s"]))
    if sc.get("esc_ers_3s"):
        risk_lines.append(_md_kv("ESC/ERS 3-Strata", sc["esc_ers_3s"]))
    if sc.get("reveal_lite2"):
        cat = sc.get("reveal_lite2")
        pts = sc.get("reveal_lite2_points")
        if cat == "nicht berechenbar":
            missing = sc.get("reveal_lite2_missing") or []
            miss_txt = ", ".join(missing) if missing else "Parameter unvollständig"
            risk_lines.append(_md_kv("REVEAL Lite 2", f"nicht berechenbar (fehlend: {miss_txt})"))
        else:
            cat_de = {"low": "niedrig", "intermediate": "intermediär", "high": "hoch"}.get(str(cat), str(cat))
            pts_txt = str(pts) if pts is not None else "—"
            risk_lines.append(_md_kv("REVEAL Lite 2", f"{pts_txt} Punkte ({cat_de})"))
    if der.get("hfpef_category"):
        risk_lines.append(_md_kv("HFpEF (H2FPEF)", f"{der['hfpef_category']} (~{_fmt(der.get('hfpef_percent'),0)}%)"))
    risk_block = "\n".join(risk_lines) if risk_lines else "Keine Risikostratifizierung möglich (Daten fehlen)."

    # Modules – ausschließlich bewusst gewählte P-Module (Single Source of Truth)
    selected = _normalize_module_ids(ui.get("modules") or [])

    policy = der.get("p_module_policy") or {}
    eff_policy = pmods_apply_overrides(policy, pmods_get_force_optional(ui))
    disabled_mods: Dict[str, str] = eff_policy.get("disabled") or {}
    allowed_order: List[str] = eff_policy.get("allowed") or list(_ALL_P_MODULE_IDS)

    all_mods = list(selected)

    skipped_mods = [m for m in all_mods if m in disabled_mods]
    all_mods = [m for m in all_mods if m not in disabled_mods]

    order_index = {mid: i for i, mid in enumerate(allowed_order)}
    all_mods = sorted(all_mods, key=lambda m: order_index.get(m, 10_000))

    modules_txts: List[str] = []
    for mid in all_mods:
        if mid == "P01":
            txt = render_p01_dynamic(env)
            modules_txts.append(f"**{mid} – {blocks.get(mid, TextBlock(mid, mid, '', 'module')).title}**\n{txt}")
            continue
        if mid in blocks:
            txt = render_block(blocks[mid], ctx)
            txt = filter_module_text(txt, env)
            if txt:
                modules_txts.append(f"**{mid} – {blocks[mid].title}**\n{txt}")

    recs = dec.get("recommendations") or []

    # Empfehlungen aus dem Regelwerk können Platzhalter aus der TextDB enthalten -> jetzt mit ctx auflösen
    def _fmt_rec(x: Any) -> str:
        try:
            return str(x).format_map(SafeDict(ctx)).strip()
        except Exception:
            return str(x).strip()

    recs = [_fmt_rec(r) for r in recs if str(r).strip()]
    # Verlaufskonsequenz (falls Vor-RHK angegeben)
    tr_rec = (ctx.get("comparison_recommendation_doc") or "").strip()
    if tr_rec and (tr_rec not in recs):
        recs = list(recs) + [tr_rec]

    # Guard: do not silently include/interpret exercise/volume/vaso modules unless explicitly checked.
    if bool(der.get("exercise_values_present")) and not bool(der.get("exercise_done")):
        recs = list(recs) + [
            "Hinweis: Belastungswerte sind im Datensatz vorhanden, die Belastungshämodynamik wurde jedoch nicht als durchgeführt markiert (Checkbox nicht gesetzt). "
            "Interpretation/Übernahme erfolgt daher nicht. Bitte ggf. Modul aktivieren oder Werte entfernen."
        ]

    # Age-adapted filtering: for patients >=70 years, suppress transplant references in recommendations.
    age = _safe_float(ui.get("age"))
    if age is not None and age >= 70 and recs:
        def _keep_rec(r: str) -> bool:
            s = (r or "").lower()
            return ("transplant" not in s) and ("ltx" not in s)
        recs = [r for r in recs if _keep_rec(str(r))]



    # Concluding sentence (kann multifaktoriell sein)
    eti = der.get("ph_etiology") if isinstance(der, dict) else None
    concluding = ""
    if isinstance(eti, dict) and str(eti.get("doc_conclusion") or "").strip():
        concluding = str(eti.get("doc_conclusion") or "").strip()
    else:
        # Keep the interpretation section recommendation-free. Any procedural/therapeutic
        # suggestions belong to the dedicated "Procedere" section.
        leading_cause = dec.get("leading_cause") or "unklaren Genese"
        concluding = f"In der Zusammenschau der Befunde gehen wir von einer führenden **{leading_cause}** aus."


    # Build final report
    header = (
        "# Rechtsherzkatheter – Befundbericht\n\n"
        f"**Datum:** {_dt.date.today().strftime('%d.%m.%Y')}\n"
        f"**Tool-Version:** {APP_VERSION}\n\n"
    )
    patient_line = ""
    if ui.get("name") or ui.get("firstname"):
        patient_line = f"**Patient:** {ui.get('firstname','')} {ui.get('name','')}".strip() + "\n\n"

    summary_block = summarize_inputs(case, mode="doctor")

    relevant_vor = _build_relevante_vorerkrankungen_line(ui)

    report = [
        header,
        patient_line,
        "Relevante Vorerkrankungen: " + relevant_vor + "\n\n",
        summary_block,
        "\n## Rechtsherzkatheter\n",
        "#### Ruhehämodynamik\n",
        rest_line,
    ]
    if exercise_block:
        report.append("\n" + exercise_block)
    if volume_block:
        report.append("\n" + volume_block)
    if vaso_block:
        report.append("\n" + vaso_block)
    if stepox_block:
        report.append("\n" + stepox_block)
    if curve_block:
        report.append("\n" + curve_block)

    # Verlauf / Vergleich (optional)
    if ctx.get("comparison_table_md"):
        report.append("\n#### Verlauf / Vergleich (Vorher → Jetzt)\n")
        report.append((ctx.get("comparison_table_md") or "").strip() + "\n")

    # Diagnosis + risk + assessment
    report.append("\n## Beurteilung\n")
    report.append(beurteilung.strip() + "\n")

    # Interpretation (konsolidiert, ohne Procedere-Dopplung)
    report.append("\n## Interpretation:\n")
    ie_parts: List[str] = []
    if interpretation:
        ie_parts.append(interpretation.strip())

    # Optional DB-driven Ergänzung (section-mirrored with patient report)
    _db_ie = _report_db_text(case, audience="doctor", section="rhk_ie")
    if _db_ie:
        ie_parts.append(_db_ie)

    # Ätiologie (Differenziallogik) – ohne konkrete Maßnahmen (siehe Procedere)
    _dd = _build_ph_etiology_dd_block(der)
    if _dd:
        ie_parts.append(_dd)
    else:
        # Fallback: alte, knappe Schlussformulierung
        _conc = _sanitize_concluding(concluding)
        if _conc:
            ie_parts.append(_conc)

    # NOTE: No recommendations in the interpretation block by requirement.
    _ie_txt = "\n\n".join([p for p in ie_parts if p]).strip()
    _ie_txt = _sanitize_interpretation_block(_ie_txt)
    report.append(_ie_txt + "\n" if _ie_txt else "")
    # PH therapy course (documented in UI) – placed between Empfehlung and Procedere
    ph_tx_block = _build_ph_therapieverlauf_block(ui, der)
    if ph_tx_block:
        report.append("\n" + ph_tx_block)

    if modules_txts or skipped_mods or ui.get("procedere_free") or recs:
        report.append("\n## Procedere:\n")
        if modules_txts:
            report.append("\n\n".join(modules_txts))
        if skipped_mods:
            report.append("\n_Hinweis: Nicht übernommen (in dieser Konstellation nicht anwählbar): "
                          + ", ".join(skipped_mods) + "._")
        free = (ui.get("procedere_free") or "").strip()
        if free:
            report.append("\n**Freitext:**\n" + free)
        if recs:
            report.append("\n**Zusätzliche Hinweise:**\n")
            report.extend([f"- {r}" for r in recs])

    _res = "\n".join(report).strip()
    _cache_set('doctor_report', fp, _res)
    return _res
# =============================================================================
# Patient report (plain language, no abbreviations/numbers)
# =============================================================================

def _stable_patient_seed(case: Dict[str, Any]) -> int:
    """Deterministic seed for patient text variants.

    Goal: different cases → different wording, but same case → stable wording
    across repeated generations.
    """
    ui = case.get("ui") or {}
    der = case.get("derived") or {}
    dec = case.get("decision") or {}

    # Select only a few stable, clinically relevant discriminators.
    key = {
        "bundle": dec.get("bundle"),
        "primary_dx": dec.get("primary_dx"),
        "hemo_category": der.get("hemo_category"),
        "exercise_done": bool(der.get("exercise_done")),
        "exercise_pattern": der.get("exercise_pattern"),
        "step_up_present": bool(der.get("step_up_present")),
        "ct_ild": bool(ui.get("ct_ild")),
        "ct_emphysema": bool(ui.get("ct_emphysema")),
        "vq_defect": bool(ui.get("vq_defect")),
        "hfpef_category": der.get("hfpef_category"),
        "anemia": bool(der.get("anemia")),
        "congestion": bool(der.get("congestion_likely")),
    }
    s = json.dumps(key, sort_keys=True, ensure_ascii=False, default=str)
    h = hashlib.md5(s.encode("utf-8")).hexdigest()
    return int(h[:8], 16)


def _patient_name(ui: Dict[str, Any]) -> str:
    # Support multiple possible UI key names (historic variants)
    first = ""
    for k in ("firstname", "first_name", "vorname", "first"):
        v = ui.get(k)
        if isinstance(v, str) and v.strip():
            first = v.strip()
            break

    last = ""
    for k in ("name", "lastname", "last_name", "nachname", "surname"):
        v = ui.get(k)
        if isinstance(v, str) and v.strip():
            last = v.strip()
            break

    full = (first + " " + last).strip()
    return full



def _patient_salutation(ui: Dict[str, Any], rng: random.Random) -> str:
    """Returns a stable, formal salutation.

    Patient-facing report templates predominantly use formal address (Sie/Ihre).
    Randomly switching between "Hallo" and "Guten Tag" caused an inconsistent
    register (Hallo + Sie), which is confusing for patients.
    """
    name = _patient_name(ui)
    if name:
        return f"Guten Tag {name},"
    return "Guten Tag,"


def _load_patient_textdb() -> Tuple[Dict[str, Any], Dict[str, List[str]], Dict[str, str], Dict[str, str]]:
    """Loads patient-facing text blocks if available (flat file, no folders)."""
    import sys
    app_dir = os.path.dirname(os.path.abspath(__file__))
    if app_dir not in sys.path:
        sys.path.insert(0, app_dir)

    for mod_name in ("rhk_textdb_patient", "rhk_textdb_patient_v7"):
        try:
            mod = __import__(mod_name)  # type: ignore
            blocks = getattr(mod, "PATIENT_BLOCKS", None)
            bundles = getattr(mod, "PATIENT_BUNDLES", None)
            module_summary = getattr(mod, "PATIENT_MODULE_SUMMARY", {}) or {}
            glossary = getattr(mod, "PATIENT_GLOSSARY", {}) or {}
            if isinstance(blocks, dict) and isinstance(bundles, dict):
                if not isinstance(module_summary, dict):
                    module_summary = {}
                if not isinstance(glossary, dict):
                    glossary = {}
                return blocks, bundles, module_summary, glossary
        except Exception:
            continue
    return {}, {}, {}, {}


def _load_echo_patient_textdb() -> Tuple[Dict[str, Any], Dict[str, str]]:
    """Loads echo patient-facing text blocks if available (flat file, no folders)."""
    import sys
    app_dir = os.path.dirname(os.path.abspath(__file__))
    if app_dir not in sys.path:
        sys.path.insert(0, app_dir)

    for mod_name in ("rhk_textdb_echo_patient",):
        try:
            mod = __import__(mod_name)  # type: ignore
            blocks = getattr(mod, "ECHO_PATIENT_BLOCKS", None)
            glossary = getattr(mod, "ECHO_PATIENT_GLOSSARY", {}) or {}
            if isinstance(blocks, dict):
                if not isinstance(glossary, dict):
                    glossary = {}
                return blocks, glossary
        except Exception:
            continue
    return {}, {}


def _pick_echo_patient_template(block: Any, rng: random.Random) -> str:
    """Pick a template variant for echo patient blocks."""
    if block is None:
        return ""
    if isinstance(block, dict):
        temps = block.get("templates")
        if isinstance(temps, (list, tuple)) and temps:
            return str(rng.choice(list(temps)))
        if isinstance(temps, str) and temps.strip():
            return str(temps)
        if isinstance(block.get("template"), str):
            return str(block.get("template"))
        return ""

    temps = getattr(block, "templates", None)
    if isinstance(temps, (list, tuple)) and temps:
        return str(rng.choice(list(temps)))
    if isinstance(temps, str) and temps.strip():
        return temps
    templ = getattr(block, "template", None)
    if isinstance(templ, str):
        return templ
    return ""


def _render_echo_patient_text(block_id: str, blocks: Dict[str, Any], ctx: Dict[str, Any], rng: random.Random) -> str:
    block = blocks.get(block_id)
    templ = _pick_echo_patient_template(block, rng)
    if not templ:
        return ""
    txt = templ.format_map(SafeDict(ctx)).strip()
    txt = re.sub(r"\n{3,}", "\n\n", txt)
    txt = re.sub(r"[ \t]{2,}", " ", txt)
    return txt.strip()


# =============================================================================
# Doctor report for Word/Clipboard (compact, ordered)
# =============================================================================

def build_doctor_report_for_copy(case: Dict[str, Any], blocks: Dict[str, TextBlock]) -> str:
    fp = _case_fingerprint(case)
    cached = _cache_get('doctor_report_copy', fp)
    if cached is not None:
        return cached

    """Backward-compatible alias.

    Single source of truth for the Arztbericht is `build_doctor_report`.
    Clipboard/DOCX/UI must match exactly.
    """
    _res = build_doctor_report(case, blocks)
    _cache_set('doctor_report_copy', fp, _res)
    return _res
def build_echo_patient_report(case: Dict[str, Any]) -> str:
    """Patient*innenbericht Echokardiographie (strukturierte Interpretation).

    Implementierung liegt in `rhk_echo_report_patient.py` und wird hier nur
    gecached/wrapped, um etablierte Schnittstellen stabil zu halten.
    """
    fp = _case_fingerprint(case)
    cached = _cache_get('echo_patient_report', fp)
    if cached is not None:
        return cached

    from rhk_echo_report_patient import build_echo_patient_report as _impl
    out = _impl(case)

    _cache_set('echo_patient_report', fp, out)
    return out

def _pick_patient_template(block: Any, rng: random.Random) -> str:
    """Pick a block template variant.

    Supports:
    - dataclass with .templates: list[str] or tuple[str]
    - dataclass with .template: str
    - dict-like entries
    """
    if block is None:
        return ""

    # dict-like
    if isinstance(block, dict):
        temps = block.get("templates")
        if isinstance(temps, (list, tuple)) and temps:
            return str(rng.choice(list(temps)))
        if isinstance(temps, str) and temps.strip():
            return temps
        if isinstance(block.get("template"), str):
            return str(block.get("template"))
        return ""

    temps = getattr(block, "templates", None)
    if isinstance(temps, (list, tuple)) and temps:
        return str(rng.choice(list(temps)))
    if isinstance(temps, str) and temps.strip():
        return temps

    templ = getattr(block, "template", None)
    if isinstance(templ, str):
        return templ

    return ""


def _render_patient_text(block_id: str, blocks: Dict[str, Any], ctx: Dict[str, Any], rng: random.Random) -> str:
    block = blocks.get(block_id)
    templ = _pick_patient_template(block, rng)
    if not templ:
        return ""
    txt = templ.format_map(SafeDict(ctx)).strip()

    # Normalize whitespace a bit
    txt = re.sub(r"\n{3,}", "\n\n", txt)
    txt = re.sub(r"[ \t]{2,}", " ", txt)
    return txt.strip()


def build_patient_report(case: Dict[str, Any]) -> str:
    fp = _case_fingerprint(case)
    cached = _cache_get('patient_report', fp)
    if cached is not None:
        return cached

    """Erstellt einen patientenfreundlichen Bericht (drucktauglich, mit echtem Mehrwert).

    Leitlinien für den Patientenbericht:
    - **Klarer Nutzen**: Was bedeutet der Befund konkret? Was passiert als Nächstes? Was kann ich selbst tun?
    - **Wenig Floskeln**: lieber kurze, konkrete Sätze.
    - **Keine verwirrenden Score-Labels** (z. B. "HFpEF-likely" wird übersetzt).
    - **Zahlen nur als Orientierung**: wenige Kernwerte + verständliche Einordnung.
    - **Dynamik**: Wenn ein Vor-RHK vorliegt → Verlauf (besser/stabil/schlechter) + Konsequenz.

    Hinweis: Dieser Text ersetzt kein ärztliches Gespräch.
    """

    ui: Dict[str, Any] = case.get("ui", {}) or {}
    der: Dict[str, Any] = case.get("derived", {}) or {}
    dec: Dict[str, Any] = case.get("decision", {}) or {}
    hf: Dict[str, Any] = case.get("hfpef", {}) or {}
    sc: Dict[str, Any] = case.get("scores", {}) or {}
    sc: Dict[str, Any] = case.get("scores", {}) or {}

    blocks, bundles, module_summary, glossary = _load_patient_textdb()
    rng = random.Random(_stable_patient_seed(case))

    # ------------------------------------------------------------------
    # Helfer
    # ------------------------------------------------------------------
    def _norm(x: Any) -> str:
        return str(x).strip() if x is not None else ""

    def _fmt_val(v: Any, digits: int = 1) -> str:
        vv = _safe_float(v)
        if vv is None:
            return "—"
        return _fmt(vv, digits)

    def _qual(label: str, v: Optional[float]) -> Optional[str]:
        if v is None:
            return None
        if label == "mPAP":
            if v <= 20:
                return "normal"
            if v <= 30:
                return "erhöht"
            return "deutlich erhöht"
        if label == "PAWP":
            if v <= 15:
                return "normal"
            if v <= 20:
                return "erhöht"
            return "deutlich erhöht"
        if label == "PVR":
            if v <= 2:
                return "normal"
            if v <= 3:
                return "leicht erhöht"
            if v <= 5:
                return "erhöht"
            return "deutlich erhöht"
        if label == "CI":
            # grobe Orientierung
            if v < 2.2:
                return "niedrig"
            if v < 2.5:
                return "grenzwertig"
            return "normal"
        if label == "RAP":
            if v <= 7:
                return "normal"
            if v <= 12:
                return "erhöht"
            return "deutlich erhöht"
        return None

    def _risk_txt(cat: Optional[str]) -> Optional[str]:
        if not cat:
            return None
        c = str(cat).strip().lower()
        if c.startswith("low") or c.startswith("niedrig"):
            return (
                "Die Gesamt-Einordnung wirkt derzeit eher stabil. "
                "Wenn Ihre Beschwerden ebenfalls stabil sind, reichen häufig Kontrollen im Abstand von einigen Monaten."
            )
        if c.startswith("inter") or "mittel" in c:
            return (
                "Die Gesamt-Einordnung spricht für einen mittleren Kontrollbedarf. "
                "Oft sind Kontrollen in Wochen bis wenigen Monaten sinnvoll, um Verlauf und Therapie gemeinsam zu überprüfen."
            )
        if c.startswith("high") or "hoch" in c:
            return (
                "Die Gesamt-Einordnung spricht für einen hohen Kontrollbedarf. "
                "Engmaschige Betreuung und ggf. intensivere Therapieoptionen im spezialisierten PH‑Zentrum sind wichtig."
            )
        return None

    def _bundle_patient_blocks(bundle_id: str) -> List[str]:
        """Welche Patienten-Bausteine (PX_*) passen zu welchem Bundle (Kxx)?"""
        bids = bundles.get(bundle_id) or []
        out: List[str] = []
        for bid in bids:
            if bid in blocks:
                out.append(bid)
        return out

    # ------------------------------------------------------------------
    # Kernwerte (Ruhe)
    # ------------------------------------------------------------------
    mpap = _safe_float(der.get("mpap_rest") if der.get("mpap_rest") is not None else der.get("mpap"))
    pawp = _safe_float(der.get("pawp_rest") if der.get("pawp_rest") is not None else der.get("pawp"))
    pvr = _safe_float(der.get("pvr_rest") if der.get("pvr_rest") is not None else der.get("pvr"))
    ci = _safe_float(der.get("ci_rest") if der.get("ci_rest") is not None else der.get("ci"))
    rap = _safe_float(der.get("rap_rest") if der.get("rap_rest") is not None else der.get("rap"))

    has_ph = bool(mpap is not None and mpap > 20)
    congestion = bool(der.get("congestion_likely"))
    hemo_cat = str(der.get("hemo_category") or "").strip().lower()

    # Grobe Einordnung (aus Regelwerk/Entscheidung)
    bundle = _norm(dec.get("bundle") or "")
    primary_dx = _norm(dec.get("primary_dx") or "")
    leading_cause = _norm(dec.get("leading_cause") or "")
    leading_action = _norm(dec.get("leading_action") or "")

    # Patientenbericht-Archetypen (H1...H6) – Fokusverschiebung ohne Diagnostik
    archetype_id = str(der.get("p_archetype_id") or "H0").strip().upper()
    if not archetype_id:
        archetype_id = "H0"

    _ARCH_BLOCKS = {
        "H1": {
            "measured": "PX_ARCH_H1_FOCUS_MEASURED",
            "meaning": "PX_ARCH_H1_FOCUS_MEANING",
        },
        "H2": {
            "measured": "PX_ARCH_H2_FOCUS_MEASURED",
            "meaning": "PX_ARCH_H2_FOCUS_MEANING",
        },
        "H3": {
            "measured": "PX_ARCH_H3_FOCUS_MEASURED",
            "meaning": "PX_ARCH_H3_FOCUS_MEANING",
        },
        "H4": {
            "measured": "PX_ARCH_H4_FOCUS_MEASURED",
            "meaning": "PX_ARCH_H4_FOCUS_MEANING",
        },
        "H5": {
            "measured": "PX_ARCH_H5_FOCUS_MEASURED",
            "meaning": "PX_ARCH_H5_FOCUS_MEANING",
        },
        "H6": {
            "measured": "PX_ARCH_H6_FOCUS_MEASURED",
            "meaning": "PX_ARCH_H6_FOCUS_MEANING",
        },
    }

    

    # ------------------------------------------------------------------
    # Vertikale Verfeinerung (Sub-Layer): Symptomprofil und Diskrepanzen
    # ------------------------------------------------------------------
    def _symptom_profile() -> Dict[str, Any]:
        """Return a lightweight symptom profile.

        Robustness note:
        - Example payloads (and occasionally UI integrations) may provide boolean values for symptom fields.
        - We must never assume string input here, otherwise example-loading can crash the whole app.
        """

        who = str(ui.get("who_fc") or "").strip().upper()

        def _to_bool(v: Any, truthy: set[str]) -> bool:
            if v is None:
                return False
            if isinstance(v, bool):
                return bool(v)
            s = str(v).strip().lower()
            return s in truthy

        syn = _to_bool(ui.get("syncope"), {"ja", "yes", "true", "1", "gelegentlich", "manchmal"})
        diz = _to_bool(ui.get("dizziness"), {"ja", "yes", "true", "1"})
        # stairs_flights: grobe Belastungsdyspnoe-Orientierung
        stairs = _safe_float(ui.get("stairs_flights"))

        sev = "unknown"
        if who in {"IV"}:
            sev = "high"
        elif who in {"III"}:
            sev = "high"
        elif who in {"II"}:
            sev = "moderate"
        elif who in {"I"}:
            sev = "low"

        # Wenn WHO fehlt, nutze Treppenangabe grob
        if sev == "unknown" and stairs is not None:
            if stairs <= 1:
                sev = "low"
            elif stairs <= 2:
                sev = "moderate"
            else:
                sev = "high"

        return {
            "who_fc": who,
            "syncope": syn,
            "dizziness": diz,
            "stairs_flights": stairs,
            "severity": sev,
        }

    def _discordance_flags(symp: Dict[str, Any]) -> Dict[str, bool]:
        # 1) hoher Druck + niedriger BNP
        d1 = bool((mpap is not None and mpap > 30) and (bio_qual == "niedrig"))

        # 2) eher niedriger Druck + starke Symptome
        strong_symp = bool(symp.get("severity") == "high" or symp.get("syncope"))
        d2 = bool((mpap is not None and mpap <= 25) and strong_symp)

        # 3) Echo wirkt beruhigend, Katheter hoch
        pasp_echo = _safe_float(ui.get("pasp_echo"))
        trv = _safe_float(ui.get("trv_ms"))
        echo_low = bool((pasp_echo is not None and pasp_echo < 40) or (trv is not None and trv < 2.8))
        d3 = bool(echo_low and (mpap is not None and mpap > 30))

        return {
            "high_mpap_low_bnp": d1,
            "low_pressure_high_symptoms": d2,
            "echo_ok_cath_high": d3,
        }

    def _arch_text(kind: str) -> str:
        """Optionaler Fokus-Text je Archetyp (Fallback: leer)."""
        bid = (_ARCH_BLOCKS.get(archetype_id) or {}).get(kind)
        if not bid:
            return ""
        return _render_patient_text(bid, blocks, ctx, rng)

    def _patientize_cause(txt: str) -> str:
        t = (txt or "").strip()
        if not t:
            return ""
        low = t.lower()
        # Vereinfachen: Gruppen-/Jargon entfernen und in Alltagssprache übersetzen
        if ("gruppe 2" in low) or ("linkskard" in low) or ("hfpef" in low):
            return "Hinweise, dass die linke Herzhälfte mitbeteiligt ist (Rückstau in die Lunge)."
        if ("gruppe 3" in low) or ("copd" in low) or ("ild" in low) or ("fibrose" in low) or ("hypox" in low):
            return "Hinweise, dass eine Lungenerkrankung/Atemwegsproblematik mitbeteiligt sein könnte."
        if ("gruppe 4" in low) or ("cteph" in low) or ("embol" in low) or ("thrombo" in low):
            return "Hinweise, dass ältere Blutgerinnsel in den Lungengefäßen eine Rolle spielen könnten."
        if ("gruppe 1" in low) or ("pah" in low) or ("pulmonal-arter" in low):
            return "Hinweise, dass vor allem die Lungengefäße selbst betroffen sind (pulmonal-arterielle Form)."
        # Fallback: Klammern mit "Gruppe" entfernen
        t = re.sub(r"\s*\(.*grupp(e)?\s*\d.*?\)\s*", " ", t, flags=re.IGNORECASE).strip()
        return t

    cause_patient = _patientize_cause(leading_cause or primary_dx)

    # Verlaufstrend (optional)
    trend_info = _compare_rhk_trend(ui, der)

    # Module: ausschließlich bewusst gewählte P-Module (Single Source of Truth)
    selected_mods = _normalize_module_ids(ui.get("modules") or [])

    policy = der.get("p_module_policy") or {}
    eff_policy = pmods_apply_overrides(policy, pmods_get_force_optional(ui))
    disabled_mods: Dict[str, str] = eff_policy.get("disabled") or {}
    allowed_order: List[str] = eff_policy.get("allowed") or list(_ALL_P_MODULE_IDS)

    all_mods = [m for m in selected_mods if m not in disabled_mods]

    order_index = {mid: i for i, mid in enumerate(allowed_order)}
    all_mods = sorted(all_mods, key=lambda m: order_index.get(m, 10_000))

    # Warnhinweise aus Plausibilität (z. B. CTEPH/Antikoag)
    anticoag_status = (ui.get("anticoag_status") or "").strip().lower()
    clot_hint = bool(ui.get("vq_defect") or ui.get("ct_pe") or ui.get("pe_history"))
    ild = bool(ui.get("ct_ild"))
    antifib_status = (ui.get("antifib_status") or "").strip().lower()

    warn_lines: List[str] = []
    if clot_hint and anticoag_status and anticoag_status not in {"ja", "yes", "true"}:
        warn_lines.append(
            "Es gibt Hinweise, die (auch) zu älteren Gerinnseln in den Lungengefäßen passen könnten. "
            "Bitte klären Sie zeitnah mit Ihrem Behandlungsteam, ob eine Blutverdünnung (Antikoagulation) notwendig ist."
        )
    if ild and antifib_status in {"", "nein", "unklar"}:
        warn_lines.append(
            "Bei Hinweisen auf eine Lungenfibrose ist eine spezialisierte Mitbetreuung wichtig. "
            "Bitte klären Sie, ob eine antifibrotische Therapie in Ihrem Fall sinnvoll ist."
        )

    # HFpEF (übersetzt)
    hf_cat = _norm(der.get("hfpef_category") or hf.get("hfpef_category") or "")
    hf_prob = _safe_float(der.get("hfpef_prob") or hf.get("hfpef_prob"))

    hf_txt: Optional[str] = None
    if hf_cat:
        c = hf_cat.lower()
        if "high" in c or "likely" in c:
            hf_txt = "Es gibt Hinweise, dass die linke Herzhälfte sich unter Belastung nicht optimal füllt (das kann zu einem Rückstau in die Lunge beitragen)."
        elif "inter" in c or "mid" in c:
            hf_txt = "Es gibt gewisse Hinweise, dass die linke Herzhälfte unter Belastung mitbeteiligt sein könnte."
        elif "low" in c or "unlikely" in c:
            hf_txt = "Es gibt eher keine klaren Hinweise, dass die linke Herzhälfte die Hauptursache ist."
        # Prozentangabe nur, wenn vorhanden – aber nicht als Score-Label
        if hf_txt and hf_prob is not None:
            hf_txt = hf_txt + f" (Orientierend: {int(round(hf_prob))}%)."

    # Risiko (vereinfachte Sprache)
    risk_txt = _risk_txt(der.get("risk_category"))

    # ESC/ERS Follow-up Risiko (4-Strata)
    esc4 = sc.get("esc_ers_4s")
    esc4_n = sc.get("esc_ers_4s_n")
    esc4_missing = sc.get("esc_ers_4s_missing") or []

    # BNP/NT-proBNP (patientenfreundliche Einordnung; keine harten Diagnosen)
    bnp_kind = (ui.get("bnp_kind") or "BNP/NT-proBNP")
    bnp_val = _safe_float(ui.get("bnp_value"))
    entresto = bool(ui.get("entresto"))

    def _bio_qual(kind: str, v: Optional[float]) -> Optional[str]:
        if v is None:
            return None
        k = (kind or "").upper()
        # Grobe Orientierung – bewusst weich formuliert
        if "NT" in k:
            if v < 300:
                return "niedrig"
            if v < 1400:
                return "erhöht"
            return "deutlich erhöht"
        # BNP
        if v < 100:
            return "niedrig"
        if v < 300:
            return "erhöht"
        return "deutlich erhöht"

    bio_qual = _bio_qual(str(bnp_kind), bnp_val)

    # Diskrepanz: hoher Druck, aber niedriger BNP/NT-proBNP
    try:
        _mp = _safe_float(der.get('mpap_rest') if der.get('mpap_rest') is not None else der.get('mpap'))
        disc['high_mpap_low_bnp'] = bool(_mp is not None and _mp > 30 and bio_qual == 'niedrig')
    except Exception:
        pass

    # ------------------------------------------------------------------
    # Bericht zusammensetzen
    # ------------------------------------------------------------------
    lines: List[str] = []
    pname = _patient_name(ui)
    salutation = _patient_salutation(ui, rng)

    # Kontext für patientenfreundliche Textbausteine
    ctx = {
        "name": pname,
        "salutation": salutation,
    }

    lines.append("# Patientenbericht zum Rechtsherzkatheter")
    meta = []
    if pname:
        meta.append(f"**Name:** {pname}")
    meta.append(f"**Datum:** {_dt.date.today().strftime('%d.%m.%Y')}")
    meta.append(f"**Version:** {APP_VERSION}")
    lines.append(" · ".join(meta))
    lines.append("")

    # 1) Kurzfazit (Nutzen)
    lines.append("## Das Wichtigste auf einen Blick")

    # Maximal 5 kurze Sätze, konsequent fallbezogen (keine Floskeln).
    overview = []  # type: list[str]

    # Satz 1: Kernbefund (mit Wert + Schwelle, wenn vorhanden)
    if has_ph and mpap is not None:
        overview.append(
            f"Der mittlere Druck in Ihren Lungengefäßen (mPAP) liegt bei {_fmt(mpap,0)} mmHg und ist damit deutlich erhöht (Lungenhochdruck ab >20 mmHg)."
        )
    elif has_ph:
        overview.append("Die Messwerte sprechen für eine Druckerhöhung in den Lungengefäßen (Lungenhochdruck).")
    else:
        overview.append("In der Messung finden sich keine Hinweise auf eine relevante Druckerhöhung in den Lungengefäßen.")

    # Satz 2: Muster (prä oder postkapillär), wenn ableitbar
    if has_ph and hemo_cat:
        if hemo_cat == "precap":
            if pvr is not None:
                overview.append(
                    f"Das Muster passt zu einer präkapillären Form, dabei ist der Widerstand in den Lungengefäßen erhöht (PVR {_fmt(pvr,1)} WU, erhöht ab >2 WU)."
                )
            else:
                overview.append("Das Muster passt zu einer präkapillären Form, dabei steht der Widerstand in den Lungengefäßen im Vordergrund.")
        elif hemo_cat in {"ipcph", "cpcph"}:
            if pawp is not None:
                overview.append(
                    f"Es gibt Hinweise, dass die linke Herzseite mitbeteiligt sein könnte (PAWP {_fmt(pawp,0)} mmHg, häufig erhöht ab >15 mmHg)."
                )
            else:
                overview.append("Es gibt Hinweise, dass die linke Herzseite mitbeteiligt sein könnte.")

    # Satz 3: Was fällt besonders auf (BNP oder Pumpfunktion)
    if bnp_val is not None:
        q = _bio_qual(str(bnp_kind), bnp_val)
        q_txt = f" ({q})" if q else ""
        if q == "niedrig":
            overview.append(
                f"Der Blutwert {bnp_kind} liegt bei {_fmt(bnp_val,0)} pg/ml{q_txt}. Das spricht eher gegen eine aktuell stark erhöhte Herzbelastung (der Wert wird aber auch von Alter, Nierenfunktion und akuten Infekten beeinflusst)."
            )
        elif q in {"erhöht", "deutlich erhöht"}:
            overview.append(
                f"Der Blutwert {bnp_kind} liegt bei {_fmt(bnp_val,0)} pg/ml{q_txt}. Das passt dazu, dass das Herz derzeit stärker belastet ist und wird im Verlauf als wichtiger Orientierungspunkt genutzt."
            )
        else:
            overview.append(
                f"Der Blutwert {bnp_kind} liegt bei {_fmt(bnp_val,0)} pg/ml. Wir nutzen diesen Wert zusammen mit Symptomen und Belastbarkeit, um den Verlauf besser einzuordnen."
            )
    elif ci is not None and _qual('CI', ci) in {"niedrig", "grenzwertig"}:
        overview.append(f"Die Pumpleistung ist eher reduziert (CI {_fmt(ci,2)} l/min/m²).")

    # Satz 4: Risiko (kurz, ohne langen Exkurs)
    if esc4:
        overview.append(f"Die Risikoeinstufung liegt aktuell bei {esc4}. Das beeinflusst, wie eng wir Therapie und Kontrollen planen.")

    # Satz 5: Unsicherheit + nächster Schritt (sofort erklären, warum und wie es weitergeht)
    eti = der.get("ph_etiology") if isinstance(der, dict) else None
    cand_n = len(eti.get("candidates") or []) if isinstance(eti, dict) else 0
    ambiguous = bool(cand_n > 1)

    if ambiguous:
        if leading_action:
            overview.append(
                f"Welche Ursache am meisten beiträgt, ist anhand der bisherigen Daten noch nicht sicher. Als nächstes klären wir gezielt {leading_action}, damit wir die Behandlung passend ausrichten können."
            )
        else:
            overview.append(
                "Welche Ursache am meisten beiträgt, ist anhand der bisherigen Daten noch nicht sicher. Deshalb ergänzen wir weitere Untersuchungen, um die Hauptursache zu klären und die Behandlung gezielt auszurichten."
            )
    else:
        if leading_action:
            overview.append(f"Als nächster Schritt klären wir gezielt {leading_action}, damit wir die Behandlung passend ausrichten können.")

    # Ausgabe (max. 5 Sätze)
    for s in overview[:5]:
        lines.append(s)
    lines.append("")

    # Optional DB-driven Ergänzung (section-mirrored with doctor report)
    _db_ie_p = _report_db_text(case, audience="patient", section="rhk_ie")
    if _db_ie_p:
        lines.append("### Ergänzende Einordnung")
        lines.append(_db_ie_p)
        lines.append("")

    # ------------------------------------------------------------------
    # Narrativer Kernteil (fallzentriert, nicht generisch)
    # ------------------------------------------------------------------
    lines.append("## Was wurde bei Ihnen gemessen – und warum ist das wichtig?")

    if has_ph:
        # 1) Druck + Widerstand + Einordnung als zusammenhängende Geschichte
        if mpap is not None:
            lines.append(
                f"Bei Ihnen wurde ein erhöhter Druck im Lungenkreislauf gemessen (mPAP {_fmt(mpap,0)} mmHg, Lungenhochdruck ab >20 mmHg). "
                "Das bedeutet: Das rechte Herz muss Blut gegen einen höheren Widerstand in Richtung Lunge pumpen."
            )
        else:
            lines.append(
                "Bei Ihnen zeigen die Messungen einen Lungenhochdruck. Das bedeutet: Das rechte Herz muss Blut gegen einen höheren Widerstand in Richtung Lunge pumpen."
            )

        if hemo_cat == "precap" and (pawp is not None) and (pvr is not None):
            lines.append(
                f"Der Druck vor der linken Herzhälfte ist dabei nicht erhöht (PAWP {_fmt(pawp,0)} mmHg). "
                f"Gleichzeitig ist der Widerstand in den Lungengefäßen deutlich erhöht (PVR {_fmt(pvr,1)} WU, erhöht ab >2 WU). "
                "Das Muster spricht eher für eine Ursache im Lungenkreislauf selbst oder im Zusammenhang mit einer Lungenerkrankung."
            )
        elif hemo_cat in {"ipcph", "cpcph"} and pawp is not None:
            lines.append(
                f"Der Druck vor der linken Herzhälfte ist erhöht (PAWP {_fmt(pawp,0)} mmHg). "
                "Das kann einen Rückstau in die Lunge begünstigen und wird bei der Einordnung mit berücksichtigt."
            )
        elif hemo_cat:
            # Muster bekannt, Werte teils nicht
            if hemo_cat == "precap":
                lines.append("Das Messmuster passt eher zu einer Form, bei der die Lungengefäße oder die Lunge selbst im Vordergrund stehen.")
            else:
                lines.append("Das Messmuster passt eher zu einer Form, bei der die linke Herzseite mitbeteiligt sein kann.")

        # 2) Pumpfunktion / Rückstau – nur wenn Werte vorhanden
        if ci is not None:
            qci = _qual("CI", ci)
            if qci in {"niedrig", "grenzwertig"}:
                lines.append(
                    f"Die Pumpleistung des Herzens ist dabei eher reduziert (CI {_fmt(ci,2)} l/min/m²). "
                    "Das kann erklären, warum Belastung schneller schwerfällt oder Schwindel auftreten kann."
                )
            else:
                lines.append(f"Die Pumpleistung ist im Rahmen der Messung nicht klar vermindert (CI {_fmt(ci,2)} l/min/m²).")

        if rap is not None:
            qrap = _qual("RAP", rap)
            if qrap in {"erhöht", "deutlich erhöht"}:
                lines.append(
                    f"Der Druck im rechten Vorhof (RAP) liegt bei {_fmt(rap,0)} mmHg und ist erhöht. "
                    "Das kann ein Hinweis auf eine stärkere Belastung der rechten Herzhälfte sein."
                )

        # 3) Biomarker – korrekt, fallbezogen, kurz
        if bnp_val is not None and bio_qual:
            if bio_qual == "niedrig":
                lines.append(
                    f"Der Blutwert {bnp_kind} ist bei Ihnen niedrig ({_fmt(bnp_val,0)} pg/ml). "
                    "Das spricht eher gegen eine aktuell ausgeprägte Herzüberlastung – wichtig ist aber immer die Gesamtschau mit Beschwerden und Messwerten."
                )
            elif bio_qual == "erhöht":
                lines.append(
                    f"Der Blutwert {bnp_kind} ist bei Ihnen erhöht ({_fmt(bnp_val,0)} pg/ml). "
                    "Das passt dazu, dass das Herz aktuell stärker arbeiten muss."
                )
            elif bio_qual == "deutlich erhöht":
                lines.append(
                    f"Der Blutwert {bnp_kind} ist bei Ihnen deutlich erhöht ({_fmt(bnp_val,0)} pg/ml). "
                    "Das ist ein Warnsignal dafür, dass das Herz stärker belastet ist und wird im Verlauf eng beobachtet."
                )

    else:
        lines.append(
            "In Ruhe zeigen die Messwerte keinen Lungenhochdruck. Wenn Beschwerden vor allem unter Belastung auftreten, kann das trotzdem abgeklärt werden, "
            "weil manche Veränderungen erst unter Belastung sichtbar werden."
        )

    # Archetyp-spezifische Fokusverschiebung (ohne neue Fakten, nur Schwerpunkt)
    t_arch = _arch_text("measured")
    if t_arch:
        lines.append(t_arch)

    lines.append("")


    # Volumenchallenge / Vasoreaktivität (falls durchgeführt) – patientengerecht
    if der.get("vol_challenge_done"):
        lines.append("## Zusatztest: Volumenchallenge (Flüssigkeitsbelastung)")
        t = _render_patient_text("PX_VOLUME_CHALLENGE", blocks, ctx, rng)
        if t:
            lines.append(t)

        pawp_pre = _safe_float(der.get("vol_challenge_pawp_pre"))
        pawp_post = _safe_float(der.get("vol_challenge_pawp_post"))
        d_pawp = _safe_float(der.get("vol_challenge_delta_pawp"))
        endp_ge18 = bool(der.get("vol_challenge_pawp_ge_18"))

        # Nur wenige Kernwerte, verständlich
        bits = []
        if pawp_pre is not None and pawp_post is not None:
            bits.append(f"PAWP vor/nach: {fmt_float(pawp_pre,0)} → {fmt_float(pawp_post,0)} mmHg")
        if d_pawp is not None:
            bits.append(f"Änderung: {fmt_float(d_pawp,0)} mmHg")
        if bits:
            lines.append("Orientierung: " + " | ".join(bits) + ".")
        lines.append("Einordnung: " + ("Der Druck auf der linken Herzseite steigt dabei deutlich an. Das kann zu einem Rückstau in die Lunge beitragen."
                                         if endp_ge18 else
                                         "Der Druck auf der linken Herzseite bleibt dabei eher niedrig. Das spricht eher gegen eine ausgeprägte Druckerhöhung durch Flüssigkeit allein."))
        lines.append("")

    if der.get("vaso_test_done"):
        lines.append("## Zusatztest: Vasoreaktivität")
        t = _render_patient_text("PX_VASOREACTIVITY", blocks, ctx, rng)
        if t:
            lines.append(t)

        agent = str(ui.get("vaso_agent") or "—")
        resp_desc = str(ui.get("vaso_response_desc") or "").strip()
        responder = der.get("vaso_responder")
        # Sehr kurze, patientengerechte Aussage
        if agent and agent != "—":
            lines.append(f"Testmedikament: {agent}.")
        if resp_desc:
            lines.append(f"Beobachtung: {resp_desc}.")
        if responder is True:
            lines.append("Einordnung: Es gab eine deutliche Entspannung der Lungengefäße im Test. Das kann für die weitere Therapieplanung relevant sein.")
        elif responder is False and resp_desc:
            lines.append("Einordnung: Es zeigte sich im Test keine ausgeprägte Entspannung nach den klassischen Kriterien.")
        lines.append("")
    # 2) Einordnung / Erklärung
    lines.append("## Was bedeutet das für Sie?")

    eti = der.get("ph_etiology") if isinstance(der, dict) else None
    eti_patient_line = str(eti.get("patient_cause_line") or "").strip() if isinstance(eti, dict) else ""
    cand_n = len(eti.get("candidates") or []) if isinstance(eti, dict) else 0
    ambiguous = bool(cand_n > 1)

    # Redundanz vermeiden: Die hämodynamische Grundlogik wurde oben bereits erklärt.
    # Hier Fokus auf Einordnung, offene Punkte und Konsequenzen (ohne Procedere-Details).
    if has_ph:
        if hemo_cat == "precap":
            lines.append("In der Zusammenschau ergibt sich eher ein präkapilläres Muster. Entscheidend ist nun, warum der Widerstand in den Lungengefäßen erhöht ist.")
        elif hemo_cat in {"ipcph", "cpcph"}:
            lines.append("In der Zusammenschau gibt es Hinweise auf eine Mitbeteiligung der linken Herzseite. Entscheidend ist nun, wie groß dieser Anteil ist und ob zusätzlich der Lungenkreislauf selbst betroffen ist.")
        else:
            lines.append("In der Zusammenschau ist die Einordnung möglich, aber nicht alle Teilaspekte sind eindeutig. Wir stützen uns deshalb auf mehrere Bausteine (Messwerte, Bildgebung, Belastbarkeit).")
        lines.append("")
    else:
        lines.append("Die Messwerte in Ruhe sind unauffällig. Wenn Beschwerden vor allem unter Belastung auftreten, kann das trotzdem weiter eingeordnet werden – manche Veränderungen zeigen sich erst dann.")
        lines.append("")

    # Sub-Layer: Symptom-Gewichtung (gleiche Fakten, andere Gewichtung)
    symp = _symptom_profile()
    if symp.get("syncope"):
        t = _render_patient_text("PX_SYMPTOM_PROFILE_SYNCOPE", blocks, ctx, rng)
        if t:
            lines.append(t)
            lines.append("")
    else:
        sev = symp.get("severity")
        bid = {
            "low": "PX_SYMPTOM_PROFILE_LOW",
            "moderate": "PX_SYMPTOM_PROFILE_MODERATE",
            "high": "PX_SYMPTOM_PROFILE_HIGH",
        }.get(sev or "")
        if bid:
            t = _render_patient_text(bid, blocks, ctx, rng)
            if t:
                lines.append(t)
                lines.append("")

    disc = _discordance_flags(symp)

    # Archetyp-spezifische Fokusverschiebung (Einordnung) – nur, wenn verfügbar
    t_arch2 = _arch_text("meaning")
    if t_arch2:
        lines.append(t_arch2)
        lines.append("")

    # Unsicherheit immer erklären (Frage 11a)
    if ambiguous:
        lines.append("Welche Ursache im Vordergrund steht, lässt sich anhand der vorliegenden Angaben noch nicht sicher festlegen.")
        if leading_action:
            lines.append(f"Als nächster Schritt klären wir deshalb gezielt {leading_action}. Damit wird klarer, welche Behandlung bei Ihnen am besten passt.")
        else:
            lines.append("Deshalb ergänzen wir weitere Untersuchungen. Ziel ist, die Hauptursache zu klären und die Behandlung gezielt auszurichten.")
        lines.append("")

    # Mögliche Ursachen in patientenfreundlicher Sprache (nur wenn vorhanden)
    if eti_patient_line:
        # vermeiden von doppelten Einleitungen wie "Hinweise auf ... Hinweise auf ..."
        cleaned = eti_patient_line
        cleaned = cleaned.replace("Hinweise auf:", "")
        cleaned = re.sub(r"\bHinweise auf\b", "", cleaned).strip()
        cleaned = re.sub(r"\s{2,}", " ", cleaned)
        if cleaned:
            lines.append(f"Mögliche Ursachen, die wir in Ihrem Fall prüfen: {cleaned}")
            lines.append("")

    # Shunt/Step up (wenn vorhanden)
    if bool(ui.get("chd_pos")) or bool(ui.get("step_up_present")):
        lines.append("Die Messungen geben Hinweise auf eine zusätzliche Verbindung zwischen Herzhöhlen. Das kann den Blutfluss beeinflussen und wird deshalb gezielt abgeklärt.")
        lines.append("")

    # Linke Herzseite (HFpEF Hinweis)
    if hf_txt:
        lines.append(hf_txt)
        lines.append("")

    # 3) Werte zur Orientierung (kompakt, aber konkret)

    lines.append("## Wichtige Werte zur Orientierung")
    hemo_items = []  # type: list[str]

    if mpap is not None:
        hemo_items.append(
            f"- **mPAP** (mittlerer Druck in den Lungengefäßen): {_fmt(mpap,0)} mmHg ({_qual('mPAP', mpap)}; Lungenhochdruck ab >20 mmHg)"
        )
    if pawp is not None:
        hemo_items.append(
            f"- **PAWP** (Druck vor der linken Herzhälfte): {_fmt(pawp,0)} mmHg ({_qual('PAWP', pawp)}; häufig erhöht ab >15 mmHg)"
        )
    if pvr is not None:
        hemo_items.append(
            f"- **PVR** (Widerstand in den Lungengefäßen): {_fmt(pvr,1)} WU ({_qual('PVR', pvr)}; erhöht ab >2 WU)"
        )
    if ci is not None:
        hemo_items.append(
            f"- **CI** (Pumpleistung bezogen auf die Körpergröße): {_fmt(ci,2)} l/min/m² ({_qual('CI', ci)})"
        )
    if rap is not None:
        hemo_items.append(
            f"- **RAP** (Druck im rechten Vorhof): {_fmt(rap,0)} mmHg ({_qual('RAP', rap)}; häufig erhöht ab >8 mmHg)"
        )
    if bnp_val is not None:
        q = _bio_qual(str(bnp_kind), bnp_val)
        q_txt = f"{q}" if q else ""
        if q_txt:
            q_txt = f" ({q_txt})"
        hemo_items.append(
            f"- **{bnp_kind}** (Blutwert bei Herzbelastung): {_fmt(bnp_val,0)} pg/ml{q_txt}"
        )

    if hemo_items:
        lines.extend(hemo_items)
    else:
        lines.append("Keine Kernwerte verfügbar.")

    lines.append("")

    # Hinweis zur Einordnung (kurz, ohne Glossar)
    lines.append("Wichtig: Entscheidend ist die Kombination dieser Werte und der Verlauf. Eine einzelne Zahl erklärt Beschwerden selten vollständig.")
    lines.append("")
    # Diskrepanz-Erklärungen (Sub-Layer)
    disc_blocks: List[str] = []
    if disc.get("high_mpap_low_bnp"):
        disc_blocks.append("PX_DISCORDANCE_HIGH_MPAP_LOW_BNP")
    if disc.get("low_pressure_high_symptoms"):
        disc_blocks.append("PX_DISCORDANCE_LOW_PRESSURE_HIGH_SYMPTOMS")
    if disc.get("echo_ok_cath_high"):
        disc_blocks.append("PX_DISCORDANCE_ECHO_OK_CATH_HIGH")

    if disc_blocks:
        lines.append("## Wenn Werte und Beschwerden nicht gut zusammenpassen")
        lines.append(
            "Das kommt bei Herz und Lungenerkrankungen häufiger vor. Wichtig ist dann, dass wir gezielt erklären, welcher Teil des Befundes für Sie im Alltag entscheidend ist."
        )
        for bid in disc_blocks:
            t = _render_patient_text(bid, blocks, ctx, rng)
            if t:
                lines.append(t)
        lines.append("")

    # 4) Verlauf / Vergleich (wenn vorhanden)
    if trend_info.get("has_prev"):
        lines.append("## Verlauf im Vergleich")
        # Kontext Therapie (falls angegeben)
        tx_txt = (trend_info.get("tx_txt") or "").strip()
        if ui.get("prev_is_initial"):
            lines.append("Diese Untersuchung dient (auch) als **Verlaufskontrolle nach einer Ausgangsmessung**.")
        if tx_txt:
            lines.append(f"Seit der Voruntersuchung wurde als Therapie angegeben: **{tx_txt}**.")
        lines.append(trend_info.get("sentence_patient") or "")
        if trend_info.get("detail_patient"):
            lines.append(trend_info.get("detail_patient"))
        recp = (trend_info.get("rec_patient") or "").strip()
        subtype_pat = (trend_info.get("subtype_patient") or "").strip()
        if subtype_pat:
            lines.append("")
            lines.append(subtype_pat)

        if recp:
            lines.append("")
            lines.append(f"**Was bedeutet das praktisch?** {recp}")
        lines.append("")

    # 5) Nächste Schritte / Module (konkret + Nutzen)
    lines.append("## Wie geht es weiter?")

    # Level/Sortierung aus Policy (v24.1+)
    levels_map: Dict[str, int] = (policy.get("levels") or {}) if isinstance(policy, dict) else {}

    # Kandidaten-Gruppen (für kurze, patientenfreundliche Begründungen)
    eti_groups: List[int] = []
    if isinstance(eti, dict) and isinstance(eti.get("candidates"), list):
        for c in (eti.get("candidates") or [])[:5]:
            try:
                eti_groups.append(int(c.get("group")))
            except Exception:
                continue

    risk_cat_local = str(der.get("risk_category") or "").lower()

    # "Stabil / nicht-hochrisikant" (für alltagsbezogene Studien-Hinweise)
    # - High / intermediate-high: keine Lifestyle-Hinweise im Patientenbericht
    # - sonst: erlaubt (inkl. intermediate-low/low)
    def _is_high_risk(cat: str) -> bool:
        c = (cat or "").strip().lower()
        if not c:
            return False
        # robust gegen unterschiedliche Schreibweisen
        return (
            c.startswith("high") or "high" in c or "hoch" in c or
            "intermediate-high" in c or "intermediate high" in c or "intermediatehigh" in c
        )

    lifestyle_allowed = not _is_high_risk(risk_cat_local)

    def _module_level(mid: str) -> int:
        try:
            lvl = int(levels_map.get(mid, 3))
        except Exception:
            lvl = 3
        return lvl if lvl in (1, 2, 3) else 3

    def _module_reason(mid: str) -> str:
        """Kurze, patientenfreundliche Begründung (nur wenn passend).

        Wichtig: Nicht jede Maßnahme hat eine eindeutige, einzelne Ursache.
        Die Gründe werden daher bewusst zurückhaltend formuliert.
        """
        # Rückstau / Wasser
        if mid == "P02" and congestion:
            return "weil es Hinweise auf Wassereinlagerungen bzw. Rückstau gibt"

        # Blutarmut
        if mid == "P13" and bool(der.get("anemia")):
            return "weil die Blutwerte auf eine Blutarmut hindeuten können"

        # Lunge/Atemwege
        if mid in ("P08", "P12") and (
            bool(der.get("ct_ild")) or bool(der.get("ct_emphysema")) or
            bool(der.get("lufu_restrictive")) or bool(der.get("lufu_obstructive")) or bool(der.get("lufu_diffusion"))
        ):
            return "weil Befunde an Lunge/Atemwegen auffällig sein können und wir das genauer einordnen möchten"

        # Linkes Herz
        if mid == "P09" and (2 in eti_groups or (pawp is not None and pawp > 15)):
            return "weil es Hinweise auf eine Beteiligung der linken Herzseite geben kann"

        # Gerinnsel/Embolien (CTEPH-/V/Q-Logik)
        if mid in ("P05", "P10") and (4 in eti_groups or bool(der.get("vq_defect")) or bool(der.get("ct_embolie")) or bool(der.get("ct_pe"))):
            return "weil Hinweise auf (ältere) Blutgerinnsel/Embolien eine Rolle spielen könnten"

        if mid == "P10" and (4 in eti_groups) and str(ui.get("anticoag_status") or "").lower() in ("nein", "unklar", ""):
            return "weil in diesem Zusammenhang die Frage nach einer Blutverdünnung besonders wichtig ist"

        # Autoimmun / Virologie / Genetik
        if mid == "P17" and bool(ui.get("immunology_pos")):
            return "weil bestimmte Autoimmun-/Rheuma-Erkrankungen Lungenhochdruck mit verursachen können"
        if mid == "P18" and bool(ui.get("virology_pos")):
            return "weil bestimmte Virusinfektionen in seltenen Fällen mit Lungenhochdruck zusammenhängen"
        if mid == "P20" and bool(ui.get("mutation_pos")):
            return "weil genetische Faktoren bei manchen Formen von Lungenhochdruck eine Rolle spielen können"

        # Advanced Therapies (nur wenn Gesamtlage eher schwer)
        if mid == "P25" and (risk_cat_local.startswith("high") or "hoch" in risk_cat_local):
            return "weil wir bei einer eher schweren Gesamtsituation frühzeitig auch weiterführende Optionen im Spezialzentrum mitdenken"

        return ""

    if all_mods:
        by_level: Dict[int, List[str]] = {1: [], 2: [], 3: []}
        for mid in all_mods:
            by_level[_module_level(mid)].append(mid)

        level_titles = {
            1: "Level I – prioritäre Empfehlungen",
            2: "Level II – sinnvolle Ergänzungen",
            3: "Level III – optional (je nach Kontext)",
        }

        lines.append(
            "Die folgenden Schritte sind, je nach Gesamtbild, geplant oder sinnvoll. "
            "Falls verfügbar, steht darunter kurz, warum das in Ihrer Situation relevant sein kann."
        )
        lines.append("")

        for lvl in (1, 2, 3):
            mids = by_level.get(lvl) or []
            if not mids:
                continue
            lines.append(f"### {level_titles.get(lvl, f'Level {lvl}')}")
            for mid in mids:
                txt = (module_summary.get(mid) or "").strip()
                if not txt:
                    # Fallback: Titel aus Arzt-TextDB
                    try:
                        from rhk_textdb import ALL_BLOCKS as _ALL
                        blk = _ALL.get(mid)
                        if blk is not None:
                            txt = str(blk.title)
                    except Exception:
                        txt = txt or mid

                reason = _module_reason(mid)
                if reason:
                    lines.append(f"- {txt}  \n  Warum bei Ihnen: {reason}.")
                else:
                    lines.append(f"- {txt}")
            lines.append("")
    else:
        # generischer Baustein
        t = _render_patient_text("PX_NEXT_STEPS", blocks, ctx, rng)
        if t:
            lines.append(t)
            lines.append("")

    # Alltag & Sicherheit (fallbezogen, ohne Coaching-Floskeln)
    lines.append("## Alltag und Sicherheit")

    if leading_action:
        lines.append(f"Der nächste Schwerpunkt in Ihrem Fall ist: {leading_action}.")
        lines.append("")

    # Therapiehistorie (nur Fakten, keine Empfehlung)
    eps = _get_ph_tx_episodes(ui, der)
    if eps:
        hist = [format_ph_tx_episode_line(e) for e in eps if str(e.get("status") or "").strip().lower() in ("früher", "abgesetzt", "pausiert")]
        cur = [format_ph_tx_episode_line(e) for e in eps if str(e.get("status") or "").strip().lower() == "aktuell"]
        planned = [format_ph_tx_episode_line(e) for e in eps if str(e.get("status") or "").strip().lower() == "geplant"]
        if hist:
            lines.append("PH Therapie Historie: " + ", ".join([x for x in hist if x]) + ".")
        if cur:
            lines.append("PH Therapie aktuell: " + ", ".join([x for x in cur if x]) + ".")
        if planned:
            lines.append("PH Therapie geplant: " + ", ".join([x for x in planned if x]) + ".")
        if hist or cur or planned:
            lines.append("")

    # Personalisierte Hinweise (nicht redundant, nicht belehrend)
    # Robust gegen bool / numerische Werte aus Beispielen oder externen Importen.
    def _to_bool2(v: Any, truthy: set[str]) -> bool:
        if v is None:
            return False
        if isinstance(v, bool):
            return bool(v)
        s = str(v).strip().lower()
        return s in truthy

    syn = _to_bool2(ui.get("syncope"), {"ja", "yes", "true", "1", "gelegentlich", "manchmal"})
    diz = _to_bool2(ui.get("dizziness"), {"ja", "yes", "true", "1"})

    if syn:
        lines.append("Da bei Ihnen Ohnmacht oder Beinahe Ohnmacht angegeben wurde, ist das ein besonders wichtiges Warnsignal. Bitte melden Sie sich bei erneuten Episoden zeitnah.")
    elif diz:
        lines.append("Da bei Ihnen Schwindel angegeben wurde, ist wichtig, Belastung so zu dosieren, dass keine Beinahe Ohnmacht auftritt. Bei deutlicher Zunahme bitte frühzeitig Rücksprache halten.")

    if congestion:
        lines.append("Es gibt Hinweise auf Rückstau. Neue Schwellungen oder eine rasche Gewichtszunahme über wenige Tage sollten zeitnah besprochen werden.")
    else:
        lines.append("Wenn neue Schwellungen, rasche Gewichtszunahme oder deutlich zunehmende Luftnot auftreten, sollte das frühzeitig abgeklärt werden.")

    # Studienbasierte Alltagshinweise nur bei stabilen / nicht-hochrisikanten Konstellationen
    if lifestyle_allowed:
        lines.append("")
        lines.append("### Evidenzbasierte Alltagshinweise")

        lines.append(
            "Studien bei Patient*innen mit Lungengefäßerkrankungen zeigen, dass regelmäßige, moderate Bewegung "
            "die Belastbarkeit und Lebensqualität verbessern kann. Entscheidend ist nicht Tempo, sondern "
            "eine gut verträgliche Regelmäßigkeit."
        )
        lines.append(
            "Als alltagsnahe Orientierung kann ein tägliches Gehziel im Bereich von etwa 7.000 bis 10.000 Schritten "
            "hilfreich sein, sofern dies ohne deutliche Luftnot, Schwindel oder Brustdruck möglich ist. "
            "Wenn Beschwerden unter Belastung zunehmen, ist eine geringere Dosis oft besser verträglich als "
            "ein seltener hoher Aufwand."
        )

        if congestion:
            lines.append(
                "Bei Hinweisen auf Rückstau kann eine individuell abgestimmte Trinkmenge und Salzaufnahme "
                "zur Entlastung beitragen. Konkrete Zielwerte werden im Gespräch festgelegt, weil sie von Nierenfunktion, "
                "Medikamenten und dem klinischen Verlauf abhängen."
            )

    if warn_lines:
        lines.append("")
        lines.append("**Wichtiger Hinweis:**")
        for w in warn_lines:
            lines.append(f"- {w}")

    lines.append("")

    # 7) Safety net
    lines.append("## Wann sollten Sie sich sofort melden?")
    t = _render_patient_text("PX_SAFETY_NET", blocks, ctx, rng)
    if t:
        lines.append(t)
    else:
        lines.append("- starke oder plötzlich zunehmende Luftnot in Ruhe")
        lines.append("- Brustschmerz/Brustdruck")
        lines.append("- Ohnmacht oder beinahe Ohnmacht")
        lines.append("- blutiger Auswurf/Husten von Blut")
        lines.append("- rasche Gewichtszunahme oder stark zunehmende Schwellungen")
    lines.append("")

    # 9) Disclaimer
    t = _render_patient_text("PX_DISCLAIMER", blocks, ctx, rng)
    if t:
        lines.append(t)

    # Clean spacing
    out = "\n".join([ln.rstrip() for ln in lines]).strip()
    out = re.sub(r"\n{3,}", "\n\n", out)
    _res = out
    _cache_set('patient_report', fp, _res)
    return _res

def build_internal_report(case: Dict[str, Any]) -> str:
    fp = _case_fingerprint(case)
    cached = _cache_get('internal_report', fp)
    if cached is not None:
        return cached

    env = case.get("env") or {}
    dec = case.get("decision") or {}
    debug = case.get("debug") or {}
    warns = case.get("warnings") or debug.get("warnings") or []
    rule_trace = debug.get("rule_trace") or {}
    fired = rule_trace.get("fired") or []
    errors = rule_trace.get("errors") or []

    lines = [
        "## Internal Debug",
        f"- Bundle: {dec.get('bundle')}",
        f"- Primary DX: {dec.get('primary_dx')}",
        f"- Tags: {', '.join(dec.get('tags') or [])}",
        f"- Missing (Regelwerk): {', '.join(dec.get('missing_fields') or [])}",
        f"- Warnungen (Plausibilität): {len(warns)}",
        "",
        "### Plausibilitätswarnungen (Auszug)",
    ]

    if not warns:
        lines.append("- keine")
    else:
        for w in warns[:12]:
            if isinstance(w, dict):
                try:
                    sev = str(w.get("severity") or "warn").upper()
                    msg = str(w.get("message") or "").strip()
                    flds = w.get("fields") or []
                    ftxt = f" (Felder: {', '.join([str(x) for x in flds])})" if flds else ""
                    lines.append(f"- [{sev}] {msg}{ftxt}")
                except Exception:
                    lines.append(f"- {_format_warning_item(w)}")
            else:
                msg = _format_warning_item(w)
                if msg:
                    lines.append(f"- {msg}")
        if len(warns) > 12:
            lines.append(f"- … weitere {len(warns) - 12} Warnungen")

    lines += [
        "",
        "### Regelwerk – Trace",
        f"- Ausgelöste Regeln: {len(fired)}",
        f"- Regel-Fehler: {len(errors)}",
        "",
        "#### Ausgelöste Regeln (Auszug)",
    ]

    if not fired:
        lines.append("- keine")
    else:
        for r in fired[:20]:
            try:
                rid = r.get("id")
                pr = r.get("priority")
                wh = str(r.get("when") or "")
                wh_short = (wh[:160] + "…") if len(wh) > 160 else wh
                lines.append(f"- {rid} (prio {pr}): {wh_short}")
            except Exception:
                continue
        if len(fired) > 20:
            lines.append(f"- … weitere {len(fired) - 20} Regeln")

    if errors:
        lines += ["", "#### Regel-Fehler (Auszug)"]
        for e in errors[:12]:
            try:
                rid = e.get("id")
                pr = e.get("priority")
                err = str(e.get("error") or "")
                err_short = (err[:180] + "…") if len(err) > 180 else err
                lines.append(f"- {rid} (prio {pr}): {err_short}")
            except Exception:
                continue
        if len(errors) > 12:
            lines.append(f"- … weitere {len(errors) - 12} Fehler")

    lines += [
        "",
        "### Env (Auszug)",
    ]

    keys = [
        "mpap", "pawp_rest", "pvr", "ci", "tpg", "dpg",
        "hemo_category", "precap", "ipcph", "cpcph",
        "hfpef_category", "hfpef_percent",
        "congestion_likely", "step_up_present", "step_up_from_to",
        "mpap_co_slope", "pawp_co_slope", "exercise_pattern",
        "adaptation_type",
        "s_prime_raai",
        "warnings_count",
    ]
    for k in keys:
        lines.append(f"- {k}: {env.get(k)}")
    _res = "\n".join(lines)
    _cache_set('internal_report', fp, _res)
    return _res
# =============================================================================
# Random example generation (now with lab constellations)
# =============================================================================

def random_example(scenario: Optional[str] = None, seed: Optional[int] = None) -> Dict[str, Any]:
    """
    Liefert ein zufälliges, aber in sich stimmiges Beispiel.

    Ziel:
    - Beispiele sollen möglichst viele Features abdecken (RHK Ruhe/Belastung, Volumen, Vaso, Step-up,
      Echo/CT/VQ/Lufu, Labor-Konstellationen).
    - Module werden teils bewusst vorselektiert, um die Procedere-Logik sichtbar zu machen.
    """
    today = _dt.date.today()

    rng = random if seed is None else random.Random(int(seed))


    scenarios = [
        "no_ph",            # normale Hämodynamik
        "pah_pre",          # präkapilläre PH (PAH-typisch)
        "cteph",            # CTEPH-Konstellation
        "ild_ph",           # ILD/Hypoxie
        "hfpef_ipcph",      # iPcPH (HFpEF-typisch)
        "cpcph",            # cPcPH
        "shunt_asd",        # Shunt/Step-up
    ]
    scen = scenario if (isinstance(scenario, str) and scenario in scenarios) else rng.choice(scenarios)

    ui: Dict[str, Any] = {}

    # --- Demografie ---
    if scen in ("pah_pre", "shunt_asd"):
        age = rng.choice([28, 34, 41])
        sex = rng.choice(["weiblich", "weiblich", "männlich"])
    elif scen in ("hfpef_ipcph", "cpcph"):
        age = rng.choice([62, 68, 74, 79])
        sex = rng.choice(["weiblich", "männlich"])
    else:
        age = rng.choice([45, 52, 58, 66, 72])
        sex = rng.choice(["weiblich", "männlich"])

    ui["firstname"] = rng.choice(["Anna", "Max", "Sofia", "Leon", "Mara", "Jonas"])
    ui["name"] = rng.choice(["Beispiel", "Muster", "Patient", "Testfall"])
    ui["age"] = age
    ui["sex"] = sex
    ui["height_cm"] = rng.choice([160, 168, 175, 182])
    ui["weight_kg"] = rng.choice([58, 72, 86, 98])

    ui["bp_sys"] = rng.choice([105, 115, 125, 135, 145])
    ui["bp_dia"] = rng.choice([65, 70, 75, 80, 85])
    ui["hr"] = rng.choice([55, 65, 75, 85, 95])

    ui["story"] = rng.choice([
        "Belastungsdyspnoe seit Monaten, reduzierte Belastbarkeit.",
        "Zunehmende Luftnot, gelegentlich Schwindel.",
        "Kontrolle nach PH-Verdachtsdiagnose.",
        "Therapieevaluation bei bekannter PH.",
    ])

    ui["ph_known"] = scen in ("pah_pre", "cteph", "cpcph")
    ui["ph_suspected"] = not ui["ph_known"]

    # Default: CHD/Shunt-Anamnese in Beispielen explizit setzen (UI-Gating)
    ui["chd_pos"] = (scen == "shunt_asd")
    ui["chd_type"] = "ASD (Vorhofseptumdefekt)" if ui["chd_pos"] else "keine Angabe"
    ui["chd_desc"] = "" if not ui["chd_pos"] else "Bekannter ASD, Shuntkonstellation in der Stufenoxymetrie."

    # --- PH-Status Konsistenz + Details ---
    # In der Praxis schließen sich "PH-Diagnose bekannt" und "PH-Verdachtsdiagnose" gegenseitig aus.
    # Bei bekannter PH füllen wir deshalb zusätzliche Kontextfelder, damit die UI nicht "leer" wirkt.
    if ui.get("ph_known"):
        if scen == "pah_pre":
            ui["ph_known_dx"] = "PAH (Gruppe 1)"
            ui["ph_known_subtype"] = rng.choice([
                "SOP: Systemsklerose-assoziierte PAH",
                "idiopathische PAH",
                "portopulmonale Hypertonie",
            ])
            ui["ph_current_meds"] = rng.choice([
                ["PDE‑5‑Hemmer", "Endothelin‑Rezeptorantagonist (ERA)"],
                ["PDE‑5‑Hemmer"],
                ["sGC‑Stimulator (Riociguat)", "Endothelin‑Rezeptorantagonist (ERA)"],
            ])
        elif scen == "cteph":
            ui["ph_known_dx"] = "CTEPH (Gruppe 4)"
            ui["ph_known_subtype"] = rng.choice([
                "inoperable CTEPH (BPA-Evaluation)",
                "Status nach LE mit Residuen",
                "CTED/CTEPH im Verlauf",
            ])
            ui["ph_current_meds"] = rng.choice([
                ["sGC‑Stimulator (Riociguat)"],
                ["sGC‑Stimulator (Riociguat)", "Diuretikum"],
            ])
            ui["ph_interventions"] = rng.choice([
                ["BPA (Ballonangioplastie, Katheter)"],
                ["PEA (Pulmonalisendarteriektomie, OP)"],
                [],
            ])
        elif scen == "cpcph":
            ui["ph_known_dx"] = "PH bei Linksherzerkrankung / HFpEF (Gruppe 2)"
            ui["ph_known_subtype"] = rng.choice([
                "HFpEF mit postkapillärer PH",
                "cPcPH bei HFpEF (Mischkomponente wahrscheinlich)",
                "Linksherzerkrankung im Verlauf",
            ])
            ui["ph_current_meds"] = rng.choice([
                ["Diuretikum"],
                ["Diuretikum", "Sauerstofftherapie"],
            ])
        else:
            ui["ph_known_dx"] = "Sonstige/unklar (Gruppe 5)"
            ui["ph_known_subtype"] = "unklar"
            ui["ph_current_meds"] = []

        # Pflichtähnliche Felder (in Beispielen immer gefüllt)
        ui["ph_first_dx"] = rng.choice(["03/2020", "09/2021", "01/2022", "06/2023"])
        ui["ph_reason_rhk"] = rng.choice(["Verlaufskontrolle", "Therapieentscheidung", "Neusymptomatik"])
        ui["ph_prev_meds"] = ui.get("ph_prev_meds") or []
        # Bei bekannter Diagnose ist Verdacht nicht gesetzt
        ui["ph_suspected"] = False
    else:
        # Kein bekannter PH-Status: Verdacht ist möglich, Details bleiben leer
        ui["ph_known_dx"] = None
        ui["ph_known_subtype"] = ""
        ui["ph_first_dx"] = ""
        ui["ph_reason_rhk"] = None
        ui["ph_current_meds"] = ui.get("ph_current_meds") or []
        ui["ph_prev_meds"] = ui.get("ph_prev_meds") or []
        ui["ph_interventions"] = ui.get("ph_interventions") or []

    # --- Klinik/Funktion ---
    ui["who_fc"] = rng.choice(["II", "III"]) if scen != "no_ph" else rng.choice(["I", "II"])
    ui["six_mwd_m"] = rng.choice([240, 320, 420]) if scen != "no_ph" else rng.choice([420, 480, 520])

    # --- EKG (Beispiele sollen fehlende vs vorhandene Daten demonstrieren) ---
    # ekg_present: True/False/None (None = keine Angabe)
    if rng.random() < 0.75:
        ui["ekg_present"] = True
        ui["ekg_rhs_signs"] = rng.sample(
            ["P pulmonale", "Rechtsachsenabweichung", "RV-Hypertrophie", "RBBB/in kompletter RSB", "S1Q3T3"],
            k=rng.choice([0, 1, 2]),
        )
        ui["ekg_other_text"] = ""
    else:
        ui["ekg_present"] = False
        ui["ekg_rhs_signs"] = []
        ui["ekg_other_text"] = ""
    ui["stairs_flights"] = rng.choice([0, 1, 2, 3])
    ui["syncope"] = rng.choices(["keine", "gelegentlich", "wiederholt"], weights=[0.83, 0.14, 0.03], k=1)[0]
    ui["hemoptysis"] = (scen == "cteph") and (rng.random() < 0.15)
    ui["dizziness"] = rng.choice([False, True])

    # --- Labor ---
    lab_mode = rng.choice(["normal", "inflammation", "anemia", "renal"])
    ui["crp_mg_l"] = rng.choice([2, 4, 6]) if lab_mode != "inflammation" else rng.choice([25, 60])
    ui["leukocytes_g_l"] = rng.choice([6.5, 7.8, 9.1]) if lab_mode != "inflammation" else rng.choice([12.0, 15.5])
    ui["creatinine_mg_dl"] = rng.choice([0.8, 1.0, 1.2]) if lab_mode != "renal" else rng.choice([1.8, 2.2])
    ui["egfr"] = rng.choice([65, 75, 85, 95]) if lab_mode != "renal" else rng.choice([25, 30, 35, 40, 45])
    ui["platelets_g_l"] = rng.choice([190, 240, 320])
    ui["inr"] = rng.choice([1.0, 1.1, 1.2])
    ui["ptt_s"] = rng.choice([28, 31, 34])

    # Hb gezielt setzen: in ~50% fehlt Hb, um "nicht anwählbar" Logik zu testen
    if rng.random() < 0.5:
        ui["hb_g_dl"] = None
        ui["anemia_type"] = None
    else:
        if lab_mode == "anemia":
            ui["hb_g_dl"] = rng.choice([9.8, 10.6, 11.4])
            ui["anemia_type"] = rng.choice(["mikrozytär", "normozytär", "makrozytär"])
        else:
            ui["hb_g_dl"] = rng.choice([12.6, 13.8, 15.1])
            ui["anemia_type"] = None

    ui["bnp_kind"] = rng.choice(["NT-proBNP", "BNP"])
    if scen in ("no_ph",):
        ui["bnp_value"] = rng.choice([40, 80, 120])
    elif scen in ("hfpef_ipcph", "cpcph"):
        ui["bnp_value"] = rng.choice([380, 900, 1800])
    else:
        ui["bnp_value"] = rng.choice([120, 380, 1200, 2400])

    # --- Bildgebung/Echo ---
    ui["ct_done"] = True
    ui["ct_koronarkalk"] = rng.choice([False, True])

    ui["ct_ild"] = (scen == "ild_ph")
    ui["ct_emphysema"] = (scen == "ild_ph") and rng.choice([False, True])
    ui["ct_embolie"] = (scen == "cteph")
    ui["ct_mosaic"] = (scen == "cteph")

    ui["vq_done"] = (scen == "cteph") or (rng.random() < 0.4)
    ui["vq_defect"] = (scen == "cteph") and ui["vq_done"]
    ui["vq_desc"] = "Mehrsegmentale Perfusionsdefekte." if ui["vq_defect"] else ""

    ui["echo_done"] = True
    ui["lvef"] = 60 if scen not in ("cpcph",) else 55
    ui["la_enlarged"] = True if scen in ("hfpef_ipcph", "cpcph") else rng.choice([False, True])
    ui["ee_ratio"] = 16 if scen in ("hfpef_ipcph", "cpcph") else rng.choice([9, 11, 13])
    ui["pasp_echo"] = rng.choice([35, 45, 60]) if scen != "no_ph" else 28
    ui["tapse_mm"] = 22 if scen == "no_ph" else rng.choice([14, 16, 18, 20])
    ui["atrial_fib"] = True if scen in ("hfpef_ipcph", "cpcph") else False


    # --- Zusatzfelder für Modul-Gating (damit "nicht anwählbar" Regeln sichtbar werden) ---

    # Antikoagulation
    if scen == "cteph" or ui.get("atrial_fib"):
        ui["anticoag_status"] = "ja"
        ui["anticoag_substance"] = rng.choice(["DOAC (Apixaban, Rivaroxaban)", "VKA (Phenprocoumon/Warfarin)"])
        ui["anticoag_indication"] = "CTEPH/CTEPD" if scen == "cteph" else "Vorhofflimmern"
        ui["anticoag_since"] = rng.choice(["09/2023", "03/2024", "11/2024"])
        ui["anticoag_note"] = ""
    else:
        ui["anticoag_status"] = "nein"
        ui["anticoag_substance"] = None
        ui["anticoag_indication"] = "keine Angabe"
        ui["anticoag_since"] = ""
        ui["anticoag_note"] = ""

    # Immunologie Autoimmun
    if scen in ("pah_pre",) and rng.random() < 0.35:
        ui["immunology_pos"] = True
        ui["immunology_items"] = rng.sample([
            "Systemische Sklerose (Sklerodermie)",
            "SLE (Lupus erythematodes)",
            "MCTD (Mixed connective tissue disease)",
            "Sjögren-Syndrom",
        ], k=1)
        ui["immunology_desc"] = "Autoimmunerkrankung bekannt."
    else:
        ui["immunology_pos"] = False
        ui["immunology_items"] = []
        ui["immunology_desc"] = ""

    # Virologie Infektiologie
    if scen in ("pah_pre",) and rng.random() < 0.15:
        ui["virology_pos"] = True
        ui["virology_items"] = rng.sample([
            "HIV",
            "Hepatitis B",
            "Hepatitis C",
            "Schistosomiasis (parasitär)",
        ], k=1)
        ui["virology_desc"] = "Infektiologischer Risikofaktor dokumentiert."
    else:
        ui["virology_pos"] = False
        ui["virology_items"] = []
        ui["virology_desc"] = ""

    # Mutation Genetik
    if (scen == "pah_pre") and (age is not None) and (age < 45) and (rng.random() < 0.18):
        ui["mutation_pos"] = True
        ui["mutation_items"] = rng.sample([
            "BMPR2 Mutation",
            "ALK1 ACVRL1 Mutation",
            "EIF2AK4 Mutation",
        ], k=1)
        ui["mutation_desc"] = "Hinweis auf hereditäre Konstellation."
    else:
        ui["mutation_pos"] = False
        ui["mutation_items"] = []
        ui["mutation_desc"] = ""

    # Abdomensonographie
    ui["abd_sono_done"] = rng.random() < 0.55
    if ui["abd_sono_done"]:
        if rng.random() < 0.12:
            ui["abd_sono_desc"] = "Hinweis auf Leberzirrhose und portale Hypertension."
        else:
            ui["abd_sono_desc"] = rng.choice(["Unauffällig.", "Normalbefund.", "Kein Hinweis auf Leberzirrhose."])
    else:
        ui["abd_sono_desc"] = ""
    ui["s_prime_cm_s"] = rng.choice([9.0, 11.0, 13.0])
    ui["ra_esa_cm2"] = rng.choice([16.0, 20.0, 26.0])

    # --- Lufu ---
    ui["lufu_done"] = True
    ui["lufu_obstructive"] = bool(ui["ct_emphysema"])
    ui["lufu_restrictive"] = bool(ui["ct_ild"])
    ui["lufu_diffusion"] = bool(ui["ct_ild"]) or (rng.random() < 0.35)
    ui["fev1_l"] = rng.choice([1.4, 2.1, 2.8])
    ui["fvc_l"] = rng.choice([2.0, 2.8, 3.6])
    ui["dlco_sb"] = rng.choice([35, 52, 68])
    ui["lufu_summary"] = rng.choice(["", "Leichte Diffusionsstörung.", "Obstruktives Muster."])

    # In einem Teil der Fälle explizit unauffällige Lufu setzen, damit P12 klar deaktiviert werden kann
    if scen == "no_ph" or rng.random() < 0.18:
        ui["lufu_obstructive"] = False
        ui["lufu_restrictive"] = False
        ui["lufu_diffusion"] = False
        ui["lufu_summary"] = rng.choice(["Unauffällig.", "Normalbefund.", "Keine relevanten Auffälligkeiten."])

    # Langzeit-Sauerstofftherapie (für Ätiologie- und Modul-Logik)
    if scen == "ild_ph":
        ui["ltot"] = True
        ui["ltot_flow_l_min"] = rng.choice([1.0, 2.0, 3.0])
    else:
        ui["ltot"] = False
        ui["ltot_flow_l_min"] = None
    # --- Hämodynamik (Ruhe) ---
    if scen == "no_ph":
        spap, dpap, pawp, co, rap = 28, 10, 10, 5.2, 6
    elif scen == "hfpef_ipcph":
        spap, dpap, pawp, co, rap = 55, 25, 20, 4.5, 10
    elif scen == "cpcph":
        spap, dpap, pawp, co, rap = 70, 35, 22, 3.6, 14
    else:  # präkapillär
        spap, dpap, pawp, co, rap = 72, 30, 10, 4.0, 9

    ui["spap_rest"] = spap
    ui["dpap_rest"] = dpap
    ui["mpap_rest"] = None  # berechnen lassen
    ui["pawp_rest"] = pawp
    ui["rap_rest"] = rap
    ui["co_rest"] = co
    ui["ci_rest"] = None
    ui["pvr_rest"] = None

    # --- Belastung / Volumen ---
    ui["exercise_done"] = scen in ("pah_pre", "hfpef_ipcph", "cpcph") and (rng.random() < 0.75)
    if ui["exercise_done"]:
        ui["exercise_protocol"] = rng.choice(["WHO-Rampe", "Stufenprotokoll"])
        ui["exercise_peak_watts"] = rng.choice([75, 100, 125, 150, 175])
        ui["spap_peak"] = spap + rng.choice([25, 35])
        ui["dpap_peak"] = dpap + rng.choice([10, 15])
        ui["mpap_peak"] = None
        ui["pawp_peak"] = pawp + (rng.choice([3, 10, 15]) if scen in ("hfpef_ipcph", "cpcph") else rng.choice([2, 4, 6]))
        ui["co_peak"] = co + rng.choice([1.0, 1.8, 2.5])
    else:
        ui["exercise_protocol"] = ""
        ui["exercise_peak_watts"] = None
        ui["spap_peak"] = None
        ui["dpap_peak"] = None
        ui["mpap_peak"] = None
        ui["pawp_peak"] = None
        ui["co_peak"] = None

    ui["volume_challenge_done"] = (scen == "hfpef_ipcph") and (rng.random() < 0.6)
    if ui["volume_challenge_done"]:
        ui["volume_ml"] = rng.choice([500, 750])
        ui["pawp_post"] = pawp + rng.choice([5, 8, 12])
        ui["mpap_post"] = None
        ui["co_post"] = co + rng.choice([0.5, 1.0])
    else:
        ui["volume_ml"] = None
        ui["pawp_post"] = None
        ui["mpap_post"] = None
        ui["co_post"] = None

    # --- Vaso (nur PAH-Beispiel) ---
    ui["vaso_test_done"] = (scen == "pah_pre") and (rng.random() < 0.5)
    if ui["vaso_test_done"]:
        ui["vaso_substance"] = rng.choice(["NO", "Iloprost"])
        ui["vaso_mpap_pre"] = None
        ui["vaso_mpap_post"] = None
        ui["vaso_response_desc"] = rng.choice([
            "Kein signifikanter Abfall des mPAP.",
            "Vasoreaktivitätskriterium erreicht (Abfall mPAP, CO stabil).",
        ])
    else:
        ui["vaso_substance"] = ""
        ui["vaso_mpap_pre"] = None
        ui["vaso_mpap_post"] = None
        ui["vaso_response_desc"] = ""

    # --- Stufenoxymetrie/Step-up (Shunt) ---
    if scen == "shunt_asd":
        ui["sat_svc"] = 65
        ui["sat_ivc"] = None
        ui["sat_ra"] = 80
        ui["sat_rv"] = 80
        ui["sat_pa"] = 80
        ui["sat_ao"] = 96

        # Kontext: angeborener Herzfehler/Shunt
        ui["chd_pos"] = True
        ui["chd_type"] = "ASD (Vorhofseptumdefekt)"
        ui["chd_desc"] = "Beispiel: ASD mit Links nach Rechts Shunt."
    else:
        ui["sat_svc"] = None
        ui["sat_ivc"] = None
        ui["sat_ra"] = None
        ui["sat_rv"] = None
        ui["sat_pa"] = None
        ui["sat_ao"] = None

        ui["chd_pos"] = False
        ui["chd_type"] = "keine Angabe"
        ui["chd_desc"] = ""

        # Default: kein bekannter angeborener Shunt
        ui["chd_pos"] = ui.get("chd_pos") if ui.get("chd_pos") is not None else False
        ui["chd_type"] = ui.get("chd_type") or "keine Angabe"
        ui["chd_desc"] = ui.get("chd_desc") or ""

    # --- Kurvenflags ---
    ui["wedge_v_wave"] = (scen in ("hfpef_ipcph", "cpcph")) and (rng.random() < 0.5)
    ui["wedge_a_wave"] = (scen in ("hfpef_ipcph", "cpcph")) and (rng.random() < 0.35)
    ui["rap_a_wave"] = rng.random() < 0.2
    ui["rap_v_wave"] = rng.random() < 0.15
    ui["rv_pseudo_dip"] = rng.random() < 0.1
    ui["rv_dip_plateau"] = rng.random() < 0.05

    # --- Procedere/Module ---
    ui["procedere_free"] = ""
    ui["modules"] = []

    # Sichtbare Demo-Auswahl: ein paar Module passend zum Beispiel
    if scen == "cteph":
        ui["modules"] = ["P10"]
    elif scen == "ild_ph":
        ui["modules"] = ["P12"]
    elif scen in ("hfpef_ipcph", "cpcph"):
        ui["modules"] = ["P09"]
    elif scen == "shunt_asd":
        ui["modules"] = ["P01"]
    elif scen == "pah_pre":
        ui["modules"] = ["P14"]

    # Optional: Schwangerschaft-Modul gelegentlich vorselektieren (nur wenn weiblich und <= 50)
    if sex == "weiblich" and age <= 50 and rng.random() < 0.15:
        ui["modules"] = list(dict.fromkeys(ui["modules"] + ["P21"]))

    # Optional: Anämie-Modul vorselektieren, wenn Hb tatsächlich niedrig ist
    hb = _safe_float(ui.get("hb_g_dl"))
    hb_low = 13.0 if sex == "männlich" else 12.0
    if hb is not None and hb < hb_low:
        ui["modules"] = list(dict.fromkeys(ui["modules"] + ["P13"]))

    # In der UI existieren Level-Gruppen (modules_lvl1/2/3). Beispiele sollen diese Logik sichtbar füllen.
    # Wir legen die Vorselektion standardmäßig in Level 3 ab (robust, da Level-Policy fallabhängig variieren kann).
    ui["modules_lvl1"] = []
    ui["modules_lvl2"] = []
    ui["modules_lvl3"] = list(ui.get("modules") or [])

    # --- Vor-RHK (gelegentlich) ---
    if rng.random() < 0.35:
        ui["prev_rhk_date"] = rng.choice(["03/21", "11/22", "06/23"])
        ui["prev_label"] = rng.choice(["stabiler Verlauf", "leicht progredient", "gebessert"])
        ui["prev_mpap"] = rng.choice([18, 24, 30])
        ui["prev_pawp"] = rng.choice([7, 12, 18])
        ui["prev_ci"] = rng.choice([2.1, 2.8, 3.2])
        ui["prev_pvr"] = rng.choice([1.5, 2.6, 4.2])
    else:
        ui["prev_rhk_date"] = ""
        ui["prev_label"] = ""
        ui["prev_mpap"] = None
        ui["prev_pawp"] = None
        ui["prev_ci"] = None
        ui["prev_pvr"] = None

    return ui


# =============================================================================
# Beispielreihe (Suite)
# =============================================================================

def example_suite_case(index: Any = 0) -> Dict[str, Any]:
    """Liefert ein Beispiel aus einer festen Suite.

    Ziel: Über mehrere Beispiele hinweg sollen möglichst viele Funktionen getestet werden,
    ohne Zufall und ohne implizite Datenannahmen.

    - Wiederholtes Klicken lädt das nächste Beispiel (Index modulo Suite-Länge).
    - Jede Suite belegt andere Pfade (RHK Ruhe/Belastung, Volumen, Vaso, Step-up,
      CT/VQ/Lufu, CPET, PH Therapieepisoden inkl. Restart und Sotatercept, Legacy-Import).
    """

    try:
        idx = int(index or 0)
    except Exception:
        idx = 0

    def _tx(lines: List[List[str]]) -> str:
        # 6 Spalten: Medikament, Status, seit, bis, Grund, Kommentar
        out_lines: List[str] = []
        for row in lines:
            row = (row or []) + [""] * (6 - len(row or []))
            out_lines.append("\t".join([str(c or "").strip() for c in row[:6]]).strip())
        return "\n".join([ln for ln in out_lines if ln])

    SUITE: List[Dict[str, Any]] = [
        {
            "id": "E01",
            "label": "PAH Restart und Sotatercept",
            "scenario": "pah_pre",
            "modules": ["P14"],
            "story": "Beispiel E01: PAH (präkapillär) mit Belastung und Vasoreaktivität. Therapieepisoden mit Restart und Sotatercept.",
            "overrides": {
                "ph_known": True,
                "ph_known_dx": "PAH (Gruppe 1)",
                "ph_known_subtype": "idiopathische PAH",
                "ph_first_dx": "03/2022",
                "ph_reason_rhk": "Verlaufskontrolle",
                "co_method": "Thermodilution",
                "exercise_done": True,
                "vaso_test_done": True,
                "vaso_substance": "NO",
                "vaso_response_desc": "Vasoreaktivitätskriterium erreicht (Abfall mPAP, CO stabil).",
                "ph_tx_table": _tx([
                    ["Opsumit (Macitentan)", "aktuell", "01/2024", "", "", ""],
                    ["Sildenafil", "abgesetzt", "05/2023", "10/2023", "Unverträglichkeit/Nebenwirkung", "Kopfschmerz"],
                    ["Sildenafil", "aktuell", "11/2023", "", "", "Restart"],
                    ["Sotatercept", "geplant", "02/2026", "", "", "Therapieplanung"],
                ]),
                "ph_current_meds": [],
                "ph_prev_meds": [],
                "ph_new_meds": [],
                "ph_stopped_meds": [],
                "consent_done": True,
                "access_route": "V. jugularis rechts",
                "allergies_present": True,
                "allergies_list": ["Pflaster"],
                "cpet_done": True,
                "cpet_peak_vo2_ml_kg_min": 13.8,
                "cpet_peak_vo2_pct_pred": 56,
                "cpet_ve_vco2_slope": 42,
                "cpet_petco2_vt1_mmhg": 30,
                "cpet_spo2_nadir_pct": 90,
                "cpet_rer_peak": 1.18,
                "cpet_hr_peak_bpm": 156,
            },
        },
        {
            "id": "E02",
            "label": "HIV assoziierte PAH",
            "scenario": "pah_pre",
            "modules": ["P14"],
            "story": "Beispiel E02: Präkapilläre PH mit Risikofaktor HIV. Differenzialblock soll Gruppe 1 (HIV) ausweisen.",
            "overrides": {
                "ph_known": True,
                "ph_known_dx": "PAH (Gruppe 1)",
                "ph_known_subtype": "HIV assoziierte PAH",
                "virology_pos": True,
                "virology_items": ["HIV"],
                "virology_desc": "HIV positiv.",
                "immunology_pos": False,
                "immunology_items": [],
                "immunology_desc": "",
                "vq_done": True,
                "vq_defect": False,
                "ct_embolie": False,
                "ct_mosaic": False,
                "ph_tx_table": _tx([
                    ["Opsumit (Macitentan)", "aktuell", "06/2024", "", "", ""],
                    ["Tadalafil", "aktuell", "06/2024", "", "", ""],
                ]),
                "consent_done": True,
                "access_route": "V. jugularis rechts",
                "cpet_done": True,
                "cpet_peak_vo2_ml_kg_min": 12.0,
                "cpet_peak_vo2_pct_pred": 49,
                "cpet_ve_vco2_slope": 45,
                "cpet_petco2_vt1_mmhg": 28,
                "cpet_spo2_nadir_pct": 92,
            },
        },
        {
            "id": "E03",
            "label": "CTEPH",
            "scenario": "cteph",
            "modules": ["P10"],
            "story": "Beispiel E03: CTEPH Konstellation mit V/Q Defekten und CT Mosaik. Antikoagulation aktiv. Therapie mit Adempas.",
            "overrides": {
                "ph_known": True,
                "ph_known_dx": "CTEPH (Gruppe 4)",
                "ph_known_subtype": "inoperable CTEPH (BPA Evaluation)",
                "vq_done": True,
                "vq_defect": True,
                "vq_desc": "Mehrsegmentale Perfusionsdefekte.",
                "ct_embolie": True,
                "ct_mosaic": True,
                "anticoag_status": "ja",
                "anticoag_indication": "CTEPH/CTEPD",
                "anticoag_substance": "DOAC (Apixaban, Rivaroxaban)",
                "ph_tx_table": _tx([
                    ["Adempas (Riociguat)", "aktuell", "08/2024", "", "", ""],
                ]),
                "consent_done": True,
                "access_route": "V. jugularis rechts",
            },
        },
        {
            "id": "E04",
            "label": "Gruppe 3 ILD Hypoxie",
            "scenario": "ild_ph",
            "modules": ["P12"],
            "story": "Beispiel E04: PH Verdacht bei ILD. LTOT aktiv, Lufu restriktiv und DLCO reduziert. Differenzialblock soll Gruppe 3 priorisieren.",
            "overrides": {
                "ph_known": False,
                "ph_suspected": True,
                "ct_done": True,
                "ct_ild": True,
                "ct_emphysema": False,
                "ltot": True,
                "ltot_flow_l_min": 2.0,
                "lufu_done": True,
                "lufu_restrictive": True,
                "lufu_diffusion": True,
                "dlco_sb": 42,
                "vq_done": False,
                "virology_pos": False,
                "immunology_pos": False,
                "consent_done": True,
                "access_route": "V. jugularis rechts",
                "cpet_done": True,
                "cpet_peak_vo2_ml_kg_min": 11.2,
                "cpet_peak_vo2_pct_pred": 45,
                "cpet_ve_vco2_slope": 38,
                "cpet_spo2_nadir_pct": 84,
                "cpet_o2_supp_l_min": 2.0,
            },
        },
        {
            "id": "E05",
            "label": "HFpEF iPcPH mit Volumen und Belastung",
            "scenario": "hfpef_ipcph",
            "modules": ["P09"],
            "story": "Beispiel E05: iPcPH HFpEF Konstellation. Volumenbelastung und Belastungshämodynamik aktiv.",
            "overrides": {
                "ph_known": True,
                "ph_known_dx": "PH bei Linksherzerkrankung / HFpEF (Gruppe 2)",
                "ph_known_subtype": "HFpEF mit postkapillärer PH",
                "exercise_done": True,
                "volume_challenge_done": True,
                "volume_ml": 500,
                "atrial_fib": True,
                "la_enlarged": True,
                "anticoag_status": "ja",
                "anticoag_indication": "Vorhofflimmern",
                "anticoag_substance": "DOAC (Apixaban, Rivaroxaban)",
                "consent_done": True,
                "access_route": "V. jugularis rechts",
            },
        },
        {
            "id": "E06",
            "label": "HFpEF cPcPH",
            "scenario": "cpcph",
            "modules": ["P09"],
            "story": "Beispiel E06: cPcPH Muster bei Linksherzerkrankung. Erhöhte PAWP und erhöhte PVR.",
            "overrides": {
                "ph_known": True,
                "ph_known_dx": "PH bei Linksherzerkrankung / HFpEF (Gruppe 2)",
                "ph_known_subtype": "cPcPH bei HFpEF",
                "atrial_fib": True,
                "la_enlarged": True,
                "consent_done": True,
                "access_route": "V. jugularis rechts",
            },
        },
        {
            "id": "E07",
            "label": "Shunt ASD",
            "scenario": "shunt_asd",
            "modules": ["P01"],
            "story": "Beispiel E07: Step up in der Stufenoxymetrie bei ASD. Testet Shunt Logik.",
            "overrides": {
                "chd_pos": True,
                "chd_type": "ASD (Vorhofseptumdefekt)",
                "chd_desc": "ASD V.a. bzw. bekannt.",
                "consent_done": True,
                "access_route": "V. jugularis rechts",
            },
        },
        {
            "id": "E08",
            "label": "Legacy PH Therapie Import",
            "scenario": "pah_pre",
            "modules": ["P14"],
            "story": "Beispiel E08: Legacy PH Therapie Felder gefüllt (Mehrfachlisten). Testet Button: Legacy Therapie in Episoden übernehmen.",
            "overrides": {
                "ph_known": True,
                "ph_known_dx": "PAH (Gruppe 1)",
                "ph_known_subtype": "Systemsklerose assoziierte PAH",
                "immunology_pos": True,
                "immunology_items": ["Systemische Sklerose (Sklerodermie)"],
                "immunology_desc": "Autoimmunerkrankung bekannt.",
                "ph_tx_table": "",
                "ph_current_meds": ["PDE‑5‑Hemmer", "Endothelin‑Rezeptorantagonist (ERA)"],
                "ph_prev_meds": ["Prostazyklin‑Therapie / -Analogon"],
                "ph_tx_status": "eskaliert",
                "ph_new_meds": ["Sotatercept (BMPR2/Activin-Pfad)"],
                "ph_stopped_meds": ["Prostazyklin‑Therapie / -Analogon"],
                "ph_stop_reason": "Unverträglichkeit/Nebenwirkung",
                "ph_stop_reason_text": "Beispiel: Flush und Hypotonie.",
                "consent_done": True,
                "access_route": "V. jugularis rechts",
            },
        },
    ]

    cfg = SUITE[idx % len(SUITE)]

    # Deterministische Basis aus random_example, dann gezielte Overrides
    ui = random_example(scenario=str(cfg.get("scenario") or ""), seed=10_000 + (idx % 10_000))

    # Sichtbarkeit/Orientierung in der UI
    ui["firstname"] = "Test"
    ui["name"] = f"{cfg.get('id')} {cfg.get('label')}"
    ui["story"] = str(cfg.get("story") or ui.get("story") or "")

    # Module
    mods = list(cfg.get("modules") or [])
    ui["modules"] = mods
    ui["modules_lvl1"] = []
    ui["modules_lvl2"] = []
    ui["modules_lvl3"] = mods

    # Standard: keine impliziten Angaben
    ui.setdefault("allergies_present", False)
    ui.setdefault("allergies_list", [])
    ui.setdefault("allergies_other_text", "")
    ui.setdefault("lsb_present", False)
    ui.setdefault("lsb_reason", "")
    ui.setdefault("anticoag_paused", False)

    # Apply overrides (last write wins)
    for k, v in (cfg.get("overrides") or {}).items():
        ui[k] = v

    # Konsistenz: bekannte PH impliziert kein Verdacht
    if bool(ui.get("ph_known")):
        ui["ph_suspected"] = False

    return ui





# =============================================================================
# JSON export/import helpers
# =============================================================================


# ---------------------------------------------------------------------------
# Performance: cache expensive clipboard conversions (example loads / repeated generate)
# ---------------------------------------------------------------------------

@lru_cache(maxsize=256)
def _markdown_to_plain_cached(s: str) -> str:
    return markdown_to_plain(s)

@lru_cache(maxsize=128)
def _markdown_to_word_html_cached(s: str) -> str:
    return markdown_to_word_html(s)

@lru_cache(maxsize=256)
def _extract_markdown_section_cached(md: str, start: str, end: str) -> str:
    return extract_markdown_section(md, start, end)


def markdown_to_plain(md: Any) -> str:
    """Best-effort Markdown -> plain text.

    Goal: copy/paste into Arztbrief systems without formatting artifacts.
    This is intentionally conservative and avoids clever formatting.
    """
    try:
        s = "" if md is None else str(md)
    except Exception:
        return ""

    # Normalize line endings
    s = s.replace("\r\n", "\n").replace("\r", "\n")

    # Remove code fences (keep content)
    s = re.sub(r"```[a-zA-Z0-9_-]*\n", "", s)
    s = s.replace("```", "")

    # Tables: replace pipes with tabs and strip separator rows
    lines: List[str] = []
    for ln in s.split("\n"):
        if re.match(r"^\s*\|?\s*[:-]+\s*\|", ln):
            continue
        if "|" in ln:
            ln = ln.strip().strip("|")
            ln = "\t".join([c.strip() for c in ln.split("|")])
        lines.append(ln)
    s = "\n".join(lines)

    # Headings: strip leading hashes
    s = re.sub(r"^\s{0,3}#{1,6}\s+", "", s, flags=re.M)

    # Bold/italic/underline markers
    s = s.replace("**", "").replace("__", "").replace("*", "").replace("_", "")

    # Defensive: remove any leftover placeholder artifacts that could appear
    # if a preprocessed string (e.g. clipboard conversion) is accidentally routed
    # through this function.
    s = re.sub(r"@@?BOPEN@@?", "", s)
    s = re.sub(r"@@?BCLOSE@@?", "", s)
    s = s.replace("BOPEN", "").replace("BCLOSE", "")

    # Links: [text](url) -> text
    s = re.sub(r"\[([^\]]+)\]\(([^\)]+)\)", r"\1", s)

    # Inline code
    s = s.replace("`", "")

    # Collapse extra spaces but keep intentional newlines
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()



def markdown_to_word_html(md: Any) -> str:
    """Best-effort Markdown -> HTML fragment suitable for pasting into MS Word.

    Notes:
    - Preserves headings, paragraphs, simple lists, and simple tables.
    - Avoids italics (no <em>) by stripping single * / _ emphasis.
    - Uses minimal inline styling to match Word defaults.

    Returns a full HTML document string. For clipboard usage, it includes
    <!--StartFragment--> / <!--EndFragment--> markers.
    """
    import html as _html

    try:
        s = "" if md is None else str(md)
    except Exception:
        s = ""

    s = s.replace("\r\n", "\n").replace("\r", "\n")

    # Remove code fences (keep content)
    s = re.sub(r"```[a-zA-Z0-9_-]*\n", "", s)
    s = s.replace("```", "")

    # Inline helpers
    def _inline(x: str) -> str:
        x = "" if x is None else str(x)

        # Links: [text](url) -> text
        x = re.sub(r"\[([^\]]+)\]\(([^\)]+)\)", r"\1", x)

        # Bold: **text** or __text__
        BOPEN = "@@BOPEN@@"
        BCLOSE = "@@BCLOSE@@"
        x = re.sub(r"\*\*(.+?)\*\*", lambda m: f"{BOPEN}{m.group(1)}{BCLOSE}", x)
        x = re.sub(r"__(.+?)__", lambda m: f"{BOPEN}{m.group(1)}{BCLOSE}", x)

        # Italics: *text* or _text_ (single markers only)
        x = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", lambda m: m.group(1), x)
        x = re.sub(r"(?<!_)_(?!_)(.+?)(?<!_)_(?!_)", lambda m: m.group(1), x)

        # Inline code: `x` -> x
        x = x.replace("`", "")

        # Escape HTML
        x = _html.escape(x, quote=False)

        # Restore bold placeholders
        x = x.replace(BOPEN, "<strong>").replace(BCLOSE, "</strong>")
        return x

    lines = s.split("\n")

    out = []
    out.append("<html><body>")
    out.append("<!--StartFragment-->")
    # Clipboard (Word): Arial, 10pt body. Headings are handled explicitly (11pt bold).
    out.append("<div style=\"font-family:Arial,sans-serif;font-size:10pt;line-height:1.25;\">")

    i = 0
    in_ul = False
    in_ol = False
    in_table = False

    def _close_lists():
        nonlocal in_ul, in_ol
        if in_ul:
            out.append("</ul>")
            in_ul = False
        if in_ol:
            out.append("</ol>")
            in_ol = False

    def _close_table():
        nonlocal in_table
        if in_table:
            out.append("</table>")
            in_table = False

    # Paragraph buffer
    para: list[str] = []

    def _flush_para():
        nonlocal para
        if not para:
            return
        _close_lists()
        _close_table()
        txt = _inline(" ".join([p.strip() for p in para if p.strip()]))
        if txt:
            out.append(f"<p style=\"margin:0 0 6pt 0;\">{txt}</p>")
        para = []

    def _is_table_sep(ln: str) -> bool:
        # e.g. |---|:---:|
        return bool(re.match(r"^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$", ln))

    while i < len(lines):
        ln = lines[i]
        raw_ln = ln
        ln = ln.rstrip("\n")
        stripped = ln.strip()

        # Blank line flushes paragraph
        if stripped == "":
            _flush_para()
            i += 1
            continue

        # Headings
        hm = re.match(r"^\s{0,3}(#{1,6})\s+(.*)$", ln)
        if hm:
            _flush_para()
            _close_lists()
            _close_table()
            level = min(len(hm.group(1)), 4)
            heading_raw = hm.group(2).strip()
            heading_text = _inline(heading_raw)

            # For some sections, Word looks much better if we collapse key-value bullet lists into a compact flow text.
            def _hkey(x: str) -> str:
                x = (x or "").strip().lower()
                x = x.replace(" / ", "/").replace(" /", "/").replace("/ ", "/")
                x = re.sub(r"\s+", " ", x)
                x = re.sub(r"[:：]+$", "", x)
                return x

            flow_keys = {
                "klinik",
                "befundübersicht",
                "stufenoxymetrie",
                "bildgebung/echo/cmr",
                "bildgebung",
            }

            if _hkey(heading_raw) in flow_keys:
                j = i + 1
                while j < len(lines) and lines[j].strip() == "":
                    j += 1
                items = []
                while j < len(lines):
                    ln2 = lines[j].rstrip("\n")
                    # Stop at next heading
                    if re.match(r"^\s{0,3}(#{1,6})\s+(.*)$", ln2):
                        break
                    m_ul2 = re.match(r"^\s*[-•\*]\s+(.*)$", ln2)
                    if m_ul2:
                        items.append(m_ul2.group(1).strip())
                        j += 1
                        continue
                    break

                if items:
                    joined = "; ".join([_inline(it) for it in items if it.strip()])
                    label = heading_text
                    if not re.search(r"[:：]\s*$", label):
                        label = label + ":"
                    out.append(f"<p style=\"margin:0 0 6pt 0;\"><strong>{label}</strong> {joined}</p>")
                    i = j
                    continue

            # Word copy requirements: headings = Arial 11pt bold, rest = Arial 10pt.
            # Use <p> instead of <h*> to avoid Word re-styling.
            out.append(
                f"<p style=\"margin:8pt 0 4pt 0;font-family:Arial,sans-serif;font-size:11pt;font-weight:700;\">{heading_text}</p>"
            )
            i += 1
            continue

                # Tables (pipe tables)
        if "|" in stripped and stripped.count("|") >= 2:
            # detect contiguous table block
            # start only if next line is separator OR looks like table row and we are already in table
            nxt = lines[i+1].strip() if i + 1 < len(lines) else ""
            if in_table or _is_table_sep(nxt):
                _flush_para()
                _close_lists()
                if not in_table:
                    out.append("<table border=\"1\" cellspacing=\"0\" cellpadding=\"4\" style=\"border-collapse:collapse;\">")
                    in_table = True
                # Skip separator rows
                if _is_table_sep(stripped):
                    i += 1
                    continue
                row = stripped.strip("|")
                cells = [c.strip() for c in row.split("|")]
                # Header heuristic: if next line is separator and we're at start of table
                is_header = False
                if i + 1 < len(lines) and _is_table_sep(lines[i+1].strip()):
                    # This line is header
                    is_header = True
                tag = "th" if is_header else "td"
                out.append("<tr>" + "".join([f"<{tag}>{_inline(c)}</{tag}>" for c in cells]) + "</tr>")
                i += 1
                continue

        # Unordered list
        m_ul = re.match(r"^\s*[-•\*]\s+(.*)$", ln)
        if m_ul:
            _flush_para()
            _close_table()
            if in_ol:
                out.append("</ol>")
                in_ol = False
            if not in_ul:
                out.append("<ul style=\"margin:0 0 6pt 18pt;padding:0;\">")
                in_ul = True
            out.append(f"<li style=\"margin:0;\">{_inline(m_ul.group(1).strip())}</li>")
            i += 1
            continue

        # Ordered list
        m_ol = re.match(r"^\s*\d+\.\s+(.*)$", ln)
        if m_ol:
            _flush_para()
            _close_table()
            if in_ul:
                out.append("</ul>")
                in_ul = False
            if not in_ol:
                out.append("<ol style=\"margin:0 0 6pt 18pt;padding:0;\">")
                in_ol = True
            out.append(f"<li style=\"margin:0;\">{_inline(m_ol.group(1).strip())}</li>")
            i += 1
            continue

        # Default: paragraph line
        para.append(raw_ln)
        i += 1

    _flush_para()
    _close_lists()
    _close_table()

    out.append("</div>")
    out.append("<!--EndFragment-->")
    out.append("</body></html>")
    html = "\n".join(out)
    # Defensive: ensure no placeholder artifacts leak into the clipboard payload
    html = html.replace("@@BOPEN@@", "").replace("@@BCLOSE@@", "")
    # In the worst case, strip bare marker words too (should not happen)
    html = html.replace("BOPEN", "").replace("BCLOSE", "")
    return html


def markdown_to_docx_file(md: Any, out_path: str) -> str:
    """Best-effort Markdown -> DOCX.

    Scope (intentionally small and stable)
    - Paragraphs
    - Simple bullet lists (- / •)
    - Bold spans (**x**)
    - Section headers as bold lines ("**Titel:**")
    - Headings (## / ###)
    - Page breaks via a dedicated marker line: [[PAGEBREAK]]

    This is used for the "DOCX" download button (copy layout). The in-app
    report remains unchanged.
    """
    from docx import Document
    from docx.shared import Pt, Cm

    try:
        s = "" if md is None else str(md)
    except Exception:
        s = ""
    s = s.replace("\r\n", "\n").replace("\r", "\n")

    doc = Document()

    # Page setup: keep Word-default look but enforce predictable margins.
    try:
        for sec in doc.sections:
            sec.top_margin = Cm(2.0)
            sec.bottom_margin = Cm(2.0)
            sec.left_margin = Cm(2.0)
            sec.right_margin = Cm(2.0)
    except Exception:
        pass

    # Copy requirements: Arial 10pt body, headings Arial 11pt bold.
    style = doc.styles["Normal"]
    try:
        style.font.name = "Arial"
        style.font.size = Pt(10)
    except Exception:
        pass

    def _set_run_font(run, size_pt: int):
        try:
            run.font.name = "Arial"
            run.font.size = Pt(size_pt)
        except Exception:
            pass

    bold_pat = re.compile(r"\*\*(.+?)\*\*")

    def _add_runs_with_bold(par, text: str):
        """Add runs to paragraph, turning **x** into bold runs."""
        if not text:
            return
        pos = 0
        for m in bold_pat.finditer(text):
            if m.start() > pos:
                r0 = par.add_run(text[pos:m.start()])
                _set_run_font(r0, 10)
            r = par.add_run(m.group(1))
            r.bold = True
            _set_run_font(r, 10)
            pos = m.end()
        if pos < len(text):
            r1 = par.add_run(text[pos:])
            _set_run_font(r1, 10)

    # Build paragraphs
    lines = [ln.rstrip() for ln in s.split("\n")]
    prev_blank = True
    for raw in lines:
        ln = (raw or "").rstrip()
        if not ln.strip():
            prev_blank = True
            continue

        # Explicit page break marker
        if ln.strip() == "[[PAGEBREAK]]":
            try:
                doc.add_page_break()
            except Exception:
                # Fallback: add spacing
                doc.add_paragraph()
            prev_blank = True
            continue

        # Headings (Word copy layout): bold, Arial 11pt
        m_h2 = re.match(r"^\s*##\s+(.+)$", ln)
        if m_h2:
            par = doc.add_paragraph()
            r = par.add_run(m_h2.group(1).strip())
            r.bold = True
            _set_run_font(r, 11)
            prev_blank = True
            continue
        m_h3 = re.match(r"^\s*###\s+(.+)$", ln)
        if m_h3:
            par = doc.add_paragraph()
            r = par.add_run(m_h3.group(1).strip())
            r.bold = True
            _set_run_font(r, 11)
            prev_blank = True
            continue

        # Bullet list (supports nesting via leading spaces: '  - item', '    - item', ...)
        m_b = re.match(r"^(\s*)(?:[-•]\s+)(.+)$", ln)
        if m_b:
            indent = len(m_b.group(1) or "")
            # Common Markdown convention: 2 spaces per nesting level (cap at 2 for stability)
            lvl = 0
            if indent >= 2:
                lvl = 1
            if indent >= 4:
                lvl = 2
            style_name = "List Bullet"
            if lvl == 1:
                style_name = "List Bullet 2"
            elif lvl == 2:
                style_name = "List Bullet 3"
            try:
                par = doc.add_paragraph(style=style_name)
            except Exception:
                par = doc.add_paragraph(style="List Bullet")
            _add_runs_with_bold(par, m_b.group(2).strip())
            prev_blank = False
            continue

        # Section header line: **Titel:**
        m_h = re.match(r"^\s*\*\*(.+?)\*\*\s*$", ln)
        if m_h and (m_h.group(1).endswith(":") or m_h.group(1).endswith(" :")):
            par = doc.add_paragraph()
            r = par.add_run(m_h.group(1).strip())
            r.bold = True
            try:
                par.paragraph_format.space_after = Pt(6)
            except Exception:
                pass
            prev_blank = False
            continue

        # Normal paragraph (start new paragraph after blank lines)
        if prev_blank:
            par = doc.add_paragraph()
            _add_runs_with_bold(par, ln.strip())
        else:
            # Continue previous paragraph (Word-like flow)
            par = doc.paragraphs[-1]
            par.add_run(" ")
            _add_runs_with_bold(par, ln.strip())
        prev_blank = False

    doc.save(out_path)
    return out_path


def extract_markdown_section(md: Any, start_heading: str, end_heading: Optional[str] = None) -> str:
    """Extract a section from markdown by headings (best-effort).

    Returns the substring starting at the first occurrence of start_heading
    (as a Markdown heading line) until end_heading (exclusive) if provided.
    """
    try:
        s = "" if md is None else str(md)
    except Exception:
        return ""
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    # Match headings like "## Rechtsherzkatheter"
    start_pat = re.compile(rf"^\s*#+\s*{re.escape(start_heading)}\s*$", re.M)
    m = start_pat.search(s)
    if not m:
        # fallback: plain substring search
        idx = s.find(start_heading)
        if idx < 0:
            return s
        s2 = s[idx:]
        if end_heading and end_heading in s2:
            return s2.split(end_heading, 1)[0]
        return s2

    start_idx = m.start()
    s2 = s[start_idx:]
    if end_heading:
        end_pat = re.compile(rf"^\s*#+\s*{re.escape(end_heading)}\s*$", re.M)
        m2 = end_pat.search(s2)
        if m2:
            return s2[:m2.start()].strip()
        # fallback substring
        if end_heading in s2:
            return s2.split(end_heading, 1)[0].strip()
    return s2.strip()


def build_summary_dict(case: Dict[str, Any], rulebook_meta: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    """Structured, stable JSON summary for studies/registries/QA."""
    fp = _case_fingerprint(case)
    cached = _cache_get('summary_dict', fp)
    if cached is not None:
        return cached
    ui = case.get("ui") or {}
    der = case.get("derived") or {}
    scores = case.get("scores") or {}
    dec = case.get("decision") or {}
    warns = case.get("warnings") or []

    def _num(x: Any) -> Optional[float]:
        try:
            if x is None or x == "":
                return None
            return float(x)
        except Exception:
            return None

    # Slim warnings (message + severity + code if present)
    wslim: List[Dict[str, Any]] = []
    if isinstance(warns, list):
        for w in warns:
            if not isinstance(w, dict):
                continue
            wslim.append(
                {
                    "severity": w.get("severity"),
                    "code": w.get("code"),
                    "message": w.get("message"),
                }
            )

    rb = rulebook_meta or {}
    rb_meta = {
        "version": (rb.get("version") if isinstance(rb, dict) else None),
        "updated": (rb.get("updated") if isinstance(rb, dict) else None),
    }

    # Dates
    today = _dt.datetime.now().isoformat(timespec="seconds")

    # Core hemodynamics
    hemo = {
        "rap_rest_mmHg": _num(der.get("rap_rest")),
        "spap_rest_mmHg": _num(der.get("spap_rest")),
        "dpap_rest_mmHg": _num(der.get("dpap_rest")),
        "mpap_rest_mmHg": _num(der.get("mpap_rest")),
        "pawp_rest_mmHg": _num(der.get("pawp_rest")),
        "co_rest_L_min": _num(der.get("co")),
        "ci_rest_L_min_m2": _num(der.get("ci")),
        "pvr_rest_WU": _num(der.get("pvr_rest")),
        "pvri_rest_WU_m2": _num(der.get("pvri")),
        "tpg_mmHg": _num(der.get("tpg")),
        "dpg_mmHg": _num(der.get("dpg")),
    }

    # Classification / risk
    classification = {
        "hemo_category": der.get("hemo_category"),
        "primary_dx": dec.get("primary_dx"),
        "bundle": dec.get("bundle"),
        "risk_category": der.get("risk_category"),
        "esc_ers_4s": scores.get("esc_ers_4s"),
        "esc_ers_3s": scores.get("esc_ers_3s"),
        "reveal_lite2": scores.get("reveal_lite2"),
        "reveal_lite2_points": scores.get("reveal_lite2_points"),
    }

    # Echo snapshot (only the main fields used in patient echo report)
    echo = {
        "lvef_percent": _num(ui.get("lvef")),
        "tapse_mm": _num(ui.get("tapse_mm")),
        "s_prime_cm_s": _num(ui.get("s_prime_cm_s")),
        "pasp_echo_mmHg": _num(ui.get("pasp_echo")),
        "ra_esa_cm2": _num(ui.get("ra_esa_cm2")),
        "ee_ratio": _num(ui.get("ee_ratio")),
        "trv_ms": _num(ui.get("trv_ms")),
    }

    labs = {
        "hb_g_dl": _num(ui.get("hb_g_dl")),
        "crp_mg_l": _num(ui.get("crp_mg_l")),
        "creatinine_mg_dl": _num(ui.get("creatinine_mg_dl")),
        "egfr_ml_min_1_73m2": _num(ui.get("egfr")),
        "bnp_kind": ui.get("bnp_kind"),
        "bnp_value_pg_ml": _num(ui.get("bnp_value")),
    }

    patient = {
        "firstname": ui.get("firstname"),
        "name": ui.get("name"),
        "age_years": _num(ui.get("age")),
        "sex": ui.get("sex"),
        "height_cm": _num(ui.get("height_cm")),
        "weight_kg": _num(ui.get("weight_kg")),
    }

    context = {
        "who_fc": ui.get("who_fc"),
        "six_mwd_m": _num(ui.get("six_mwd_m")),
        "story": ui.get("story"),
        "ph_known": ui.get("ph_known"),
        "ph_suspected": ui.get("ph_suspected"),
        "ph_known_dx": ui.get("ph_known_dx"),
    }

    procedere = {
        "modules_selected": ui.get("modules") or [],
        "procedere_free": ui.get("procedere_free") or "",
    }

    _res = {
        "schema": "rhk_summary_v1",
        "generated_at": today,
        "app_version": APP_VERSION,
        "rulebook": rb_meta,
        "patient": patient,
        "context": context,
        "hemodynamics": hemo,
        "classification": classification,
        "echo": echo,
        "labs": labs,
        "procedere": procedere,
        "warnings": wslim,
    }
    _cache_set('summary_dict', fp, _res)
    return _res


def export_json(case: Dict[str, Any], path: str) -> str:
    with open(path, "w", encoding="utf-8") as f:
        json.dump(case, f, ensure_ascii=False, indent=2)
    return path


def export_summary_json(summary: Dict[str, Any], path: str) -> str:
    with open(path, "w", encoding="utf-8") as f:
        json.dump(summary, f, ensure_ascii=False, indent=2)
    return path


def load_case_json(file_path: str) -> Dict[str, Any]:
    with open(file_path, "r", encoding="utf-8") as f:
        return json.load(f)



def build_echo_doctor_report_extended(case: Dict[str, Any]) -> str:
    """Arztbericht Echokardiographie – strukturiert, PH-bezogen, Rechtsherz-Fokus.

    Implementierung liegt in `rhk_echo_report_doctor.py` und wird hier nur
    gecached/wrapped, um etablierte Schnittstellen stabil zu halten.
    """
    fp = _case_fingerprint(case)
    cached = _cache_get('echo_doctor_report', fp)
    if cached is not None:
        return cached

    from rhk_echo_report_doctor import build_echo_doctor_report as _impl
    out = _impl(case)

    _cache_set('echo_doctor_report', fp, out)
    return out

