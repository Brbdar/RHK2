#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Markdown/plain/html/docx conversion helpers for report exports."""

from __future__ import annotations

import re
from typing import Any, List, Optional

from rhk_logging import log_exception

_REPORT_RECOVERABLE_ERRORS = (
    AttributeError,
    KeyError,
    OSError,
    TypeError,
    ValueError,
    ImportError,
    ModuleNotFoundError,
)


def _markdown_to_plain_cached(s: str) -> str:
    # Intentionally NOT cached: input contains patient-identifiable information.
    return markdown_to_plain(s)


def _markdown_to_word_html_cached(s: str) -> str:
    # Intentionally NOT cached: input contains patient-identifiable information.
    return markdown_to_word_html(s)


def _extract_markdown_section_cached(md: str, start: str, end: str) -> str:
    # Intentionally NOT cached: input contains patient-identifiable information.
    return extract_markdown_section(md, start, end)


def markdown_to_plain(md: Any) -> str:
    """Best-effort Markdown -> plain text."""
    try:
        s = "" if md is None else str(md)
    except (TypeError, ValueError) as exc:
        log_exception(
            "RHK_REP_MD_TO_PLAIN_INPUT",
            "Markdown to plain conversion failed at input coercion.",
            exc,
        )
        return ""

    s = s.replace("\r\n", "\n").replace("\r", "\n")
    s = re.sub(r"```[a-zA-Z0-9_-]*\n", "", s)
    s = s.replace("```", "")

    lines: List[str] = []
    for ln in s.split("\n"):
        if re.match(r"^\s*\|?\s*[:-]+\s*\|", ln):
            continue
        if "|" in ln:
            ln = ln.strip().strip("|")
            ln = "\t".join([c.strip() for c in ln.split("|")])
        lines.append(ln)
    s = "\n".join(lines)

    s = re.sub(r"^\s{0,3}#{1,6}\s+", "", s, flags=re.M)
    s = s.replace("**", "").replace("__", "").replace("*", "").replace("_", "")
    s = re.sub(r"@@?BOPEN@@?", "", s)
    s = re.sub(r"@@?BCLOSE@@?", "", s)
    s = s.replace("BOPEN", "").replace("BCLOSE", "")
    s = re.sub(r"\[([^\]]+)\]\(([^\)]+)\)", r"\1", s)
    s = s.replace("`", "")
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()


def markdown_to_word_html(md: Any) -> str:
    """Markdown -> HTML fragment optimized for pasting into MS Word."""
    import html as _html

    try:
        s = "" if md is None else str(md)
    except (TypeError, ValueError) as exc:
        log_exception(
            "RHK_REP_MD_TO_HTML_INPUT",
            "Markdown to word-html conversion failed at input coercion.",
            exc,
        )
        s = ""

    s = s.replace("\r\n", "\n").replace("\r", "\n")
    s = re.sub(r"```[a-zA-Z0-9_-]*\n", "", s)
    s = s.replace("```", "")

    def _inline(x: str) -> str:
        x = "" if x is None else str(x)
        x = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1", x)

        bopen = "@@BOPEN@@"
        bclose = "@@BCLOSE@@"
        x = re.sub(r"\*\*(.+?)\*\*", lambda m: f"{bopen}{m.group(1)}{bclose}", x)
        x = re.sub(r"__(.+?)__", lambda m: f"{bopen}{m.group(1)}{bclose}", x)
        x = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", lambda m: m.group(1), x)
        x = re.sub(r"(?<!_)_(?!_)(.+?)(?<!_)_(?!_)", lambda m: m.group(1), x)
        x = x.replace("`", "")
        x = _html.escape(x, quote=False)
        x = x.replace(bopen, "<strong>").replace(bclose, "</strong>")
        return x

    lines = s.split("\n")
    out: List[str] = []
    out.append("<html><body>")
    out.append("<!--StartFragment-->")
    out.append("<div style=\"font-family:Calibri,Arial,sans-serif;font-size:10.5pt;line-height:1.25;\">")

    i = 0
    current_bullets: List[str] = []
    seen_any_heading = False
    para_buf: List[str] = []

    def _emit_heading(txt: str, _level: int) -> None:
        t = _inline(txt.strip())
        if t:
            out.append(f"<div style=\"font-size:11pt;font-weight:700;margin:10pt 0 4pt 0;\">{t}</div>")

    def _emit_bullets(items: List[str]) -> None:
        clean = [x for x in (items or []) if str(x or "").strip()]
        if not clean:
            return
        out.append("<div style=\"margin-left:14pt;margin-bottom:6pt;\">")
        out.append("<ul style=\"margin:0;padding-left:16pt;\">")
        for it in clean:
            out.append(f"<li>{_inline(it.strip())}</li>")
        out.append("</ul>")
        out.append("</div>")

    def _emit_paragraph(txt: str) -> None:
        t = _inline(txt.strip())
        if t:
            out.append(f"<div style=\"margin:0 0 6pt 0;\">{t}</div>")

    def _heading_level(line: str):
        m = re.match(r"^(#{1,6})\s+(.*)$", line.strip())
        if not m:
            return None
        return len(m.group(1)), m.group(2).strip()

    def _is_table_sep(line: str):
        return re.match(r"^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$", line) is not None

    def _parse_table(start_idx: int):
        if start_idx >= len(lines):
            return None, start_idx
        header = lines[start_idx]
        if "|" not in header or start_idx + 1 >= len(lines):
            return None, start_idx
        if not _is_table_sep(lines[start_idx + 1]):
            return None, start_idx

        def split_row(row: str) -> List[str]:
            row = row.strip()
            if row.startswith("|"):
                row = row[1:]
            if row.endswith("|"):
                row = row[:-1]
            return [c.strip() for c in row.split("|")]

        head_cells = split_row(header)
        rows: List[List[str]] = []
        j = start_idx + 2
        while j < len(lines) and "|" in lines[j] and lines[j].strip():
            rows.append(split_row(lines[j]))
            j += 1

        html_parts: List[str] = []
        html_parts.append("<table style=\"border-collapse:collapse;margin-left:14pt;margin-bottom:6pt;\">")
        html_parts.append("<tr>")
        for c in head_cells:
            html_parts.append(
                f"<th style=\"border:1px solid #999;padding:3pt 6pt;font-size:10.5pt;font-weight:700;\">{_inline(c)}</th>"
            )
        html_parts.append("</tr>")
        for row in rows:
            html_parts.append("<tr>")
            for c in row:
                html_parts.append(
                    f"<td style=\"border:1px solid #999;padding:3pt 6pt;font-size:10.5pt;\">{_inline(c)}</td>"
                )
            html_parts.append("</tr>")
        html_parts.append("</table>")
        return "".join(html_parts), j

    def _flush_current_bullets() -> None:
        nonlocal current_bullets
        _emit_bullets(current_bullets)
        current_bullets = []

    def _flush_para_into_bullets() -> None:
        nonlocal para_buf, current_bullets
        if not para_buf:
            return
        txt = " ".join([t.strip() for t in para_buf if t.strip()])
        para_buf = []
        if txt:
            current_bullets.append(txt)

    while i < len(lines):
        line = lines[i]

        tbl_html, next_i = _parse_table(i)
        if tbl_html is not None:
            _flush_para_into_bullets()
            if seen_any_heading:
                _flush_current_bullets()
            out.append(tbl_html)
            i = next_i
            continue

        hl = _heading_level(line)
        if hl is not None:
            _flush_para_into_bullets()
            if seen_any_heading:
                _flush_current_bullets()
            level, htxt = hl
            _emit_heading(htxt, level)
            seen_any_heading = True
            i += 1
            continue

        if not line.strip():
            _flush_para_into_bullets()
            i += 1
            continue

        m_ul = re.match(r"^\s*([-*•])\s+(.+)$", line)
        if m_ul:
            _flush_para_into_bullets()
            current_bullets.append(m_ul.group(2).strip())
            i += 1
            continue

        m_ol = re.match(r"^\s*(\d+)[\.\)]\s+(.+)$", line)
        if m_ol:
            _flush_para_into_bullets()
            current_bullets.append(f"{m_ol.group(1)}. {m_ol.group(2).strip()}")
            i += 1
            continue

        if seen_any_heading:
            para_buf.append(line)
        else:
            _emit_paragraph(line)
        i += 1

    _flush_para_into_bullets()
    if seen_any_heading:
        _flush_current_bullets()

    out.append("</div>")
    out.append("<!--EndFragment-->")
    out.append("</body></html>")
    return "\n".join(out)


def markdown_to_docx_file(md: Any, out_path: str) -> str:
    """Best-effort Markdown -> DOCX."""
    from docx import Document
    from docx.shared import Cm, Pt

    try:
        s = "" if md is None else str(md)
    except (TypeError, ValueError) as exc:
        log_exception(
            "RHK_REP_MD_TO_DOCX_INPUT",
            "Markdown to DOCX conversion failed at input coercion.",
            exc,
        )
        s = ""
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    doc = Document()

    try:
        for sec in doc.sections:
            sec.top_margin = Cm(2.0)
            sec.bottom_margin = Cm(2.0)
            sec.left_margin = Cm(2.0)
            sec.right_margin = Cm(2.0)
    except _REPORT_RECOVERABLE_ERRORS as exc:
        log_exception("RHK_REP_DOCX_MARGINS", "DOCX margin setup failed; continuing with defaults.", exc)

    style = doc.styles["Normal"]
    try:
        style.font.name = "Arial"
        style.font.size = Pt(10)
    except _REPORT_RECOVERABLE_ERRORS as exc:
        log_exception("RHK_REP_DOCX_STYLE", "DOCX default style setup failed; continuing.", exc)

    def _set_run_font(run, size_pt: int) -> None:
        try:
            run.font.name = "Arial"
            run.font.size = Pt(size_pt)
        except _REPORT_RECOVERABLE_ERRORS as exc:
            log_exception(
                "RHK_REP_DOCX_RUN_FONT",
                "DOCX run font setup failed for one run.",
                exc,
                size_pt=size_pt,
            )

    bold_pat = re.compile(r"\*\*(.+?)\*\*")

    def _add_runs_with_bold(par, text: str) -> None:
        if not text:
            return
        pos = 0
        for m in bold_pat.finditer(text):
            if m.start() > pos:
                r0 = par.add_run(text[pos : m.start()])
                _set_run_font(r0, 10)
            r = par.add_run(m.group(1))
            r.bold = True
            _set_run_font(r, 10)
            pos = m.end()
        if pos < len(text):
            r1 = par.add_run(text[pos:])
            _set_run_font(r1, 10)

    lines = [ln.rstrip() for ln in s.split("\n")]
    prev_blank = True
    for raw in lines:
        ln = (raw or "").rstrip()
        if not ln.strip():
            prev_blank = True
            continue

        if ln.strip() == "[[PAGEBREAK]]":
            try:
                doc.add_page_break()
            except _REPORT_RECOVERABLE_ERRORS as exc:
                log_exception(
                    "RHK_REP_DOCX_PAGEBREAK",
                    "DOCX pagebreak insertion failed; falling back to blank paragraph.",
                    exc,
                )
                doc.add_paragraph()
            prev_blank = True
            continue

        m_h2 = re.match(r"^\s*##\s+(.+)$", ln)
        if m_h2:
            par = doc.add_paragraph()
            r = par.add_run(m_h2.group(1).strip())
            r.bold = True
            _set_run_font(r, 11)
            prev_blank = True
            continue

        m_h3 = re.match(r"^\s*###\s+(.+)$", ln)
        if m_h3:
            par = doc.add_paragraph()
            r = par.add_run(m_h3.group(1).strip())
            r.bold = True
            _set_run_font(r, 11)
            prev_blank = True
            continue

        m_b = re.match(r"^(\s*)(?:[-•]\s+)(.+)$", ln)
        if m_b:
            indent = len(m_b.group(1) or "")
            lvl = 0
            if indent >= 2:
                lvl = 1
            if indent >= 4:
                lvl = 2
            style_name = "List Bullet"
            if lvl == 1:
                style_name = "List Bullet 2"
            elif lvl == 2:
                style_name = "List Bullet 3"
            try:
                par = doc.add_paragraph(style=style_name)
            except _REPORT_RECOVERABLE_ERRORS as exc:
                log_exception(
                    "RHK_REP_DOCX_BULLET_STYLE",
                    "DOCX bullet style unavailable; falling back to default bullet style.",
                    exc,
                    style_name=style_name,
                )
                par = doc.add_paragraph(style="List Bullet")
            _add_runs_with_bold(par, m_b.group(2).strip())
            prev_blank = False
            continue

        m_h = re.match(r"^\s*\*\*(.+?)\*\*\s*$", ln)
        if m_h and (m_h.group(1).endswith(":") or m_h.group(1).endswith(" :")):
            par = doc.add_paragraph()
            r = par.add_run(m_h.group(1).strip())
            r.bold = True
            try:
                par.paragraph_format.space_after = Pt(6)
            except _REPORT_RECOVERABLE_ERRORS as exc:
                log_exception("RHK_REP_DOCX_HEADER_SPACING", "DOCX header spacing setup failed; continuing.", exc)
            prev_blank = False
            continue

        if prev_blank:
            par = doc.add_paragraph()
            _add_runs_with_bold(par, ln.strip())
        else:
            par = doc.paragraphs[-1]
            par.add_run(" ")
            _add_runs_with_bold(par, ln.strip())
        prev_blank = False

    doc.save(out_path)
    return out_path


def extract_markdown_section(md: Any, start_heading: str, end_heading: Optional[str] = None) -> str:
    """Extract a section from markdown by headings (best-effort)."""
    try:
        s = "" if md is None else str(md)
    except (TypeError, ValueError) as exc:
        log_exception(
            "RHK_REP_SECTION_EXTRACT_INPUT",
            "Markdown section extraction failed at input coercion.",
            exc,
        )
        return ""
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    start_pat = re.compile(rf"^\s*#+\s*{re.escape(start_heading)}\s*$", re.M)
    m = start_pat.search(s)
    if not m:
        idx = s.find(start_heading)
        if idx < 0:
            return s
        s2 = s[idx:]
        if end_heading and end_heading in s2:
            return s2.split(end_heading, 1)[0]
        return s2

    s2 = s[m.start() :]
    if end_heading:
        end_pat = re.compile(rf"^\s*#+\s*{re.escape(end_heading)}\s*$", re.M)
        m2 = end_pat.search(s2)
        if m2:
            return s2[: m2.start()].strip()
        if end_heading in s2:
            return s2.split(end_heading, 1)[0].strip()
    return s2.strip()


__all__ = [
    "_extract_markdown_section_cached",
    "_markdown_to_plain_cached",
    "_markdown_to_word_html_cached",
    "extract_markdown_section",
    "markdown_to_docx_file",
    "markdown_to_plain",
    "markdown_to_word_html",
]
